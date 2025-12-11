"""
Страница анализа медицинских видео
Вынесена из app.py для улучшения архитектуры проекта
"""
import streamlit as st
import io
import datetime
import time
import sys
import traceback

# Импорты из claude_assistant
try:
    from claude_assistant import OpenRouterAssistant
    AI_AVAILABLE = True
except ImportError:
    AI_AVAILABLE = False
    OpenRouterAssistant = None


def show_video_analysis():
    """Страница анализа медицинских видео"""
    if not AI_AVAILABLE:
        st.error("❌ ИИ-модуль недоступен. Проверьте файл `claude_assistant.py` и API-ключ.")
        return
    
    st.header("🎬 Анализ медицинских видео")
    st.info("💡 Загрузите видео медицинской процедуры, функционального теста или динамического исследования для анализа через Gemini 2.5 Flash")
    
    # Выбор типа исследования
    study_type = st.selectbox(
        "Тип исследования:",
        ["", "fgds", "colonoscopy", "echo", "abdominal_us", "gynecology_us", "mri_brain", "mri_universal", "chest_ct"],
        format_func=lambda x: {
            "": "Выберите тип исследования",
            "fgds": "🔬 ФГДС (эзофагогастродуоденоскопия)",
            "colonoscopy": "🔬 Колоноскопия",
            "echo": "🫀 ЭхоКГ (эхокардиография)",
            "abdominal_us": "🔍 УЗИ органов брюшной полости",
            "gynecology_us": "🩺 Гинекологическое УЗИ",
            "mri_brain": "🧠 МРТ головного мозга",
            "mri_universal": "🧲 МРТ (универсальный)",
            "chest_ct": "🫁 КТ органов грудной клетки"
        }.get(x, x),
        help="Выберите тип исследования для использования специализированного промпта"
    )
    
    # Загрузка видео
    uploaded_video = st.file_uploader(
        "Загрузите видео-файл",
        type=["mp4", "mov", "avi", "webm", "mkv"],
        help="Поддерживаются форматы: MP4, MOV, AVI, WebM, MKV (максимум 100MB)"
    )
    
    if uploaded_video:
        # Показываем превью видео
        st.subheader("📹 Превью видео")
        st.video(uploaded_video)
        
        # Метаданные (опционально)
        st.subheader("📋 Метаданные (опционально)")
        col1, col2 = st.columns(2)
        
        with col1:
            patient_age = st.number_input("Возраст пациента", min_value=0, max_value=150, value=None, help="Укажите возраст для более точного анализа")
            specialty = st.selectbox(
                "Специализация",
                ["", "Терапия", "Хирургия", "Ортопедия", "Неврология", "Кардиология", "Педиатрия", "Онкология", "Другое"],
                help="Выберите специализацию для контекста анализа"
            )
        
        with col2:
            urgency = st.selectbox(
                "Срочность",
                ["", "Плановая", "Срочная", "Критическая"],
                help="Укажите уровень срочности"
            )
        
        # Дополнительный контекст (особенно для КТ ОГК)
        additional_context = ""
        if study_type == "chest_ct":
            st.subheader("📋 Дополнительные параметры для КТ ОГК")
            col_ct1, col_ct2, col_ct3 = st.columns(3)
            with col_ct1:
                ct_type = st.selectbox("Тип КТ", ["Нативное", "С контрастом", "КТЛА", "ВРКТ"])
            with col_ct2:
                clinical = st.text_input("Клиника", placeholder="Кашель, одышка, лихорадка...")
            with col_ct3:
                covid_suspicion = st.checkbox("Подозрение на COVID-19")
            
            if ct_type:
                additional_context += f"Тип КТ: {ct_type}\n"
            if clinical:
                additional_context += f"Клинические данные: {clinical}\n"
            if covid_suspicion:
                additional_context += "ВАЖНО: Оцени CT severity score для COVID-19!\n"
        else:
            additional_context = st.text_area(
                "Дополнительный контекст",
                placeholder="Опишите клиническую ситуацию, жалобы пациента, цель исследования...",
                help="Любая дополнительная информация, которая поможет в анализе"
            )
        
        # Кнопка анализа
        if st.button("🎬 Анализировать видео", type="primary", use_container_width=True):
            # Нормализуем study_type: пустая строка становится None
            # Проверяем явно на пустую строку и None
            if not study_type or study_type == "" or study_type.strip() == "":
                study_type_for_request = None
                st.info("💡 Тип исследования не выбран. Будет использован базовый промпт для анализа.")
            else:
                study_type_for_request = study_type
                # Показываем, какой тип исследования выбран
                study_type_names = {
                    "fgds": "🔬 ФГДС",
                    "colonoscopy": "🔬 Колоноскопия",
                    "echo": "🫀 ЭхоКГ",
                    "abdominal_us": "🔍 УЗИ органов брюшной полости",
                    "gynecology_us": "🩺 Гинекологическое УЗИ",
                    "mri_brain": "🧠 МРТ головного мозга",
                    "mri_universal": "🧲 МРТ (универсальный)",
                    "chest_ct": "🫁 КТ органов грудной клетки"
                }
                selected_name = study_type_names.get(study_type, study_type)
                st.success(f"✅ Используется специализированный промпт: {selected_name}")
            # Показываем прогресс
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            try:
                status_text.info("🔄 Подготовка видео...")
                progress_bar.progress(10)
                
                assistant = OpenRouterAssistant()
                
                # Подготавливаем метаданные
                metadata = {}
                if patient_age:
                    metadata['patient_age'] = patient_age
                if specialty:
                    metadata['specialty'] = specialty
                if urgency:
                    metadata['urgency'] = urgency
                if additional_context:
                    metadata['additional_context'] = additional_context
                
                # Формируем дополнительный промпт из контекста, если есть
                context_prompt = None
                if metadata:
                    context_parts = []
                    if patient_age:
                        context_parts.append(f"Возраст пациента: {patient_age} лет")
                    if specialty:
                        context_parts.append(f"Специализация: {specialty}")
                    if urgency:
                        context_parts.append(f"Срочность: {urgency}")
                    if additional_context:
                        context_parts.append(f"Дополнительный контекст: {additional_context}")
                    
                    if context_parts:
                        context_prompt = "\n\nКОНТЕКСТ:\n" + "\n".join(context_parts)
                
                # Двухэтапный анализ видео
                status_text.info("🔄 Этап 1: Специализированный анализ через Gemini 2.5 Flash...")
                progress_bar.progress(20)
                
                # Этап 1: Специализированный анализ
                with st.spinner("⏳ Анализ видео через Gemini..."):
                    results = assistant.send_video_request_two_stage(
                        prompt=context_prompt,
                        video_data=uploaded_video,
                        metadata=metadata if metadata else None,
                        study_type=study_type_for_request
                    )
                
                progress_bar.progress(50)
                
                # Показываем промежуточный результат (специализированный анализ)
                if results.get('specialized'):
                    st.subheader("📋 Промежуточный результат: Специализированный анализ")
                    with st.expander("🔍 Показать специализированный анализ (Gemini 2.5 Flash)", expanded=True):
                        st.markdown(results['specialized'])
                
                # Этап 2: Итоговое заключение от профессора
                if results.get('final') and not results['final'].startswith("❌"):
                    status_text.info("🔄 Этап 2: Итоговое заключение от профессора (Claude Opus)...")
                    progress_bar.progress(70)
                    
                    # Результат уже получен в двухэтапном методе, просто показываем прогресс
                    time.sleep(0.5)  # Небольшая задержка для визуализации прогресса
                    
                    progress_bar.progress(100)
                    status_text.empty()
                    progress_bar.empty()
                    
                    # Показываем финальное заключение
                    st.subheader("🎓 Итоговое заключение")
                    st.markdown(results['final'])
                elif results.get('final') and results['final'].startswith("❌"):
                    # Если была ошибка на этапе 2, показываем её
                    progress_bar.progress(100)
                    status_text.empty()
                    progress_bar.empty()
                    st.warning(f"⚠️ {results['final']}")
                    st.info("💡 Специализированный анализ доступен выше")
                else:
                    progress_bar.progress(100)
                    status_text.empty()
                    progress_bar.empty()
                    st.info("💡 Итоговое заключение не было сформировано. Доступен только специализированный анализ.")
                
                # Экспорт в DOC формат
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                
                # Формируем содержимое для DOC
                study_type_names = {
                    "fgds": "ФГДС",
                    "colonoscopy": "Колоноскопия",
                    "echo": "ЭхоКГ",
                    "abdominal_us": "УЗИ органов брюшной полости",
                    "gynecology_us": "Гинекологическое УЗИ",
                    "mri_brain": "МРТ головного мозга",
                    "mri_universal": "МРТ (универсальный)",
                    "chest_ct": "КТ органов грудной клетки"
                }
                study_name = study_type_names.get(study_type_for_request, "Видео-анализ") if study_type_for_request else "Видео-анализ"
                
                # Создаем DOC документ
                try:
                    from docx import Document
                    from docx.shared import Pt, Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    
                    doc = Document()
                    
                    # Заголовок
                    title = doc.add_heading(f"Анализ видео: {study_name}", level=0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Метаданные
                    doc.add_paragraph(f"Дата анализа: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}")
                    if patient_age:
                        doc.add_paragraph(f"Возраст пациента: {patient_age} лет")
                    if specialty:
                        doc.add_paragraph(f"Специализация: {specialty}")
                    if urgency:
                        doc.add_paragraph(f"Срочность: {urgency}")
                    doc.add_paragraph()
                    
                    # Раздел 1: Специализированный анализ
                    if results.get('specialized'):
                        doc.add_heading("СПЕЦИАЛИЗИРОВАННЫЙ АНАЛИЗ (Gemini 2.5 Flash)", level=1)
                        # Убираем markdown форматирование для чистого текста
                        specialized_text = results['specialized'].replace('**', '').replace('🎬', '').strip()
                        doc.add_paragraph(specialized_text)
                        doc.add_paragraph()
                    
                    # Раздел 2: Итоговое заключение
                    if results.get('final'):
                        doc.add_heading("ИТОГОВОЕ ЗАКЛЮЧЕНИЕ (Профессор, Claude Opus 4.5)", level=1)
                        final_text = results['final'].replace('**', '').replace('🎓', '').strip()
                        doc.add_paragraph(final_text)
                    
                    # Сохраняем в BytesIO для скачивания
                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    # Кнопка скачивания DOC
                    doc_filename = f"video_analysis_{study_name.replace(' ', '_')}_{timestamp}.docx"
                    st.download_button(
                        label="📥 Скачать полный отчет (.docx)",
                        data=doc_buffer.getvalue(),
                        file_name=doc_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except ImportError:
                    # Если python-docx не установлен, предлагаем TXT
                    st.warning("⚠️ Для экспорта в DOC формат требуется python-docx. Установите: pip install python-docx")
                    # Альтернатива: TXT файл
                    full_text = f"АНАЛИЗ ВИДЕО: {study_name}\n"
                    full_text += f"Дата: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}\n\n"
                    if results.get('specialized'):
                        full_text += "=" * 50 + "\n"
                        full_text += "СПЕЦИАЛИЗИРОВАННЫЙ АНАЛИЗ (Gemini 2.5 Flash)\n"
                        full_text += "=" * 50 + "\n"
                        full_text += results['specialized'] + "\n\n"
                    if results.get('final'):
                        full_text += "=" * 50 + "\n"
                        full_text += "ИТОГОВОЕ ЗАКЛЮЧЕНИЕ (Профессор, Claude Opus 4.5)\n"
                        full_text += "=" * 50 + "\n"
                        full_text += results['final'] + "\n"
                    
                    txt_filename = f"video_analysis_{timestamp}.txt"
                    st.download_button(
                        label="📥 Скачать отчет (.txt)",
                        data=full_text,
                        file_name=txt_filename,
                        mime="text/plain"
                    )
                
            except Exception as e:
                progress_bar.empty()
                status_text.empty()
                st.error(f"❌ Ошибка анализа видео: {e}")
                import traceback
                with st.expander("🔍 Детали ошибки"):
                    st.code(traceback.format_exc())
    else:
        st.info("👆 Загрузите видео-файл для начала анализа")



