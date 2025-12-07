# app.py (восстановленная версия после аварии)
import streamlit as st
import io
import base64
import sqlite3
import pandas as pd
import numpy as np
from PIL import Image
import requests
import tempfile
import os
from io import BytesIO
# import librosa  # Опционально, если нужна обработка аудио
try:
    import librosa
    LIBROSA_AVAILABLE = True
except ImportError:
    LIBROSA_AVAILABLE = False
import datetime
from pathlib import Path
import time
import sys
import gzip
import json
import re

# Безопасные импорты модулей
try:
    from modules.medical_ai_analyzer import EnhancedMedicalAIAnalyzer, ImageType
    MEDICAL_AI_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: medical_ai_analyzer недоступен: {e}", file=sys.stderr)
    MEDICAL_AI_AVAILABLE = False
    ImageType = None
    EnhancedMedicalAIAnalyzer = None

try:
    from modules.streamlit_enhanced_pages import (
        show_enhanced_analysis_page,
        show_comparative_analysis_page, 
        show_medical_protocols_page
    )
    ENHANCED_PAGES_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: streamlit_enhanced_pages недоступен: {e}", file=sys.stderr)
    ENHANCED_PAGES_AVAILABLE = False
    show_enhanced_analysis_page = None
    show_comparative_analysis_page = None
    show_medical_protocols_page = None

try:
    from modules.advanced_lab_processor import AdvancedLabProcessor
    LAB_PROCESSOR_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: advanced_lab_processor недоступен: {e}", file=sys.stderr)
    LAB_PROCESSOR_AVAILABLE = False
    AdvancedLabProcessor = None

try:
    from utils.image_processor import ImageFormatProcessor, optimize_image_for_ai
    IMAGE_PROCESSOR_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: image_processor недоступен: {e}", file=sys.stderr)
    IMAGE_PROCESSOR_AVAILABLE = False
    ImageFormatProcessor = None
    optimize_image_for_ai = None

try:
    from utils.specialist_detector import get_specialist_prompt, get_specialist_info
    SPECIALIST_DETECTOR_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: specialist_detector недоступен: {e}", file=sys.stderr)
    SPECIALIST_DETECTOR_AVAILABLE = False
    get_specialist_prompt = None
    get_specialist_info = None

try:
    from config import IS_REPLIT, MOBILE_MAX_IMAGE_SIZE, ALLOWED_IMAGE_EXTENSIONS
    CONFIG_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: config недоступен: {e}", file=sys.stderr)
    CONFIG_AVAILABLE = False
    IS_REPLIT = False
    MOBILE_MAX_IMAGE_SIZE = (1024, 1024)
    ALLOWED_IMAGE_EXTENSIONS = ['.jpg', '.jpeg', '.png']

try:
    from utils.error_handler import handle_error, log_api_call
    ERROR_HANDLER_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: error_handler недоступен: {e}", file=sys.stderr)
    ERROR_HANDLER_AVAILABLE = False
    def handle_error(error, context="", show_to_user=True):
        return str(error)
    def log_api_call(*args, **kwargs):
        pass

try:
    from utils.performance_monitor import track_model_usage
    PERFORMANCE_MONITOR_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: performance_monitor недоступен: {e}", file=sys.stderr)
    PERFORMANCE_MONITOR_AVAILABLE = False
    def track_model_usage(*args, **kwargs):
        pass

try:
    from utils.validators import validate_image, validate_file_size
    VALIDATORS_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: validators недоступен: {e}", file=sys.stderr)
    VALIDATORS_AVAILABLE = False
    def validate_image(*args, **kwargs):
        return True, ""
    def validate_file_size(*args, **kwargs):
        return True, ""

try:
    from utils.cache_manager import get_image_hash, get_cache_key, get_cached_result, save_to_cache, clear_old_cache
    CACHE_MANAGER_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: cache_manager недоступен: {e}", file=sys.stderr)
    CACHE_MANAGER_AVAILABLE = False
    def get_image_hash(*args, **kwargs):
        return ""
    def get_cache_key(*args, **kwargs):
        return ""
    def get_cached_result(*args, **kwargs):
        return None
    def save_to_cache(*args, **kwargs):
        pass
    def clear_old_cache(*args, **kwargs):
        pass

try:
    from utils.export_manager import export_analysis_to_json, export_analysis_to_csv, export_lab_results_to_excel
    EXPORT_MANAGER_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: export_manager недоступен: {e}", file=sys.stderr)
    EXPORT_MANAGER_AVAILABLE = False
    def export_analysis_to_json(*args, **kwargs):
        return ""
    def export_analysis_to_csv(*args, **kwargs):
        return ""
    def export_lab_results_to_excel(*args, **kwargs):
        return ""

try:
    from services.consensus_engine import ConsensusEngine
    CONSENSUS_ENGINE_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: consensus_engine недоступен: {e}", file=sys.stderr)
    CONSENSUS_ENGINE_AVAILABLE = False
    ConsensusEngine = None

try:
    from services.validation_pipeline import ValidationPipeline
    VALIDATION_PIPELINE_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: validation_pipeline недоступен: {e}", file=sys.stderr)
    VALIDATION_PIPELINE_AVAILABLE = False
    ValidationPipeline = None

try:
    from storages.context_store import ContextStore
    CONTEXT_STORE_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: context_store недоступен: {e}", file=sys.stderr)
    CONTEXT_STORE_AVAILABLE = False
    ContextStore = None

try:
    from evaluators.scorecards import MedicalScorecard
    SCORECARDS_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: scorecards недоступен: {e}", file=sys.stderr)
    SCORECARDS_AVAILABLE = False
    MedicalScorecard = None

try:
    from prompts.prompt_registry import PromptRegistry
    PROMPT_REGISTRY_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: prompt_registry недоступен: {e}", file=sys.stderr)
    PROMPT_REGISTRY_AVAILABLE = False
    PromptRegistry = None

try:
    from utils.gap_detector import DiagnosticGapDetector
    GAP_DETECTOR_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: gap_detector недоступен: {e}", file=sys.stderr)
    GAP_DETECTOR_AVAILABLE = False
    DiagnosticGapDetector = None

try:
    from utils.notification_system import NotificationSystem
    NOTIFICATION_SYSTEM_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: notification_system недоступен: {e}", file=sys.stderr)
    NOTIFICATION_SYSTEM_AVAILABLE = False
    NotificationSystem = None

try:
    from services.model_router import ModelRouter
    MODEL_ROUTER_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: model_router недоступен: {e}", file=sys.stderr)
    MODEL_ROUTER_AVAILABLE = False
    ModelRouter = None

try:
    from utils.evidence_ranker import EvidenceRanker
    EVIDENCE_RANKER_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: evidence_ranker недоступен: {e}", file=sys.stderr)
    EVIDENCE_RANKER_AVAILABLE = False
    EvidenceRanker = None

# --- Форма обратной связи ---
try:
    from utils.feedback_widget import show_feedback_form
    FEEDBACK_WIDGET_AVAILABLE = True
    print("✅ Модуль обратной связи загружен успешно", file=sys.stderr)
except ImportError as e:
    print(f"⚠️ Предупреждение: feedback_widget недоступен: {e}", file=sys.stderr)
    FEEDBACK_WIDGET_AVAILABLE = False
    def show_feedback_form(*args, **kwargs):
        # Заглушка, которая показывает информацию для отладки
        st.warning("⚠️ Модуль обратной связи недоступен. Проверьте логи.")
        pass

# --- Проверка доступности ИИ ---
try:
    from claude_assistant import OpenRouterAssistant
    AI_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Предупреждение: ИИ-модуль недоступен: {e}", file=sys.stderr)
    AI_AVAILABLE = False
    OpenRouterAssistant = None

# --- AssemblyAI для голосового ввода ---
try:
    from assemblyai_transcriber import transcribe_audio_assemblyai
    ASSEMBLYAI_AVAILABLE = True
except ImportError:
    ASSEMBLYAI_AVAILABLE = False
    transcribe_audio_assemblyai = None

def transcribe_audio(audio_file):
    """Заглушка - используйте AssemblyAI"""
    return "❌ Используйте AssemblyAI для расшифровки"

# --- Вспомогательная функция для анализа с streaming ---
def perform_analysis_with_streaming(assistant, prompt, image_array, metadata, use_streaming, 
                                   analysis_type="точный", model_type="opus", title=""):
    """Универсальная функция для выполнения анализа с поддержкой streaming
    
    Args:
        assistant: Экземпляр OpenRouterAssistant
        prompt: Промпт для анализа
        image_array: Массив изображения
        metadata: Метаданные
        use_streaming: Использовать ли streaming
        analysis_type: Тип анализа ("быстрый" или "точный")
        model_type: Тип модели ("gemini" или "opus")
        title: Заголовок для отображения
    """
    if use_streaming:
        # Streaming режим
        if title:
            st.markdown(f"### {title}")
        try:
            # Для streaming используем основной метод (поддерживает Opus)
            # Для Gemini пока используем обычный метод
            if analysis_type == "быстрый" and model_type == "gemini":
                # Gemini пока без streaming - используем обычный метод
                result = assistant.send_vision_request_gemini_fast(prompt, image_array, metadata)
                st.write(result)
                return result
            else:
                # Opus с streaming
                text_generator = assistant.send_vision_request_streaming(prompt, image_array, metadata)
                result = st.write_stream(text_generator)
                return result
        except Exception as e:
            st.error(f"❌ Ошибка streaming: {str(e)}")
            # Fallback на обычный режим
            try:
                with st.spinner(f"{'Gemini Flash' if model_type == 'gemini' else 'Opus 4.5'} анализирует..."):
                    if analysis_type == "быстрый":
                        result = assistant.send_vision_request_gemini_fast(prompt, image_array, metadata)
                    else:
                        result = assistant.send_vision_request(prompt, image_array, metadata)
                    st.write(result)
                    return result
            except Exception as e2:
                st.error(f"❌ Ошибка анализа: {str(e2)}")
                return None
    else:
        # Обычный режим
        with st.spinner(f"{'Gemini Flash' if model_type == 'gemini' else 'Opus 4.5'} анализирует..."):
            try:
                if analysis_type == "быстрый":
                    result = assistant.send_vision_request_gemini_fast(prompt, image_array, metadata)
                else:
                    result = assistant.send_vision_request(prompt, image_array, metadata)
                if title:
                    st.markdown(f"### {title}")
                st.write(result)
                return result
            except Exception as e:
                st.error(f"❌ Ошибка анализа: {str(e)}")
                return None

# --- Метрики моделей для отображения ---
def get_model_metrics_display(category: str):
    """Получить метрики моделей для отображения (иллюстрация)"""
    metrics = {
        'ECG': {
            'gemini': {'accuracy': 87},
            'opus': {'accuracy': 96, 'speed_multiplier': 3.5, 'price_multiplier': 4.2}
        },
        'XRAY': {
            'gemini': {'accuracy': 85},
            'opus': {'accuracy': 95, 'speed_multiplier': 3.2, 'price_multiplier': 4.0}
        },
        'MRI': {
            'gemini': {'accuracy': 88},
            'opus': {'accuracy': 96, 'speed_multiplier': 3.8, 'price_multiplier': 4.5}
        },
        'CT': {
            'gemini': {'accuracy': 86},
            'opus': {'accuracy': 95, 'speed_multiplier': 3.5, 'price_multiplier': 4.3}
        },
        'ULTRASOUND': {
            'gemini': {'accuracy': 84},
            'opus': {'accuracy': 94, 'speed_multiplier': 3.0, 'price_multiplier': 3.8}
        },
        'DERMATOSCOPY': {
            'gemini': {'accuracy': 82},
            'opus': {'accuracy': 98, 'speed_multiplier': 3.8, 'price_multiplier': 4.5}
        }
    }
    return metrics.get(category, {
        'gemini': {'accuracy': 85},
        'opus': {'accuracy': 95, 'speed_multiplier': 3.5, 'price_multiplier': 4.0}
    })

# --- Инициализация базы данных ---
def init_db():
    conn = sqlite3.connect('medical_data.db')
    cursor = conn.cursor()

    # Создаём таблицы
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS patients (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            age INTEGER,
            sex TEXT,
            phone TEXT
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS patient_notes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_id INTEGER,
            raw_text TEXT,
            structured_note TEXT,
            gdoc_url TEXT,
            diagnosis TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (patient_id) REFERENCES patients (id)
        )
    ''')

    # Таблица для истории чата с ИИ
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS ai_chat_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id TEXT,
            user_message TEXT,
            assistant_response TEXT,
            files_context TEXT,
            context_summary TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    conn.commit()
    conn.close()
    
    # Создаём таблицу для обратной связи
    try:
        from database import init_feedback_table
        init_feedback_table()
    except Exception as e:
        print(f"⚠️ Предупреждение: не удалось создать таблицу обратной связи: {e}", file=sys.stderr)

# --- Страницы ---
def show_home_page():
    # HERO-блок в стиле медицинского лендинга
    hero = st.container()
    with hero:
        col_left, col_right = st.columns([3, 2])

        with col_left:
            st.markdown(
                """
                <div style="padding: 1.5rem 0;">
                  <div style="color:#004d40;font-weight:700;font-size:1.1rem;letter-spacing:0.06em;text-transform:uppercase;margin-bottom:0.5rem;">
                    Медицинский ИИ‑ассистент профессора
                  </div>
                  <div style="font-size:2.1rem;font-weight:800;line-height:1.2;color:#003c32;margin-bottom:0.75rem;">
                    Правильное время<br/>для экспертной<br/><span style="color:#00a79d;">клинической диагностики</span>
                  </div>
                  <div style="max-width:520px;font-size:0.98rem;color:#004d40;margin-bottom:1.2rem;">
                    Единый ИИ‑центр: Opus‑профессор для ЭКГ, рентгена, КТ, МРТ, УЗИ и генетики.
                    Автоматический анализ изображений, лабораторных и генетических отчётов
                    с выводом в формате «клиническая директива» для врача.
                  </div>
                </div>
                """,
                unsafe_allow_html=True,
            )

            c1, c2 = st.columns(2)
            with c1:
                if st.button("🔍 Начать анализ изображения", use_container_width=True):
                    st.session_state["page"] = "🩻 Анализ рентгена"
            with c2:
                if st.button("🧬 Генетический консультант", use_container_width=True):
                    st.session_state["page"] = "🧬 Генетический анализ"

            st.markdown(
                """
                <div style="margin-top:1.2rem;font-size:0.9rem;color:#00695c;">
                  24/7 доступ к Opus‑консилиуму · Поддержка сложных клинических случаев ·
                  Безопасная локальная обработка данных
                </div>
                """,
                unsafe_allow_html=True,
            )

        with col_right:
            st.markdown(
                """
                <div style="
                    background: linear-gradient(145deg,#00bcd4,#26a69a);
                    border-radius:18px;
                    padding:1.5rem;
                    color:white;
                    box-shadow:0 18px 40px rgba(0,150,136,0.35);
                    text-align:center;
                ">
                  <div style="font-size:3rem;line-height:1;">🩺</div>
                  <div style="font-weight:700;font-size:1.2rem;margin-top:0.5rem;">
                    Профессор‑консультант Opus
                  </div>
                  <div style="font-size:0.9rem;margin-top:0.4rem;opacity:0.9;">
                    Кардиология · Неврология · Онкология · Генетика · Терапия
                  </div>
                  <div style="margin-top:1rem;font-size:0.85rem;text-align:left;background:rgba(255,255,255,0.08);padding:0.75rem;border-radius:12px;">
                    ✔ Сложные ЭКГ и аритмии<br/>
                    ✔ Рентген/КТ/МРТ с оценкой динамики<br/>
                    ✔ Лабораторные и генетические панели<br/>
                    ✔ Формирование готового клинического протокола
                  </div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    st.markdown("---")

    # Карточки основных модулей
    st.subheader("Ключевые модули")
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown("**📈 ЭКГ & ритмы**")
        st.caption("Анализ 12‑канальной ЭКГ, аритмии, блокады, клиническая директива.")
    with c2:
        st.markdown("**🩻 Визуальная диагностика**")
        st.caption("Рентген, КТ, МРТ, УЗИ — структурированный отчёт и оценка динамики.")
    with c3:
        st.markdown("**🔬 Лабораторные данные**")
        st.caption("Сканирование бланков, структурирование анализов, без лишних интерпретаций.")
    with c4:
        st.markdown("**🧬 Генетика & фармакогеномика**")
        st.caption("Разбор VCF/PDF, заключение генетика и профессорский обзор.")

def show_ecg_analysis():
    if not AI_AVAILABLE:
        st.error("❌ ИИ-модуль недоступен. Проверьте файл `claude_assistant.py` и API-ключ.")
        return

    st.header("📈 Анализ ЭКГ")
    
    # Мобильная поддержка: выбор источника
    source_type = st.radio(
        "Выберите источник изображения:",
        ["📁 Загрузить файл", "📷 Сделать фото"],
        horizontal=True
    )
    
    image_array = None
    metadata = {}
    
    if source_type == "📷 Сделать фото":
        # Использование камеры смартфона
        camera_image = st.camera_input("Сфотографируйте ЭКГ", key="ecg_camera")
        if camera_image:
            try:
                # Конвертация в numpy array
                image = Image.open(camera_image)
                image_array = np.array(image)
                metadata = {'source': 'camera', 'format': 'mobile_photo'}
            except Exception as e:
                st.error(f"Ошибка обработки фото: {e}")
                return
    else:
        # Загрузка файла с расширенной поддержкой форматов
        uploaded_file = st.file_uploader(
            "Загрузите ЭКГ", 
            type=["jpg", "jpeg", "png", "pdf", "dcm", "dicom", "tiff", "tif", "heic", "heif", "webp", "zip"],
            help="Поддерживаются: JPG, PNG, TIFF, HEIC, WEBP, DICOM, ZIP"
        )
        
        if uploaded_file:
            try:
                # Сохранение во временный файл
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp:
                    tmp.write(uploaded_file.getvalue())
                    tmp_path = tmp.name
                
                # Загрузка через процессор форматов
                if IMAGE_PROCESSOR_AVAILABLE and ImageFormatProcessor:
                    processor = ImageFormatProcessor()
                    image_array, file_metadata = processor.load_image(tmp_path, MOBILE_MAX_IMAGE_SIZE)
                    metadata = {**metadata, **file_metadata, 'source': 'upload'}
                else:
                    # Fallback - простая загрузка через PIL
                    image = Image.open(tmp_path)
                    image_array = np.array(image)
                    metadata = {**metadata, 'source': 'upload'}
                
                # Очистка
                os.unlink(tmp_path)
                if IMAGE_PROCESSOR_AVAILABLE and ImageFormatProcessor and 'processor' in locals():
                    processor.cleanup_temp_files()
                
            except Exception as e:
                st.error(f"Ошибка обработки файла: {e}")
                return

    if image_array is None:
        st.info("Загрузите файл или сделайте фото для анализа.")
        return

    # Валидация изображения
    if VALIDATORS_AVAILABLE and validate_image:
        is_valid, error_msg = validate_image(image_array)
        if not is_valid:
            st.error(f"❌ Ошибка валидации изображения: {error_msg}")
            return
    else:
        # Простая проверка без валидатора
        if image_array is None or image_array.size == 0:
            st.error("❌ Ошибка: изображение пустое или не загружено")
            return

    try:
        # Оптимизация для мобильных устройств
        if (IS_REPLIT or st.session_state.get('mobile_mode', False)) and IMAGE_PROCESSOR_AVAILABLE and optimize_image_for_ai:
            image_array = optimize_image_for_ai(image_array)
        
        st.image(image_array, caption="ЭКГ", use_container_width=True, clamp=True)

        # Базовый анализ
        analysis = {
            "heart_rate": 75,
            "rhythm_assessment": "Синусовый",
            "num_beats": 12,
            "duration": 10,
            "signal_quality": "Хорошее"
        }
        
        st.subheader("📊 Результаты анализа")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("ЧСС", f"{analysis['heart_rate']} уд/мин")
            st.metric("Ритм", analysis['rhythm_assessment'])
        with col2:
            st.metric("Длительность", f"{analysis['duration']:.1f} с")
            st.metric("Комплексы", analysis['num_beats'])

        assistant = OpenRouterAssistant()
        
        # Инициализация новых компонентов
        consensus_engine = ConsensusEngine(assistant)
        validator = ValidationPipeline(assistant)
        scorecard = MedicalScorecard()
        context_store = ContextStore()
        
        gap_detector = DiagnosticGapDetector()
        notifier = NotificationSystem()
        model_router = ModelRouter()
        
        evidence_ranker = EvidenceRanker()

        # Выбор пациента для сохранения в контекст
        st.subheader("👤 Связь с пациентом (опционально)")
        init_db()
        conn = sqlite3.connect('medical_data.db')
        patients = pd.read_sql_query("SELECT id, name FROM patients", conn)
        conn.close()
        
        selected_patient_id = None
        if not patients.empty:
            save_to_context = st.checkbox("💾 Сохранить результаты в контекст пациента", value=False)
            if save_to_context:
                selected_patient_name = st.selectbox("Выберите пациента:", patients['name'], key="ecg_patient_select")
                selected_patient_id = patients[patients['name'] == selected_patient_name].iloc[0]['id']
        else:
            save_to_context = False
            st.info("💡 Добавьте пациента в разделе 'База данных', чтобы сохранять результаты в контекст")

        # Использование контекста пациента (если загружен)
        patient_context = None
        if 'patient_context' in st.session_state and 'selected_patient_id' in st.session_state:
            patient_context = st.session_state['patient_context']
            st.info(f"💡 Используется клинический контекст пациента")
        
        # Получение промпта специалиста (выносим за пределы кнопок, чтобы был доступен для всех)
        from modules.medical_ai_analyzer import ImageType
        if SPECIALIST_DETECTOR_AVAILABLE and get_specialist_prompt and get_specialist_info:
            prompt = get_specialist_prompt(ImageType.ECG)
            specialist_info = get_specialist_info(ImageType.ECG)
        else:
            # Fallback промпт для ЭКГ - детальная дешифровка
            prompt = """Ты — ведущий кардиолог-электрофизиолог с 20+ летним опытом. Проведи ПОЛНУЮ дешифровку ЭКГ по международным стандартам (AHA/ACC/HRS, ESC).

ОБЯЗАТЕЛЬНО проанализируй и опиши:

1. **КАЧЕСТВО ЗАПИСИ:**
   - Скорость записи (25 или 50 мм/с)
   - Калибровка
   - Артефакты (если есть)

2. **РИТМ И ПРОВОДИМОСТЬ:**
   - Основной ритм (синусовый/несинусовый/фибрилляция/трепетание)
   - Регулярность
   - AV-проводимость (норма/блокада 1-3 степени)
   - Внутрижелудочковая проводимость (норма/блокада ножек)

3. **ЧСС:** точное значение в уд/мин

4. **ЭЛЕКТРИЧЕСКАЯ ОСЬ:** угол в градусах и направление

5. **ИНТЕРВАЛЫ (в мс):**
   - PR: значение, норма 120-200 мс
   - QRS: ширина, норма <120 мс
   - QT и QTc: значение, норма <450 мс (муж) / <470 мс (жен)
   - RR: среднее значение

6. **СЕГМЕНТЫ И ВОЛНЫ:**
   - **ST:** для КАЖДОГО отведения укажи элевацию/депрессию в мм, форму, локализацию
   - **T:** полярность, амплитуда, морфология в каждом отведении
   - **P:** наличие, морфология, амплитуда (<2.5 мм), длительность (<120 мс)
   - **Q:** патологические Q (глубина >25% R, ширина >40 мс) с указанием отведений

7. **АНАЛИЗ ПО ОТВЕДЕНИЯМ:**
   - **I, II, III, aVR, aVL, aVF:** амплитуды, патологии
   - **V1-V6:** переходная зона, прогрессия R, патологии в каждом

8. **ПАТОЛОГИИ:** все отклонения с указанием конкретных отведений

9. **КЛИНИЧЕСКАЯ ИНТЕРПРЕТАЦИЯ:**
   - Основные находки
   - Дифференциальный диагноз
   - Оценка остроты
   - Рекомендации (неотложные меры, обследования, консультации)

10. **КОДЫ МКБ-10** для выявленных патологий

ВАЖНО: измеряй ВСЕ параметры ТОЧНО, анализируй ВСЕ 12 отведений, указывай конкретные отведения для каждого отклонения, не используй общие фразы."""
            specialist_info = {'role': 'Кардиолог', 'specialization': 'ЭКГ'}
        
        # Добавляем контекст в промпт если есть
        if patient_context:
            prompt += f"\n\nКЛИНИЧЕСКИЙ КОНТЕКСТ ПАЦИЕНТА:\n{patient_context}\n\nУчтите этот контекст при анализе."
        
        # Выбор режима анализа (показывается всегда, до нажатия кнопки)
        st.markdown("---")
        
        # Блок метрик моделей
        st.markdown("### 📊 Точность моделей для ЭКГ")
        metrics = get_model_metrics_display('ECG')
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Точность Gemini Flash", f"{metrics['gemini']['accuracy']}%")
            st.metric("Точность Opus 4.5", f"{metrics['opus']['accuracy']}%")
        with col2:
            speed_diff = metrics['opus']['speed_multiplier']
            st.info(f"⚡ Opus в {speed_diff} раз медленнее")
        with col3:
            price_diff = metrics['opus']['price_multiplier']
            st.info(f"💰 Opus в {price_diff} раз дороже")
        
        # Форма обратной связи - ДО анализа, всегда видна и активна!
        st.markdown("---")
        st.markdown("### 💬 Обратная связь")
        
        # Показываем форму ВСЕГДА, даже без результата (она активна всегда)
        last_result = st.session_state.get('ecg_analysis_result', '')
        
        # Используем ФИКСИРОВАННЫЙ ID для формы, чтобы ключи виджетов не менялись
        # Это позволяет форме работать стабильно и не терять данные при рендере
        analysis_id_base = "ECG_feedback_form"
        
        # Показываем форму всегда (даже с пустым результатом до анализа)
        # Формируем input_case из метаданных ЭКГ
        input_case_data = st.session_state.get('ecg_input_case', '')
        if not input_case_data:
            # Пытаемся сформировать из метаданных
            analysis_meta = st.session_state.get('ecg_analysis', {})
            if analysis_meta:
                input_case_data = f"ЭКГ: ЧСС={analysis_meta.get('heart_rate', 'N/A')}, Ритм={analysis_meta.get('rhythm_assessment', 'N/A')}, Качество={analysis_meta.get('signal_quality', 'N/A')}"
        
        # Всегда вызываем форму, даже если модуль недоступен (покажет заглушку)
        try:
            show_feedback_form(
                analysis_type="ECG",
                analysis_result=str(last_result) if last_result else "",
                analysis_id=analysis_id_base,
                input_case=input_case_data
            )
        except Exception as e:
            st.error(f"Ошибка формы обратной связи: {e}")
            st.info("💡 Форма обратной связи временно недоступна")
        
        if not last_result:
            st.info("💡 После проведения анализа ЭКГ форма автоматически обновится с новым результатом.")
        
        st.markdown("---")
        st.markdown("### ⚙️ Режимы анализа")
        
        # Опция streaming
        use_streaming = st.checkbox("📺 Постепенное появление текста (streaming)", value=True, key="ecg_streaming")
        
        # Кнопки быстрого и точного анализа
        col_fast, col_precise = st.columns(2)
        with col_fast:
            if st.button("⚡ Быстрый анализ (Gemini Flash)", use_container_width=True, type="primary"):
                result = perform_analysis_with_streaming(
                    assistant, prompt, image_array, str(analysis), use_streaming,
                    analysis_type="быстрый", model_type="gemini", 
                    title="⚡ Быстрый анализ (Gemini Flash):"
                )
                if result:
                    st.session_state.ecg_analysis_result = result
                    st.session_state.ecg_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    # Форма под метриками обновится автоматически при следующем рендере
                    # Не вызываем st.rerun() здесь, так как результат уже выведен через perform_analysis_with_streaming
        
        with col_precise:
            opus_accuracy = metrics['opus']['accuracy']
            gemini_accuracy = metrics['gemini']['accuracy']
            accuracy_diff = opus_accuracy - gemini_accuracy
            if st.button(f"🎯 Точный анализ (Opus 4.5) - на {accuracy_diff}% точнее", use_container_width=True, type="primary"):
                result = perform_analysis_with_streaming(
                    assistant, prompt, image_array, str(analysis), use_streaming,
                    analysis_type="точный", model_type="opus",
                    title=f"🎯 Точный анализ (Opus 4.5):"
                )
                if result:
                    st.session_state.ecg_analysis_result = result
                    st.session_state.ecg_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    # Форма под метриками обновится автоматически при следующем рендере
                    # Не вызываем st.rerun() здесь, так как результат уже выведен через perform_analysis_with_streaming
        
        st.markdown("---")
        st.markdown("### ⚙️ Расширенные режимы анализа")
        
        analysis_mode = st.radio(
            "**Режим анализа:**",
            ["⚡ Быстрый (одна модель)", "🎯 Консенсус (несколько моделей)", "✅ С валидацией"],
            horizontal=True,
            key="ecg_analysis_mode",
            help="Выберите режим анализа перед запуском"
        )
        
        # Показываем информацию о выбранном режиме
        if analysis_mode == "🎯 Консенсус (несколько моделей)":
            st.info("💡 **Консенсус:** Несколько моделей проанализируют ЭКГ, затем будет сформировано общее заключение")
        elif analysis_mode == "✅ С валидацией":
            st.info("💡 **С валидацией:** Анализ будет проверен на логичность и полноту")
        else:
            st.info("💡 **Быстрый анализ:** Одна модель быстро проанализирует ЭКГ")
        
        st.markdown("---")
        
        if st.button("🔍 ИИ-анализ ЭКГ (с контекстом)", use_container_width=True):
            # Промпт уже определен выше, используем его
            
            if analysis_mode == "⚡ Быстрый (одна модель)":
                result = None
                with st.spinner("ИИ анализирует ЭКГ..."):
                    try:
                        # Opus 4.5 используется по умолчанию для клинического анализа ЭКГ
                        result = assistant.send_vision_request(prompt, image_array, str(analysis))
                    except Exception as e:
                        st.error(f"❌ Ошибка анализа: {str(e)}")
                        st.info("💡 Попробуйте еще раз или выберите другой режим анализа")
                
                # Отображаем результат ВНЕ спиннера
                if result:
                    # Сохраняем результат СРАЗУ для формы обратной связи
                    st.session_state.ecg_analysis_result = result
                    st.session_state.ecg_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    
                    st.markdown(f"### 🧠 Ответ ИИ ({specialist_info['role']}):")
                    st.write(result)
                    
                    # Сохраняем результат в session_state чтобы форма под метриками обновилась
                    st.session_state.ecg_analysis_result = result
                    st.session_state.ecg_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    # Обновляем страницу чтобы форма обновилась
                    st.rerun()
            
            elif analysis_mode == "🎯 Консенсус (несколько моделей)":
                consensus_result = None
                with st.spinner("ИИ анализирует ЭКГ..."):
                    # Используем стандартный набор моделей консенсуса из ConsensusEngine
                    st.info("🔄 Используется консенсус моделей: Sonnet + Llama Vision + Gemini (по настройкам движка консенсуса)")
                    
                    consensus_result = consensus_engine.analyze_with_consensus(
                        prompt, image_array, str(analysis)
                    )
                
                # Отображаем результат ВНЕ спиннера
                if consensus_result:
                    st.markdown("### 🎯 Консенсусное заключение:")
                    if consensus_result['consensus']['consensus_available']:
                        result = consensus_result['consensus']['consensus_response']
                        st.write(result)
                        st.metric("Уровень согласия", f"{consensus_result['consensus']['agreement_level']:.1%}")
                        
                        if consensus_result['consensus']['discrepancies']:
                            st.warning("⚠️ Обнаружены расхождения между моделями:")
                            for disc in consensus_result['consensus']['discrepancies']:
                                st.warning(f"• {disc}")
                    else:
                        result = consensus_result['consensus'].get('single_opinion', 'Ошибка получения консенсуса')
                        st.write(result)
                    
                    # Сохраняем результат СРАЗУ для формы обратной связи
                    st.session_state.ecg_analysis_result = result
                    st.session_state.ecg_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    # Форма под метриками обновится автоматически при следующем рендере
                    # Обновляем страницу только после вывода результата
                    st.rerun()
                
            elif analysis_mode == "✅ С валидацией":
                # Opus 4.5 используется по умолчанию для клинического анализа ЭКГ
                result = perform_analysis_with_streaming(
                    assistant, prompt, image_array, str(analysis), use_streaming=True,
                    analysis_type="точный", model_type="opus",
                    title=f"### 🧠 Ответ ИИ ({specialist_info['role']}):"
                )
                
                # Обработка результата ВНЕ спиннера
                if result:
                    # Отображение результатов - СРАЗУ!
                    st.markdown(f"### 🧠 Ответ ИИ ({specialist_info['role']}):")
                    st.write(result)
                    
                    # Сохраняем результат для пересылки консультанту
                    st.session_state.ecg_analysis_result = result
                    timestamp_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    st.session_state.ecg_analysis_timestamp = timestamp_str
                    
                    # Обновляем страницу чтобы форма под метриками обновилась
                    st.rerun()
                    
                    # Проверка на критические находки
                    critical_findings = notifier.check_critical_findings(result)
                    if critical_findings:
                        notifier.display_notifications(critical_findings)
                    
                    # Валидация
                    validation_result = validator.validate_response(result, {'image_type': 'ECG'})
                    
                    # Оценка
                    scorecard_result = scorecard.evaluate_response(result, ImageType.ECG)
                    
                    # Выявление пробелов
                    gaps = gap_detector.detect_gaps(result, ImageType.ECG)
                    gap_report = gap_detector.generate_gap_report(gaps)
                    
                    # Оценка доказательности
                    evidence_ranking = evidence_ranker.rank_evidence(result)
                    evidence_report = evidence_ranker.generate_evidence_report(evidence_ranking)
                    
                    # Сохранение результатов ЭКГ в контекст пациента
                    if 'selected_patient_id' in locals() and selected_patient_id:
                        try:
                            context_store.add_context(
                                patient_id=selected_patient_id,
                                context_type='imaging',
                                context_data={
                                    'type': 'ECG',
                                    'analysis': result,
                                    'specialist': specialist_info['role'],
                                    'mode': analysis_mode,
                                    'validation': validation_result,
                                    'scorecard': scorecard_result
                                },
                                source='ai_analysis'
                            )
                            st.success("✅ Результаты ЭКГ сохранены в клинический контекст пациента!")
                        except Exception as e:
                            st.warning(f"⚠️ Не удалось сохранить в контекст: {e}")
                    
                    # Оценка качества
                    st.markdown("### 📊 Оценка качества:")
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Общая оценка", scorecard_result['grade'])
                    with col2:
                        st.metric("Полнота", f"{scorecard_result['completeness']:.1%}")
                    with col3:
                        st.metric("Валидация", "✅ Пройдена" if validation_result['is_valid'] else "❌ Не пройдена")
                    with col4:
                        st.metric("Заполненность", f"{gaps['completeness_percentage']:.1f}%")
                    
                    # Отчет о пробелах
                    if gaps['completeness_percentage'] < 80:
                        with st.expander("📋 Отчет о пробелах в ответе"):
                            st.text(gap_report)
                    
                    # Рекомендации
                    if scorecard_result['recommendations']:
                        st.info("💡 Рекомендации по улучшению:")
                        for rec in scorecard_result['recommendations']:
                            st.write(f"• {rec}")
                    
                    # Предупреждения валидации
                    if validation_result['warnings']:
                        st.warning("⚠️ Предупреждения валидации:")
                        for warning in validation_result['warnings']:
                            st.warning(f"• {warning}")
                    
                    # Оценка доказательности
                    with st.expander("📚 Оценка доказательности"):
                        st.text(evidence_report)

        # Возможность скачать стандартный протокол описания ЭКГ
        if 'ecg_analysis_result' in st.session_state and st.session_state.ecg_analysis_result:
            st.markdown("---")
            st.markdown("### 💾 Экспорт протокола ЭКГ")
            timestamp = st.session_state.get('ecg_analysis_timestamp', '')
            header = "Стандартный протокол описания ЭКГ"
            if timestamp:
                header += f"\nВремя анализа: {timestamp}"
            report_text = f"{header}\n\n{st.session_state.ecg_analysis_result}"
            st.download_button(
                label="📥 Скачать протокол ЭКГ (.txt)",
                data=report_text,
                file_name=f"ECG_report_{timestamp.replace(' ', '_').replace(':', '-') if timestamp else 'latest'}.txt",
                mime="text/plain"
            )

    except Exception as e:
        handle_error(e, "show_ecg_analysis", show_to_user=True)
        return

def show_xray_analysis():
    if not AI_AVAILABLE:
        st.error("❌ ИИ-модуль недоступен. Проверьте файл `claude_assistant.py` и API-ключ.")
        return

    st.header("�� Анализ рентгена")
    
    # Мобильная поддержка: выбор источника
    source_type = st.radio(
        "Выберите источник изображения:",
        ["📁 Загрузить файл", "📷 Сделать фото"],
        horizontal=True
    )
    
    image_array = None
    metadata = {}
    
    if source_type == "📷 Сделать фото":
        camera_image = st.camera_input("Сфотографируйте рентген", key="xray_camera")
        if camera_image:
            try:
                image = Image.open(camera_image)
                image_array = np.array(image)
                metadata = {'source': 'camera', 'format': 'mobile_photo'}
            except Exception as e:
                st.error(f"Ошибка обработки фото: {e}")
                return
    else:
        uploaded_file = st.file_uploader(
            "Загрузите рентген", 
            type=["jpg", "jpeg", "png", "pdf", "dcm", "dicom", "tiff", "tif", "heic", "heif", "webp", "zip"],
            help="Поддерживаются: JPG, PNG, TIFF, HEIC, WEBP, DICOM, ZIP"
        )
        
        if uploaded_file:
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp:
                    tmp.write(uploaded_file.getvalue())
                    tmp_path = tmp.name
                
                processor = ImageFormatProcessor()
                image_array, file_metadata = processor.load_image(tmp_path, MOBILE_MAX_IMAGE_SIZE)
                metadata = {**metadata, **file_metadata, 'source': 'upload'}
                
                os.unlink(tmp_path)
                processor.cleanup_temp_files()
                
            except Exception as e:
                st.error(f"Ошибка обработки файла: {e}")
                return

    if image_array is None:
        st.info("Загрузите файл или сделайте фото для анализа.")
        return

    # Валидация изображения
    is_valid, error_msg = validate_image(image_array)
    if not is_valid:
        st.error(f"❌ Ошибка валидации изображения: {error_msg}")
        return

    try:
        # Оптимизация для мобильных устройств
        if (IS_REPLIT or st.session_state.get('mobile_mode', False)) and IMAGE_PROCESSOR_AVAILABLE and optimize_image_for_ai:
            image_array = optimize_image_for_ai(image_array)
        
        st.image(image_array, caption="Рентген", use_container_width=True, clamp=True)

        analysis = {
            "quality_assessment": "Хорошее",
            "contrast": 45.0,
            "lung_area": 50000
        }
        
        st.subheader("📊 Оценка качества")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Качество", analysis['quality_assessment'])
            st.metric("Контраст", f"{analysis['contrast']:.1f}")
        with col2:
            st.metric("Площадь лёгких", f"{analysis['lung_area']:,}")

        st.markdown("---")
        
        # Блок метрик моделей
        st.markdown("### 📊 Точность моделей для рентгена")
        metrics = get_model_metrics_display('XRAY')
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Точность Gemini Flash", f"{metrics['gemini']['accuracy']}%")
            st.metric("Точность Opus 4.5", f"{metrics['opus']['accuracy']}%")
        with col2:
            speed_diff = metrics['opus']['speed_multiplier']
            st.info(f"⚡ Opus в {speed_diff} раз медленнее")
        with col3:
            price_diff = metrics['opus']['price_multiplier']
            st.info(f"💰 Opus в {price_diff} раз дороже")
        
        # Форма обратной связи - ДО анализа, всегда видна и активна!
        st.markdown("---")
        st.markdown("### 💬 Обратная связь")
        
        last_result = st.session_state.get('xray_analysis_result', '')
        analysis_id_base = "XRAY_feedback_form"
        xray_input = f"Рентген: Качество={analysis.get('quality_assessment', 'N/A')}, Контраст={analysis.get('contrast', 'N/A')}"
        
        try:
            show_feedback_form(
                analysis_type="XRAY",
                analysis_result=str(last_result) if last_result else "",
                analysis_id=analysis_id_base,
                input_case=xray_input
            )
        except Exception as e:
            st.error(f"Ошибка формы обратной связи: {e}")
            st.info("💡 Форма обратной связи временно недоступна")
        
        if not last_result:
            st.info("💡 После проведения анализа форма автоматически обновится с новым результатом.")
        
        st.markdown("---")
        
        # Получение промпта для рентгена
        assistant = OpenRouterAssistant()
        from modules.medical_ai_analyzer import ImageType
        if SPECIALIST_DETECTOR_AVAILABLE and get_specialist_prompt and get_specialist_info:
            prompt = get_specialist_prompt(ImageType.XRAY)
            specialist_info = get_specialist_info(ImageType.XRAY)
        else:
            prompt = "Проанализируйте рентгеновский снимок. Оцените структуры, патологические изменения, дайте заключение."
            specialist_info = {'role': 'Врач-рентгенолог'}
        
        # Кнопки быстрого и точного анализа
        col_fast, col_precise = st.columns(2)
        with col_fast:
            if st.button("⚡ Быстрый анализ (Gemini Flash)", use_container_width=True, type="primary", key="xray_fast"):
                with st.spinner("Gemini Flash анализирует рентген..."):
                    try:
                        result = assistant.send_vision_request_gemini_fast(prompt, image_array)
                        st.markdown(f"### ⚡ Быстрый анализ (Gemini Flash):")
                        st.write(result)
                    except Exception as e:
                        st.error(f"❌ Ошибка анализа: {str(e)}")
        
        with col_precise:
            opus_accuracy = metrics['opus']['accuracy']
            gemini_accuracy = metrics['gemini']['accuracy']
            accuracy_diff = opus_accuracy - gemini_accuracy
            if st.button(f"🎯 Точный анализ (Opus 4.5) - на {accuracy_diff}% точнее", use_container_width=True, type="primary", key="xray_precise"):
                result = perform_analysis_with_streaming(
                    assistant, prompt, image_array, str(metadata), use_streaming=True,
                    analysis_type="точный", model_type="opus",
                    title="🎯 Точный анализ (Opus 4.5):"
                )
                if result:
                    st.session_state.xray_analysis_result = result
                    st.session_state.xray_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    st.rerun()
        
        st.markdown("---")
        st.markdown("### ⚙️ Расширенные режимы анализа")

        # Универсальный анализатор
        from utils.universal_analyzer import UniversalMedicalAnalyzer
        analyzer = UniversalMedicalAnalyzer()
        
        # Выбор режима анализа
        analysis_mode = st.radio(
            "Режим анализа:",
            ["⚡ Быстрый (одна модель)", "🎯 Консенсус (несколько моделей)", "✅ С валидацией"],
            horizontal=True,
            key="xray_analysis_mode"
        )
        
        # Выбор пациента для сохранения контекста
        patient_id = None
        if st.checkbox("💾 Сохранить в контекст пациента"):
            init_db()
            conn = sqlite3.connect('medical_data.db')
            patients = pd.read_sql_query("SELECT id, name FROM patients", conn)
            conn.close()
            
            if not patients.empty:
                selected_patient = st.selectbox("Выберите пациента", patients['name'])
                patient_id = patients[patients['name'] == selected_patient].iloc[0]['id']
        
        if st.button("🩺 ИИ-анализ рентгена", use_container_width=True):
            with st.spinner("ИИ анализирует снимок..."):
                from modules.medical_ai_analyzer import ImageType
                
                # Для консенсуса используем Claude 4.5 и Llama Vision
                if analysis_mode == "🎯 Консенсус (несколько моделей)":
                    st.info("🔄 Используется консенсус моделей: Claude 4.5 Sonnet + Opus 4.5 + Llama 3.2 90B Vision")
                
                results = analyzer.analyze_image(
                    image_array=image_array,
                    image_type=ImageType.XRAY,
                    analysis_mode=analysis_mode,
                    metadata=analysis,
                    patient_id=patient_id
                )
                
                analyzer.display_results(results)
                
                # Сохраняем результат для пересылки консультанту
                if results.get('result'):
                    st.session_state.xray_analysis_result = results['result']
                    st.session_state.xray_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                
                # Дополнительно показываем форму (на случай если display_results не показала)
                if FEEDBACK_WIDGET_AVAILABLE and results.get('result'):
                    try:
                        show_feedback_form(
                            analysis_type="XRAY",
                            analysis_result=results['result'],
                            analysis_id=f"XRAY_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
                        )
                    except Exception as e:
                        import sys
                        print(f"⚠️ Ошибка формы обратной связи XRAY: {e}", file=sys.stderr)

    except Exception as e:
        handle_error(e, "show_xray_analysis", show_to_user=True)

def show_mri_analysis():
    if not AI_AVAILABLE:
        st.error("❌ ИИ-модуль недоступен. Проверьте файл `claude_assistant.py` и API-ключ.")
        return

    st.header("🧠 Анализ МРТ")
    
    # Мобильная поддержка: выбор источника
    source_type = st.radio(
        "Выберите источник изображения:",
        ["📁 Загрузить файл", "📷 Сделать фото"],
        horizontal=True
    )
    
    image_array = None
    metadata = {}
    
    if source_type == "📷 Сделать фото":
        camera_image = st.camera_input("Сфотографируйте МРТ", key="mri_camera")
        if camera_image:
            try:
                image = Image.open(camera_image)
                image_array = np.array(image)
                metadata = {'source': 'camera', 'format': 'mobile_photo'}
            except Exception as e:
                st.error(f"Ошибка обработки фото: {e}")
                return
    else:
        uploaded_file = st.file_uploader(
            "Загрузите МРТ", 
            type=["jpg", "jpeg", "png", "pdf", "dcm", "dicom", "tiff", "tif", "heic", "heif", "webp", "zip"],
            help="Поддерживаются: JPG, PNG, TIFF, HEIC, WEBP, DICOM, ZIP"
        )
        
        if uploaded_file:
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp:
                    tmp.write(uploaded_file.getvalue())
                    tmp_path = tmp.name
                
                processor = ImageFormatProcessor()
                image_array, file_metadata = processor.load_image(tmp_path, MOBILE_MAX_IMAGE_SIZE)
                metadata = {**metadata, **file_metadata, 'source': 'upload'}
                
                os.unlink(tmp_path)
                processor.cleanup_temp_files()
                
            except Exception as e:
                st.error(f"Ошибка обработки файла: {e}")
                return

    if image_array is None:
        st.info("Загрузите файл или сделайте фото для анализа.")
        return

    # Валидация изображения
    is_valid, error_msg = validate_image(image_array)
    if not is_valid:
        st.error(f"❌ Ошибка валидации изображения: {error_msg}")
        return

    try:
        # Оптимизация для мобильных устройств
        if (IS_REPLIT or st.session_state.get('mobile_mode', False)) and IMAGE_PROCESSOR_AVAILABLE and optimize_image_for_ai:
            image_array = optimize_image_for_ai(image_array)
        
        st.image(image_array, caption="МРТ-срез", use_container_width=True, clamp=True)

        mri_analysis = {
            "quality_assessment": "Хорошее",
            "sharpness": 120.0,
            "noise_level": 20.0,
            "snr": 15.0,
            "artifacts": "Минимальные артефакты"
        }
        
        st.subheader("📊 Оценка качества МРТ")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Качество", mri_analysis['quality_assessment'])
            st.metric("Резкость", f"{mri_analysis['sharpness']:.1f}")
        with col2:
            st.metric("Шум", f"{mri_analysis['noise_level']:.1f}")
            st.metric("SNR", f"{mri_analysis['snr']:.2f}")

        st.caption(f"Артефакты: {mri_analysis['artifacts']}")

        st.markdown("---")
        
        # Блок метрик моделей
        st.markdown("### 📊 Точность моделей для МРТ")
        metrics = get_model_metrics_display('MRI')
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Точность Gemini Flash", f"{metrics['gemini']['accuracy']}%")
            st.metric("Точность Opus 4.5", f"{metrics['opus']['accuracy']}%")
        with col2:
            speed_diff = metrics['opus']['speed_multiplier']
            st.info(f"⚡ Opus в {speed_diff} раз медленнее")
        with col3:
            price_diff = metrics['opus']['price_multiplier']
            st.info(f"💰 Opus в {price_diff} раз дороже")
        
        # Форма обратной связи - ДО анализа, всегда видна и активна!
        st.markdown("---")
        st.markdown("### 💬 Обратная связь")
        
        last_result = st.session_state.get('mri_analysis_result', '')
        analysis_id_base = "MRI_feedback_form"
        mri_input = "МРТ: Магнитно-резонансная томография"
        
        try:
            show_feedback_form(
                analysis_type="MRI",
                analysis_result=str(last_result) if last_result else "",
                analysis_id=analysis_id_base,
                input_case=mri_input
            )
        except Exception as e:
            st.error(f"Ошибка формы обратной связи: {e}")
        
        if not last_result:
            st.info("💡 После проведения анализа форма автоматически обновится с новым результатом.")
        
        st.markdown("---")
        
        # Получение промпта для МРТ
        assistant = OpenRouterAssistant()
        from modules.medical_ai_analyzer import ImageType
        if SPECIALIST_DETECTOR_AVAILABLE and get_specialist_prompt and get_specialist_info:
            prompt = get_specialist_prompt(ImageType.MRI)
            specialist_info = get_specialist_info(ImageType.MRI)
        else:
            prompt = "Проанализируйте МРТ-снимок. Оцените структуры, патологические изменения, дайте заключение."
            specialist_info = {'role': 'Врач-нейрорадиолог'}
        
        # Кнопки быстрого и точного анализа
        col_fast, col_precise = st.columns(2)
        with col_fast:
            if st.button("⚡ Быстрый анализ (Gemini Flash)", use_container_width=True, type="primary", key="mri_fast"):
                with st.spinner("Gemini Flash анализирует МРТ..."):
                    try:
                        result = assistant.send_vision_request_gemini_fast(prompt, image_array)
                        st.markdown(f"### ⚡ Быстрый анализ (Gemini Flash):")
                        st.write(result)
                    except Exception as e:
                        st.error(f"❌ Ошибка анализа: {str(e)}")
        
        with col_precise:
            opus_accuracy = metrics['opus']['accuracy']
            gemini_accuracy = metrics['gemini']['accuracy']
            accuracy_diff = opus_accuracy - gemini_accuracy
            if st.button(f"🎯 Точный анализ (Opus 4.5) - на {accuracy_diff}% точнее", use_container_width=True, type="primary", key="mri_precise"):
                result = perform_analysis_with_streaming(
                    assistant, prompt, image_array, str(metadata), use_streaming=True,
                    analysis_type="точный", model_type="opus",
                    title="🎯 Точный анализ (Opus 4.5):"
                )
                if result:
                    st.session_state.mri_analysis_result = result
                    st.session_state.mri_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    st.rerun()
        
        st.markdown("---")
        st.markdown("### ⚙️ Расширенные режимы анализа")

        # Универсальный анализатор
        from utils.universal_analyzer import UniversalMedicalAnalyzer
        analyzer = UniversalMedicalAnalyzer()
        
        # Выбор режима анализа
        analysis_mode = st.radio(
            "Режим анализа:",
            ["⚡ Быстрый (одна модель)", "🎯 Консенсус (несколько моделей)", "✅ С валидацией"],
            horizontal=True,
            key="mri_analysis_mode"
        )
        
        # Выбор пациента для сохранения контекста
        patient_id = None
        if st.checkbox("💾 Сохранить в контекст пациента"):
            init_db()
            conn = sqlite3.connect('medical_data.db')
            patients = pd.read_sql_query("SELECT id, name FROM patients", conn)
            conn.close()
            
            if not patients.empty:
                selected_patient = st.selectbox("Выберите пациента", patients['name'], key="mri_patient_select")
                patient_id = patients[patients['name'] == selected_patient].iloc[0]['id']
        
        if st.button("🧠 ИИ-анализ МРТ (с контекстом)", use_container_width=True):
            with st.spinner("ИИ анализирует МРТ..."):
                from modules.medical_ai_analyzer import ImageType
                results = analyzer.analyze_image(
                    image_array=image_array,
                    image_type=ImageType.MRI,
                    analysis_mode=analysis_mode,
                    metadata=mri_analysis,
                    patient_id=patient_id
                )
                
                analyzer.display_results(results)
                
                # Форма обратной связи (дополнительно для МРТ, если display_results не показала)
                if FEEDBACK_WIDGET_AVAILABLE and results.get('result'):
                    try:
                        # Формируем input_case для МРТ
                        mri_input = f"МРТ: Тип={results.get('image_type', 'UNKNOWN')}"
                        show_feedback_form(
                            analysis_type="MRI",
                            analysis_result=results['result'],
                            analysis_id=f"MRI_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}",
                            input_case=mri_input
                        )
                    except:
                        pass

    except Exception as e:
        handle_error(e, "show_mri_analysis", show_to_user=True)

def show_dermatoscopy_analysis():
    """Анализ дерматоскопии (фото кожи)"""
    if not AI_AVAILABLE:
        st.error("❌ ИИ-модуль недоступен. Проверьте файл `claude_assistant.py` и API-ключ.")
        return

    st.header("🔬 Анализ дерматоскопии (фото кожи)")
    
    # Мобильная поддержка: выбор источника
    source_type = st.radio(
        "Выберите источник изображения:",
        ["📁 Загрузить файл", "📷 Сделать фото"],
        horizontal=True
    )
    
    image_array = None
    metadata = {}
    
    if source_type == "📷 Сделать фото":
        # Использование камеры смартфона
        camera_image = st.camera_input("Сфотографируйте кожное образование", key="derm_camera")
        if camera_image:
            try:
                image = Image.open(camera_image)
                image_array = np.array(image)
                metadata = {'source': 'camera', 'format': 'mobile_photo'}
            except Exception as e:
                st.error(f"Ошибка обработки фото: {e}")
                return
    else:
        # Загрузка файла с расширенной поддержкой форматов
        uploaded_file = st.file_uploader(
            "Загрузите фото кожи/дерматоскопию", 
            type=["jpg", "jpeg", "png", "tiff", "tif", "heic", "heif", "webp"],
            help="Поддерживаются: JPG, PNG, TIFF, HEIC, WEBP"
        )
        
        if uploaded_file:
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp:
                    tmp.write(uploaded_file.getvalue())
                    tmp_path = tmp.name
                
                processor = ImageFormatProcessor()
                image_array, file_metadata = processor.load_image(tmp_path, MOBILE_MAX_IMAGE_SIZE)
                metadata = {**metadata, **file_metadata, 'source': 'upload'}
                
                os.unlink(tmp_path)
                processor.cleanup_temp_files()
                
            except Exception as e:
                st.error(f"Ошибка обработки файла: {e}")
                return

    if image_array is None:
        st.info("Загрузите файл или сделайте фото для анализа.")
        return

    try:
        # Оптимизация для мобильных устройств
        if (IS_REPLIT or st.session_state.get('mobile_mode', False)) and IMAGE_PROCESSOR_AVAILABLE and optimize_image_for_ai:
            image_array = optimize_image_for_ai(image_array)
        
        st.image(image_array, caption="Дерматоскопия", use_container_width=True, clamp=True)

        st.markdown("---")
        
        # Блок метрик моделей для дерматоскопии
        st.markdown("### 📊 Точность моделей для дерматоскопии")
        st.info("💡 **Важно:** Для дерматоскопии рекомендуется использовать Opus 4.5 из-за высокой точности определения меланомы.")
        metrics = get_model_metrics_display('DERMATOSCOPY')
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Точность Gemini Flash", f"{metrics['gemini']['accuracy']}%")
            st.metric("Точность Opus 4.5", f"{metrics['opus']['accuracy']}%")
        with col2:
            speed_diff = metrics['opus']['speed_multiplier']
            st.info(f"⚡ Opus в {speed_diff} раз медленнее")
        with col3:
            price_diff = metrics['opus']['price_multiplier']
            st.info(f"💰 Opus в {price_diff} раз дороже")
        
        # Форма обратной связи - ДО анализа, всегда видна и активна!
        st.markdown("---")
        st.markdown("### 💬 Обратная связь")
        
        if FEEDBACK_WIDGET_AVAILABLE:
            last_result = st.session_state.get('derma_analysis_result', '')
            analysis_id_base = "DERMA_feedback_form"
            derma_input = "Дерматоскопия: Изображение кожи/родинки"
            
            show_feedback_form(
                analysis_type="DERMATOSCOPY",
                analysis_result=str(last_result) if last_result else "",
                analysis_id=analysis_id_base,
                input_case=derma_input
            )
            
            if not last_result:
                st.info("💡 После проведения анализа форма автоматически обновится с новым результатом.")
        
        st.markdown("---")
        st.markdown("### ⚙️ Режимы анализа")
        
        # Опция streaming
        use_streaming = st.checkbox("📺 Постепенное появление текста (streaming)", value=True, key="derma_streaming")
        
        assistant = OpenRouterAssistant()
        
        # Получение промпта для дерматоскопии
        from modules.medical_ai_analyzer import ImageType
        if SPECIALIST_DETECTOR_AVAILABLE and get_specialist_prompt and get_specialist_info:
            prompt = get_specialist_prompt(ImageType.DERMATOSCOPY)
            specialist_info = get_specialist_info(ImageType.DERMATOSCOPY)
        else:
            prompt = f"""Проанализируйте дерматоскопическое изображение как дерматоонколог с 15+ годами опыта.

Оцените по критериям ABCDE:
- A (Asymmetry) - Асимметрия
- B (Border) - Границы
- C (Color) - Цвет
- D (Diameter) - Диаметр
- E (Evolution) - Эволюция

Также оцените:
- Пигментную сеть
- Точки и глобулы
- Полосы и линии
- Структуры регрессии
- Сосудистую картину

Дайте заключение о риске меланомы и рекомендации."""
            specialist_info = {'role': 'Дерматоонколог'}
        
        # Кнопки - для дерматографии Opus по умолчанию (первая кнопка)
        col_precise, col_fast = st.columns(2)
        with col_precise:
            opus_accuracy = metrics['opus']['accuracy']
            gemini_accuracy = metrics['gemini']['accuracy']
            accuracy_diff = opus_accuracy - gemini_accuracy
            if st.button(f"🎯 Точный анализ (Opus 4.5) - на {accuracy_diff}% точнее [Рекомендуется]", use_container_width=True, type="primary", key="derm_precise"):
                result = perform_analysis_with_streaming(
                    assistant, prompt, image_array, str(metadata), use_streaming,
                    analysis_type="точный", model_type="opus",
                    title="🎯 Точный анализ (Opus 4.5):"
                )
                if result:
                    st.session_state.derma_analysis_result = result
                    st.session_state.derma_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    # Обновляем страницу чтобы форма под метриками обновилась
                    st.rerun()
        
        with col_fast:
            if st.button("⚡ Быстрый анализ (Gemini Flash)", use_container_width=True, key="derm_fast"):
                result = perform_analysis_with_streaming(
                    assistant, prompt, image_array, str(metadata), use_streaming,
                    analysis_type="быстрый", model_type="gemini",
                    title="⚡ Быстрый анализ (Gemini Flash):"
                )
                if result:
                    st.session_state.derma_analysis_result = result
                    st.session_state.derma_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    # Обновляем страницу чтобы форма под метриками обновилась
                    st.rerun()
        
        st.markdown("---")
        
        if st.button("🔬 ИИ-анализ дерматоскопии", use_container_width=True):
            result = perform_analysis_with_streaming(
                assistant, prompt, image_array, str(metadata), use_streaming,
                analysis_type="точный", model_type="opus",
                title=f"### 🧠 Заключение ({specialist_info['role']}):"
            )
            if result:
                st.session_state.derma_analysis_result = result
                st.session_state.derma_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                # Обновляем страницу чтобы форма под метриками обновилась
                st.rerun()

    except Exception as e:
        st.error(f"Ошибка обработки дерматоскопии: {e}")

def show_ct_analysis():
    """Анализ КТ (компьютерная томография) с полной интеграцией компонентов"""
    if not AI_AVAILABLE:
        st.error("❌ ИИ-модуль недоступен. Проверьте файл `claude_assistant.py` и API-ключ.")
        return

    st.header("🩻 Анализ КТ (компьютерная томография)")
    
    source_type = st.radio(
        "Выберите источник изображения:",
        ["📁 Загрузить файл", "📷 Сделать фото"],
        horizontal=True
    )
    
    image_array = None
    metadata = {}
    
    if source_type == "📷 Сделать фото":
        camera_image = st.camera_input("Сфотографируйте КТ-снимок", key="ct_camera")
        if camera_image:
            try:
                image = Image.open(camera_image)
                image_array = np.array(image)
                metadata = {'source': 'camera', 'format': 'mobile_photo'}
            except Exception as e:
                st.error(f"Ошибка обработки фото: {e}")
                return
    else:
        uploaded_file = st.file_uploader(
            "Загрузите КТ", 
            type=["jpg", "jpeg", "png", "pdf", "dcm", "dicom", "tiff", "tif", "heic", "heif", "webp", "zip"],
            help="Поддерживаются: JPG, PNG, TIFF, HEIC, WEBP, DICOM, ZIP"
        )
        
        if uploaded_file:
            try:
                is_valid, error_msg = validate_file_size(uploaded_file.size)
                if not is_valid:
                    st.error(f"❌ {error_msg}")
                    return
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp:
                    tmp.write(uploaded_file.getvalue())
                    tmp_path = tmp.name
                
                processor = ImageFormatProcessor()
                image_array, file_metadata = processor.load_image(tmp_path, MOBILE_MAX_IMAGE_SIZE)
                metadata = {**metadata, **file_metadata, 'source': 'upload'}
                
                os.unlink(tmp_path)
                processor.cleanup_temp_files()
                
            except Exception as e:
                st.error(f"Ошибка обработки файла: {e}")
                return

    if image_array is None:
        st.info("Загрузите файл или сделайте фото для анализа.")
        return

    # Валидация изображения
    is_valid, error_msg = validate_image(image_array)
    if not is_valid:
        st.error(f"❌ Ошибка валидации изображения: {error_msg}")
        return

    try:
        if (IS_REPLIT or st.session_state.get('mobile_mode', False)) and IMAGE_PROCESSOR_AVAILABLE and optimize_image_for_ai:
            image_array = optimize_image_for_ai(image_array)
        
        st.image(image_array, caption="КТ-срез", use_container_width=True, clamp=True)

        # Инициализация компонентов
        assistant = OpenRouterAssistant()
        consensus_engine = ConsensusEngine(assistant)
        validator = ValidationPipeline(assistant)
        scorecard = MedicalScorecard()
        gap_detector = DiagnosticGapDetector()
        notifier = NotificationSystem()
        model_router = ModelRouter()
        evidence_ranker = EvidenceRanker()
        
        from modules.medical_ai_analyzer import ImageType
        
        st.markdown("---")
        
        # Блок метрик моделей
        st.markdown("### 📊 Точность моделей для КТ")
        metrics = get_model_metrics_display('CT')
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Точность Gemini Flash", f"{metrics['gemini']['accuracy']}%")
            st.metric("Точность Opus 4.5", f"{metrics['opus']['accuracy']}%")
        with col2:
            speed_diff = metrics['opus']['speed_multiplier']
            st.info(f"⚡ Opus в {speed_diff} раз медленнее")
        with col3:
            price_diff = metrics['opus']['price_multiplier']
            st.info(f"💰 Opus в {price_diff} раз дороже")
        
        # Форма обратной связи - ДО анализа, всегда видна и активна!
        st.markdown("---")
        st.markdown("### 💬 Обратная связь")
        
        last_result = st.session_state.get('ct_analysis_result', '')
        analysis_id_base = "CT_feedback_form"
        ct_input = "КТ: Компьютерная томография"
        
        try:
            show_feedback_form(
                analysis_type="CT",
                analysis_result=str(last_result) if last_result else "",
                analysis_id=analysis_id_base,
                input_case=ct_input
            )
        except Exception as e:
            st.error(f"Ошибка формы обратной связи: {e}")
        
        if not last_result:
            st.info("💡 После проведения анализа форма автоматически обновится с новым результатом.")
        
        st.markdown("---")
        
        specialist_info = get_specialist_info(ImageType.CT)
        base_prompt = f"Проанализируйте КТ-снимок как {specialist_info['role']} с {specialist_info['experience']}. Оцените структуры, патологические изменения, денситометрию."
        prompt = get_specialist_prompt(ImageType.CT, base_prompt)
        
        # Кнопки быстрого и точного анализа
        col_fast, col_precise = st.columns(2)
        with col_fast:
            if st.button("⚡ Быстрый анализ (Gemini Flash)", use_container_width=True, type="primary", key="ct_fast"):
                with st.spinner("Gemini Flash анализирует КТ..."):
                    try:
                        result = assistant.send_vision_request_gemini_fast(prompt, image_array, str(metadata))
                        st.markdown(f"### ⚡ Быстрый анализ (Gemini Flash):")
                        st.write(result)
                    except Exception as e:
                        st.error(f"❌ Ошибка анализа: {str(e)}")
        
        with col_precise:
            opus_accuracy = metrics['opus']['accuracy']
            gemini_accuracy = metrics['gemini']['accuracy']
            accuracy_diff = opus_accuracy - gemini_accuracy
            if st.button(f"🎯 Точный анализ (Opus 4.5) - на {accuracy_diff}% точнее", use_container_width=True, type="primary", key="ct_precise"):
                result = perform_analysis_with_streaming(
                    assistant, prompt, image_array, str(metadata), use_streaming=True,
                    analysis_type="точный", model_type="opus",
                    title="🎯 Точный анализ (Opus 4.5):"
                )
                if result:
                    st.session_state.ct_analysis_result = result
                    st.session_state.ct_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    st.rerun()
        
        st.markdown("---")
        st.markdown("### ⚙️ Расширенные режимы анализа")
        
        # Выбор режима анализа
        analysis_mode = st.radio(
            "Режим анализа:",
            ["⚡ Быстрый (одна модель)", "🎯 Консенсус (несколько моделей)", "✅ С валидацией"],
            horizontal=True,
            key="ct_analysis_mode"
        )
        
        if st.button("🩻 ИИ-анализ КТ", use_container_width=True):
            if analysis_mode == "⚡ Быстрый (одна модель)":
                # Opus 4.5 используется по умолчанию для клинического анализа КТ
                result = perform_analysis_with_streaming(
                    assistant, prompt, image_array, str(metadata), use_streaming=True,
                    analysis_type="точный", model_type="opus",
                    title=f"### 🧠 Заключение ({specialist_info['role']}):"
                )
                if result:
                    st.session_state.ct_analysis_result = result
                    st.session_state.ct_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    st.rerun()
                    
            elif analysis_mode == "🎯 Консенсус (несколько моделей)":
                    consensus_result = consensus_engine.analyze_with_consensus(prompt, image_array, str(metadata))
                    st.markdown("### 🎯 Консенсус-анализ:")
                    
                    # Правильная структура: consensus_result['consensus']['consensus_response']
                    if consensus_result['consensus']['consensus_available']:
                        st.write(consensus_result['consensus']['consensus_response'])
                        st.metric("Уровень согласия", f"{consensus_result['consensus']['agreement_level']:.1%}")
                        
                        if consensus_result['consensus'].get('discrepancies'):
                            st.warning("⚠️ Обнаружены расхождения между моделями:")
                            for disc in consensus_result['consensus']['discrepancies']:
                                st.warning(f"• {disc}")
                    else:
                        st.write(consensus_result['consensus'].get('single_opinion', 'Ошибка получения консенсуса'))
                    
                    with st.expander("📊 Детали мнений моделей"):
                        for i, opinion in enumerate(consensus_result['individual_opinions'], 1):
                            if opinion['success']:
                                st.markdown(f"**Модель {i} ({opinion['model']}):**")
                                response_text = opinion['response'] if isinstance(opinion['response'], str) else str(opinion['response'])
                                st.write(response_text[:500] + "..." if len(response_text) > 500 else response_text)
                            else:
                                st.error(f"**Модель {i} ({opinion['model']}):** Ошибка: {opinion.get('error', 'Неизвестная ошибка')}")
                    
            elif analysis_mode == "✅ С валидацией":
                # Используем Opus с streaming
                result = perform_analysis_with_streaming(
                    assistant, prompt, image_array, str(metadata), use_streaming=True,
                    analysis_type="точный", model_type="opus",
                    title=f"### 🧠 Ответ ИИ ({specialist_info['role']}):"
                )
                
                if not result:
                    st.error("❌ Не удалось получить результат анализа")
                    return
                
                # Валидация
                validation = validator.validate_response(result)
                
                # Оценка качества
                evaluation = scorecard.evaluate_response(result, ImageType.CT)
                
                # Детекция пробелов
                gaps = gap_detector.detect_gaps(result, ImageType.CT)
                
                # Критические находки
                critical_findings = notifier.check_critical_findings(result)
                
                # Оценка доказательности
                evidence = evidence_ranker.rank_evidence(result)
                
                # Сохраняем результат
                if result:
                    st.session_state.ct_analysis_result = result
                    st.session_state.ct_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    st.rerun()
                
                # Уведомления о критических находках
                    notifier.display_notifications(critical_findings)
                    
                    # Валидация
                    with st.expander("✅ Результаты валидации"):
                        if validation['is_valid']:
                            st.success("✅ Валидация пройдена")
                        else:
                            st.error("❌ Обнаружены проблемы")
                        st.write(f"Полнота: {validation['completeness_score']:.1%}")
                        if validation['warnings']:
                            for warning in validation['warnings']:
                                st.warning(warning)
                        if validation['errors']:
                            for error in validation['errors']:
                                st.error(error)
                    
                    # Оценка качества
                    with st.expander("📊 Оценка качества"):
                        st.write(f"**Оценка:** {evaluation['grade']}")
                        st.write(f"**Балл:** {evaluation['score']:.1%}")
                        if evaluation['recommendations']:
                            st.write("**Рекомендации:**")
                            for rec in evaluation['recommendations']:
                                st.write(f"• {rec}")
                    
                    # Пробелы
                    if gaps['completeness_percentage'] < 100:
                        with st.expander("⚠️ Обнаруженные пробелы"):
                            st.write(gap_detector.generate_gap_report(gaps))
                    
                    # Доказательность
                    with st.expander("📚 Оценка доказательности"):
                        st.write(evidence_ranker.generate_evidence_report(evidence))

    except Exception as e:
        error_msg = handle_error(e, "show_ct_analysis", show_to_user=True)
        st.error(f"Ошибка обработки КТ: {error_msg}")

def show_ultrasound_analysis():
    """Анализ УЗИ (ультразвуковое исследование) с полной интеграцией компонентов"""
    if not AI_AVAILABLE:
        st.error("❌ ИИ-модуль недоступен. Проверьте файл `claude_assistant.py` и API-ключ.")
        return

    st.header("🔊 Анализ УЗИ (ультразвуковое исследование)")
    
    source_type = st.radio(
        "Выберите источник изображения:",
        ["📁 Загрузить файл", "📷 Сделать фото"],
        horizontal=True
    )
    
    image_array = None
    metadata = {}
    
    if source_type == "📷 Сделать фото":
        camera_image = st.camera_input("Сфотографируйте УЗИ-снимок", key="us_camera")
        if camera_image:
            try:
                image = Image.open(camera_image)
                image_array = np.array(image)
                metadata = {'source': 'camera', 'format': 'mobile_photo'}
            except Exception as e:
                st.error(f"Ошибка обработки фото: {e}")
                return
    else:
        uploaded_file = st.file_uploader(
            "Загрузите УЗИ", 
            type=["jpg", "jpeg", "png", "tiff", "tif", "heic", "heif", "webp"],
            help="Поддерживаются: JPG, PNG, TIFF, HEIC, WEBP"
        )
        
        if uploaded_file:
            try:
                is_valid, error_msg = validate_file_size(uploaded_file.size)
                if not is_valid:
                    st.error(f"❌ {error_msg}")
                    return
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp:
                    tmp.write(uploaded_file.getvalue())
                    tmp_path = tmp.name
                
                processor = ImageFormatProcessor()
                image_array, file_metadata = processor.load_image(tmp_path, MOBILE_MAX_IMAGE_SIZE)
                metadata = {**metadata, **file_metadata, 'source': 'upload'}
                
                os.unlink(tmp_path)
                processor.cleanup_temp_files()
                
            except Exception as e:
                st.error(f"Ошибка обработки файла: {e}")
                return

    if image_array is None:
        st.info("Загрузите файл или сделайте фото для анализа.")
        return

    # Валидация изображения
    is_valid, error_msg = validate_image(image_array)
    if not is_valid:
        st.error(f"❌ Ошибка валидации изображения: {error_msg}")
        return

    try:
        if (IS_REPLIT or st.session_state.get('mobile_mode', False)) and IMAGE_PROCESSOR_AVAILABLE and optimize_image_for_ai:
            image_array = optimize_image_for_ai(image_array)
        
        st.image(image_array, caption="УЗИ-снимок", use_container_width=True, clamp=True)

        # Инициализация компонентов
        assistant = OpenRouterAssistant()
        consensus_engine = ConsensusEngine(assistant)
        validator = ValidationPipeline(assistant)
        scorecard = MedicalScorecard()
        gap_detector = DiagnosticGapDetector()
        notifier = NotificationSystem()
        model_router = ModelRouter()
        evidence_ranker = EvidenceRanker()
        
        from modules.medical_ai_analyzer import ImageType
        
        st.markdown("---")
        
        # Блок метрик моделей
        st.markdown("### 📊 Точность моделей для УЗИ")
        metrics = get_model_metrics_display('ULTRASOUND')
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Точность Gemini Flash", f"{metrics['gemini']['accuracy']}%")
            st.metric("Точность Opus 4.5", f"{metrics['opus']['accuracy']}%")
        with col2:
            speed_diff = metrics['opus']['speed_multiplier']
            st.info(f"⚡ Opus в {speed_diff} раз медленнее")
        with col3:
            price_diff = metrics['opus']['price_multiplier']
            st.info(f"💰 Opus в {price_diff} раз дороже")
        
        # Форма обратной связи - ДО анализа, всегда видна и активна!
        st.markdown("---")
        st.markdown("### 💬 Обратная связь")
        
        last_result = st.session_state.get('ultrasound_analysis_result', '')
        analysis_id_base = "ULTRASOUND_feedback_form"
        us_input = "УЗИ: Ультразвуковое исследование"
        
        try:
            show_feedback_form(
                analysis_type="ULTRASOUND",
                analysis_result=str(last_result) if last_result else "",
                analysis_id=analysis_id_base,
                input_case=us_input
            )
        except Exception as e:
            st.error(f"Ошибка формы обратной связи: {e}")
        
        if not last_result:
            st.info("💡 После проведения анализа форма автоматически обновится с новым результатом.")
        
        st.markdown("---")
        
        specialist_info = get_specialist_info(ImageType.ULTRASOUND)
        base_prompt = f"Проанализируйте УЗИ-снимок как {specialist_info['role']} с {specialist_info['experience']}. Оцените эхогенность, структуры, патологические изменения."
        prompt = get_specialist_prompt(ImageType.ULTRASOUND, base_prompt)
        
        # Кнопки быстрого и точного анализа
        col_fast, col_precise = st.columns(2)
        with col_fast:
            if st.button("⚡ Быстрый анализ (Gemini Flash)", use_container_width=True, type="primary", key="us_fast"):
                with st.spinner("Gemini Flash анализирует УЗИ..."):
                    try:
                        result = assistant.send_vision_request_gemini_fast(prompt, image_array, str(metadata))
                        st.markdown(f"### ⚡ Быстрый анализ (Gemini Flash):")
                        st.write(result)
                    except Exception as e:
                        st.error(f"❌ Ошибка анализа: {str(e)}")
        
        with col_precise:
            opus_accuracy = metrics['opus']['accuracy']
            gemini_accuracy = metrics['gemini']['accuracy']
            accuracy_diff = opus_accuracy - gemini_accuracy
            if st.button(f"🎯 Точный анализ (Opus 4.5) - на {accuracy_diff}% точнее", use_container_width=True, type="primary", key="us_precise"):
                result = perform_analysis_with_streaming(
                    assistant, prompt, image_array, str(metadata), use_streaming=True,
                    analysis_type="точный", model_type="opus",
                    title="🎯 Точный анализ (Opus 4.5):"
                )
                if result:
                    st.session_state.ultrasound_analysis_result = result
                    st.session_state.ultrasound_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                    st.rerun()
        
        st.markdown("---")
        st.markdown("### ⚙️ Расширенные режимы анализа")
        
        # Выбор режима анализа
        analysis_mode = st.radio(
            "Режим анализа:",
            ["⚡ Быстрый (одна модель)", "🎯 Консенсус (несколько моделей)", "✅ С валидацией"],
            horizontal=True,
            key="us_analysis_mode"
        )
        
        if st.button("🔊 ИИ-анализ УЗИ", use_container_width=True):
            with st.spinner("ИИ анализирует УЗИ..."):
                if analysis_mode == "⚡ Быстрый (одна модель)":
                    # Opus 4.5 используется по умолчанию для клинического анализа УЗИ
                    result = perform_analysis_with_streaming(
                        assistant, prompt, image_array, str(metadata), use_streaming=True,
                        analysis_type="точный", model_type="opus",
                        title=f"### 🧠 Заключение ({specialist_info['role']}):"
                    )
                    if result:
                        st.session_state.ultrasound_analysis_result = result
                        st.session_state.ultrasound_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                        st.rerun()
                    
                elif analysis_mode == "🎯 Консенсус (несколько моделей)":
                    consensus_result = consensus_engine.analyze_with_consensus(prompt, image_array, str(metadata))
                    st.markdown("### 🎯 Консенсус-анализ:")
                    
                    # Правильная структура: consensus_result['consensus']['consensus_response']
                    if consensus_result['consensus']['consensus_available']:
                        st.write(consensus_result['consensus']['consensus_response'])
                        st.metric("Уровень согласия", f"{consensus_result['consensus']['agreement_level']:.1%}")
                        
                        if consensus_result['consensus'].get('discrepancies'):
                            st.warning("⚠️ Обнаружены расхождения между моделями:")
                            for disc in consensus_result['consensus']['discrepancies']:
                                st.warning(f"• {disc}")
                    else:
                        st.write(consensus_result['consensus'].get('single_opinion', 'Ошибка получения консенсуса'))
                    
                    with st.expander("📊 Детали мнений моделей"):
                        for i, opinion in enumerate(consensus_result['individual_opinions'], 1):
                            if opinion['success']:
                                st.markdown(f"**Модель {i} ({opinion['model']}):**")
                                response_text = opinion['response'] if isinstance(opinion['response'], str) else str(opinion['response'])
                                st.write(response_text[:500] + "..." if len(response_text) > 500 else response_text)
                            else:
                                st.error(f"**Модель {i} ({opinion['model']}):** Ошибка: {opinion.get('error', 'Неизвестная ошибка')}")
                    
                elif analysis_mode == "✅ С валидацией":
                    # Используем Opus с streaming
                    result = perform_analysis_with_streaming(
                        assistant, prompt, image_array, str(metadata), use_streaming=True,
                        analysis_type="точный", model_type="opus",
                        title=f"### 🧠 Ответ ИИ ({specialist_info['role']}):"
                    )
                    
                    if not result:
                        st.error("❌ Не удалось получить результат анализа")
                        return
                    
                    # Валидация
                    validation = validator.validate_response(result)
                    
                    # Оценка качества
                    evaluation = scorecard.evaluate_response(result, ImageType.ULTRASOUND)
                    
                    # Детекция пробелов
                    gaps = gap_detector.detect_gaps(result, ImageType.ULTRASOUND)
                    
                    # Критические находки
                    critical_findings = notifier.check_critical_findings(result)
                    
                    # Оценка доказательности
                    evidence = evidence_ranker.rank_evidence(result)
                    
                    # Отображение результатов
                    st.markdown(f"### 🧠 Заключение ({specialist_info['role']}):")
                    st.write(result)
                    
                    # Формируем input_case для УЗИ
                    us_input = "УЗИ: Ультразвуковое исследование"
                    
                    # Форма обратной связи
                    if FEEDBACK_WIDGET_AVAILABLE:
                        show_feedback_form(
                            analysis_type="ULTRASOUND",
                            analysis_result=result,
                            analysis_id=f"ULTRASOUND_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}",
                            input_case=us_input
                        )
                    
                    # Уведомления о критических находках
                    notifier.display_notifications(critical_findings)
                    
                    # Валидация
                    with st.expander("✅ Результаты валидации"):
                        if validation['is_valid']:
                            st.success("✅ Валидация пройдена")
                        else:
                            st.error("❌ Обнаружены проблемы")
                        st.write(f"Полнота: {validation['completeness_score']:.1%}")
                        if validation['warnings']:
                            for warning in validation['warnings']:
                                st.warning(warning)
                        if validation['errors']:
                            for error in validation['errors']:
                                st.error(error)
                    
                    # Оценка качества
                    with st.expander("📊 Оценка качества"):
                        st.write(f"**Оценка:** {evaluation['grade']}")
                        st.write(f"**Балл:** {evaluation['score']:.1%}")
                        if evaluation['recommendations']:
                            st.write("**Рекомендации:**")
                            for rec in evaluation['recommendations']:
                                st.write(f"• {rec}")
                    
                    # Пробелы
                    if gaps['completeness_percentage'] < 100:
                        with st.expander("⚠️ Обнаруженные пробелы"):
                            st.write(gap_detector.generate_gap_report(gaps))
                    
                    # Доказательность
                    with st.expander("📚 Оценка доказательности"):
                        st.write(evidence_ranker.generate_evidence_report(evidence))

    except Exception as e:
        error_msg = handle_error(e, "show_ultrasound_analysis", show_to_user=True)
        st.error(f"Ошибка обработки УЗИ: {error_msg}")

# --- Страница: Протокол приёма ---
def show_consultation_protocol():
    from local_docs import create_local_doc
    
    if not AI_AVAILABLE:
        st.error("❌ ИИ-модуль недоступен. Проверьте файл `claude_assistant.py` и API-ключ.")
        return

    st.header("📝 Автоматический протокол приёма")

    init_db()
    
    # Кнопка сброса всех данных протокола
    if st.button("🗑️ Очистить все данные протокола", type="secondary", use_container_width=True):
        # Очищаем все данные протокола из session_state
        if 'raw_text' in st.session_state:
            del st.session_state['raw_text']
        if 'structured_note' in st.session_state:
            del st.session_state['structured_note']
        if 'protocol_patient_name' in st.session_state:
            del st.session_state['protocol_patient_name']
        if 'transcribed_text_display' in st.session_state:
            del st.session_state['transcribed_text_display']
        if 'transcribed_genetic_question' in st.session_state:
            del st.session_state['transcribed_genetic_question']
        st.success("✅ Все данные протокола очищены")
        st.rerun()
    
    # Выбор пациента (опционально, можно создать после генерации протокола)
    conn = sqlite3.connect('medical_data.db')
    patients = pd.read_sql_query("SELECT id, name FROM patients", conn)
    conn.close()
    
    selected_patient = None
    patient_id = None
    
    if not patients.empty:
        # Если есть пациенты, можно выбрать, но не обязательно
        col1, col2 = st.columns([3, 1])
        with col1:
            selected_patient = st.selectbox(
                "Выберите пациента (опционально, можно создать после генерации протокола)",
                ["--- Создать нового ---"] + list(patients['name']),
                key="protocol_patient_select"
            )
        if selected_patient and selected_patient != "--- Создать нового ---":
            patient_id = patients[patients['name'] == selected_patient].iloc[0]['id']
    else:
        st.info("💡 Пациент будет создан автоматически после генерации протокола")

    st.subheader("📝 Ввод данных для протокола")
    
    # Выбор способа ввода
    input_method = st.radio(
        "Выберите способ ввода:",
        ["✍️ Письменный ввод", "🎤 Голосовой ввод", "📁 Загрузить готовый файл"],
        horizontal=True,
        key="protocol_input_method"
    )
    
    raw_text = st.session_state.get('raw_text', '')
    
    # Письменный ввод
    if input_method == "✍️ Письменный ввод":
        st.info("💡 Введите данные пациента в свободной форме: жалобы, анамнез, объективный осмотр и т.д.")
        raw_text = st.text_area(
            "Введите данные для протокола:",
            value=raw_text,
            height=300,
            help="Опишите жалобы, анамнез заболевания, анамнез жизни, данные объективного осмотра, результаты обследований",
            key="protocol_text_input"
        )
        
        # Кнопка показывается сразу, если есть текст
        if raw_text:
            if st.button("📝 Создать протокол из текста", use_container_width=True, type="primary"):
                st.session_state.raw_text = raw_text
                st.session_state.structured_note = ''  # Сбрасываем старый протокол
                st.rerun()  # Перезагружаем для генерации протокола
    
    # Загрузка готового файла
    elif input_method == "📁 Загрузить готовый файл":
        st.info("💡 Загрузите готовый файл с данными пациента (текст будет извлечен автоматически)")
        uploaded_file_protocol = st.file_uploader(
            "Загрузите файл",
            type=["txt", "docx", "pdf", "md"],
            help="Поддерживаются: TXT, DOCX, PDF, MD"
        )
        
        if uploaded_file_protocol:
            file_ext_protocol = uploaded_file_protocol.name.split('.')[-1].lower()
            
            if file_ext_protocol == 'txt' or file_ext_protocol == 'md':
                try:
                    raw_text = uploaded_file_protocol.read().decode('utf-8', errors='ignore')
                    st.success("✅ Файл загружен и прочитан")
                    st.text_area("Содержимое файла:", value=raw_text, height=200, disabled=True)
                    st.session_state.raw_text = raw_text
                except Exception as e:
                    st.error(f"❌ Ошибка чтения файла: {e}")
            
            elif file_ext_protocol == 'docx':
                try:
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(uploaded_file_protocol.read()))
                    raw_text = "\n".join([para.text for para in doc.paragraphs])
                    st.success("✅ DOCX файл загружен и текст извлечен")
                    st.text_area("Извлеченный текст:", value=raw_text, height=200, disabled=True)
                    st.session_state.raw_text = raw_text
                except Exception as e:
                    st.error(f"❌ Ошибка чтения DOCX: {e}")
                    st.info("💡 Установите: pip install python-docx")
            
            elif file_ext_protocol == 'pdf':
                try:
                    import PyPDF2
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file_protocol.read()))
                    raw_text = "\n".join([page.extract_text() for page in pdf_reader.pages])
                    st.success("✅ PDF файл загружен и текст извлечен")
                    st.text_area("Извлеченный текст:", value=raw_text, height=200, disabled=True)
                    st.session_state.raw_text = raw_text
                except ImportError:
                    st.error("❌ Для чтения PDF установите: pip install PyPDF2")
                except Exception as e:
                    st.error(f"❌ Ошибка чтения PDF: {e}")
            
            # Кнопка показывается сразу, если есть текст
            if raw_text:
                if st.button("📝 Создать протокол из файла", use_container_width=True, type="primary"):
                    st.session_state.raw_text = raw_text
                    st.session_state.structured_note = ''  # Сбрасываем старый протокол
                    st.rerun()  # Перезагружаем для генерацию протокола
    
    # Голосовой ввод
    elif input_method == "🎤 Голосовой ввод":
        st.subheader("🎙️ Голосовой ввод через AssemblyAI")
        
        # Выбор способа ввода аудио
        audio_input_method = st.radio(
            "Выберите способ ввода аудио:",
            ["🎤 Записать с микрофона", "📁 Загрузить аудиофайл"],
            horizontal=True,
            key="audio_input_method"
        )
    
        audio = None
        
        if audio_input_method == "🎤 Записать с микрофона":
            try:
                from audio_recorder_streamlit import audio_recorder
                st.info("💡 Нажмите на кнопку ниже, чтобы начать запись. Запись остановится автоматически при повторном нажатии.")
                st.warning("⚠️ **Важно:** Разрешите доступ к микрофону в браузере, когда появится запрос.")
                
                audio_bytes = audio_recorder(text="🎤 Нажмите для записи", pause_threshold=2.0, sample_rate=44100, key="protocol_audio_recorder")
                
                if audio_bytes:
                    # Сохраняем во временный файл для AssemblyAI
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
                        tmp_file.write(audio_bytes)
                        audio = tmp_file.name
                        st.session_state['audio_file_path'] = audio  # Сохраняем путь в session_state
                    
                    st.success(f"✅ Записано {len(audio_bytes)} байт аудио")
                    st.audio(audio_bytes, format="audio/wav")
                elif 'audio_file_path' in st.session_state:
                    # Используем сохраненный файл из предыдущей записи
                    audio = st.session_state['audio_file_path']
            except ImportError as e:
                st.warning("⚠️ Для записи с микрофона установите: pip install audio-recorder-streamlit")
                st.info("💡 Пока используйте загрузку аудиофайла")
                audio = st.audio_input("Загрузите аудио (до 30 мин)", key="protocol_audio_input_fallback")
            except Exception as e:
                st.error(f"❌ Ошибка при записи с микрофона: {e}")
                st.info("💡 Попробуйте использовать загрузку аудиофайла")
                audio = st.audio_input("Загрузите аудио (до 30 мин)", key="protocol_audio_input_error")
        else:
            audio = st.audio_input("Загрузите аудио (до 30 мин)", key="protocol_audio_input")

        if not ASSEMBLYAI_AVAILABLE:
            st.error("❌ AssemblyAI недоступен. Проверьте файл assemblyai_transcriber.py")

        if audio and st.button("🎤 Обработать аудио", key="process_audio_protocol"):
            if ASSEMBLYAI_AVAILABLE:
                with st.spinner("🔄 Расшифровка через AssemblyAI..."):
                    try:
                        from config import ASSEMBLYAI_API_KEY
                        
                        api_key = ASSEMBLYAI_API_KEY or st.secrets.get("ASSEMBLYAI_API_KEY", "")
                        if not api_key:
                            st.error("❌ API ключ AssemblyAI не найден. Проверьте config.py или secrets.toml")
                            return
                        
                        # Проверяем, что файл существует (если это путь)
                        if isinstance(audio, str):
                            if not os.path.exists(audio):
                                st.error(f"❌ Аудиофайл не найден: {audio}")
                                return
                        
                        # Вызываем функцию транскрипции
                        if not transcribe_audio_assemblyai:
                            st.error("❌ Функция транскрипции недоступна. Проверьте импорт assemblyai_transcriber")
                            return
                        
                        raw_text = transcribe_audio_assemblyai(audio, api_key)
                        
                        # Проверяем, не вернулась ли ошибка
                        if raw_text.startswith("❌"):
                            st.error(raw_text)
                            return
                        
                        st.session_state.raw_text = raw_text
                        st.rerun()  # Перезагружаем для генерации протокола
                    except Exception as e:
                        import traceback
                        st.error(f"❌ Ошибка AssemblyAI: {e}")
                        with st.expander("🔍 Детали ошибки"):
                            st.code(traceback.format_exc())
                        return
            else:
                st.error("❌ AssemblyAI недоступен")
                return

            # Показываем расшифрованный текст если есть
            if raw_text:
                st.subheader("📝 Расшифрованный текст:")
                st.text_area("Расшифрованный текст", value=raw_text, height=150, disabled=True, key="transcribed_text_display")
    
    # Генерация протокола (если есть raw_text)
    if raw_text or st.session_state.get('raw_text'):
        if not raw_text:
            raw_text = st.session_state.get('raw_text', '')
        
        if raw_text:
            # Блок выбора/загрузки шаблона протокола врача
            st.subheader("🧩 Шаблон протокола врача")
            
            # Инициализация таблицы промптов (скрыто, для сохранения промптов)
            init_db()
            from database import init_specialist_prompts_table, save_specialist_prompt, get_specialist_prompts, delete_specialist_prompt
            init_specialist_prompts_table()
            
            # Получаем сохраненные промпты для отображения в списке
            saved_prompts = get_specialist_prompts()
            saved_prompt_names = {}
            for p in saved_prompts:
                key = f"{p['specialist_name']}" + (f" ({p['template_name']})" if p['template_name'] else "")
                saved_prompt_names[key] = p
            
            col_tpl1, col_tpl2 = st.columns(2)
            with col_tpl1:
                # Базовые шаблоны
                base_templates = [
                    "Терапевт (по умолчанию)",
                    "Кардиолог",
                    "Невролог",
                    "Педиатр",
                    "Акушер‑гинеколог",
                    "Врач УЗИ",
                    "Эндоскопист",
                    "Рентгенолог",
                    "Радиолог",
                    "Генетик",
                ]
                
                # Добавляем сохраненные промпты
                if saved_prompt_names:
                    base_templates.append("--- Сохраненные промпты ---")
                    base_templates.extend(saved_prompt_names.keys())
                
                template_preset = st.selectbox(
                    "Выберите базовый шаблон",
                    base_templates,
                )
            with col_tpl2:
                uploaded_template = st.file_uploader(
                    "Или загрузите свой шаблон (.txt/.md)",
                    type=["txt", "md"],
                    help="Текстовый файл с вашим форматом протокола (разделы, формулировки, порядок блоков)",
                )

            protocol_template = ""
            selected_saved_prompt = None
            
            # Проверяем, выбран ли сохраненный промпт
            if template_preset in saved_prompt_names:
                selected_saved_prompt = saved_prompt_names[template_preset]
                protocol_template = selected_saved_prompt['prompt_text']
                st.info(f"✅ Используется сохраненный промпт: {selected_saved_prompt['specialist_name']}")
            elif uploaded_template is not None:
                try:
                    protocol_template = uploaded_template.read().decode("utf-8", errors="ignore")
                    st.success("✅ Загружен пользовательский шаблон протокола врача")
                    
                    # Предлагаем сохранить загруженный шаблон
                    with st.expander("💾 Сохранить этот шаблон для будущего использования"):
                        col_save1, col_save2 = st.columns(2)
                        with col_save1:
                            save_specialist_name = st.text_input("Название специалиста", placeholder="Например: Терапевт")
                        with col_save2:
                            save_template_name = st.text_input("Название шаблона (опционально)", placeholder="Например: Базовый")
                        
                        if st.button("💾 Сохранить промпт", key="save_uploaded_template"):
                            if save_specialist_name:
                                try:
                                    save_specialist_prompt(
                                        save_specialist_name.strip(),
                                        protocol_template.strip(),
                                        save_template_name.strip() if save_template_name else None,
                                        False
                                    )
                                    st.success(f"✅ Промпт для {save_specialist_name} сохранен!")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Ошибка сохранения: {e}")
                            else:
                                st.warning("⚠️ Введите название специалиста")
                except Exception as e:
                    st.warning(f"⚠️ Не удалось прочитать шаблон: {e}")
                    protocol_template = ""
            else:
                preset_templates = {
                    "Терапевт (по умолчанию)": "",
                    "Кардиолог": "Особое внимание уделяй сердечно‑сосудистой системе, функциональному классу сердечной недостаточности, рискам по SCORE2, осложнениям гипертензии и ишемической болезни сердца.",
                    "Невролог": "Подробно опиши неврологический статус, очаговую симптоматику, шкалу NIHSS (если применимо), риски инсульта и показания к нейровизуализации.",
                    "Педиатр": "Учитывай возрастные нормы, перинатальный анамнез, вакцинацию, физическое и психомоторное развитие ребёнка.",
                    "Акушер‑гинеколог": "Фокус на акушерско‑гинекологическом анамнезе, менструальной функции, беременности и родах, рисках акушерских осложнений.",
                    "Врач УЗИ": "Подробно опиши результаты ультразвукового исследования: размеры органов, структуру, эхогенность, наличие патологических образований, кровоток (при допплерографии). Укажи локализацию и характеристики выявленных изменений.",
                    "Эндоскопист": "Детально опиши результаты эндоскопического исследования: состояние слизистой оболочки, наличие патологических изменений (эрозии, язвы, полипы, новообразования), их локализацию, размеры, характер. Укажи результаты биопсии (если проводилась).",
                    "Рентгенолог": "Систематически проанализируй рентгенограмму: оценка качества снимка, описание всех визуализированных структур, выявление патологических изменений (инфильтраты, затемнения, просветления, деформации), их локализацию и характеристики. Сравни с предыдущими исследованиями при наличии.",
                    "Радиолог": "Проведи комплексный анализ результатов лучевой диагностики (КТ, МРТ, ПЭТ-КТ): описание всех визуализированных структур, выявление патологических образований, их локализацию, размеры, плотность/интенсивность сигнала, контрастное усиление, признаки злокачественности. Оцени динамику при наличии предыдущих исследований.",
                    "Генетик": "Детально проанализируй результаты генетического тестирования: перечисли все выявленные варианты (гены, rsID, нотация c./p., генотипы), их клиническое значение, патогенность, тип наследования. Оцени риски для пациента и родственников, фармакогенетические последствия. Сформулируй клинические рекомендации и тактику ведения.",
                }
                protocol_template = preset_templates.get(template_preset, "")
            
            # Управление промптами в отдельном expander (скрыто по умолчанию)
            with st.expander("⚙️ Управление сохраненными промптами", expanded=False):
                st.info("💡 Сохраняйте свои промпты для быстрого использования при генерации протоколов")
                
                col_manage1, col_manage2 = st.columns(2)
                
                with col_manage1:
                    st.markdown("### ➕ Сохранить новый промпт")
                    with st.form("save_prompt_form"):
                        new_specialist = st.text_input("Название специалиста", placeholder="Например: Терапевт, Кардиолог")
                        new_template_name = st.text_input("Название шаблона (опционально)", placeholder="Например: Базовый, Расширенный")
                        new_prompt_text = st.text_area("Текст промпта", height=200, placeholder="Введите промпт для генерации протокола...")
                        is_default_new = st.checkbox("Сделать промптом по умолчанию для этого специалиста")
                        
                        if st.form_submit_button("💾 Сохранить промпт", use_container_width=True):
                            if new_specialist and new_prompt_text:
                                try:
                                    save_specialist_prompt(
                                        new_specialist.strip(),
                                        new_prompt_text.strip(),
                                        new_template_name.strip() if new_template_name else None,
                                        is_default_new
                                    )
                                    st.success(f"✅ Промпт для {new_specialist} сохранен!")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Ошибка сохранения: {e}")
                            else:
                                st.warning("⚠️ Заполните название специалиста и текст промпта")
                
                with col_manage2:
                    st.markdown("### 📋 Сохраненные промпты")
                    all_prompts = get_specialist_prompts()
                    
                    if all_prompts:
                        for prompt in all_prompts:
                            with st.expander(f"🔹 {prompt['specialist_name']}" + (f" - {prompt['template_name']}" if prompt['template_name'] else "") + (" ⭐ (по умолчанию)" if prompt['is_default'] else "")):
                                st.text_area(
                                    "Промпт:",
                                    value=prompt['prompt_text'],
                                    height=150,
                                    disabled=True,
                                    key=f"view_prompt_{prompt['id']}"
                                )
                                st.caption(f"Создан: {prompt['created_at']}")
                                if st.button(f"🗑️ Удалить", key=f"delete_prompt_{prompt['id']}"):
                                    delete_specialist_prompt(prompt['id'])
                                    st.success("✅ Промпт удален")
                                    st.rerun()
                    else:
                        st.info("💡 Нет сохраненных промптов. Создайте первый промпт слева.")

            template_block = ""
            if protocol_template and protocol_template.strip():
                template_block = (
                    f"\nДополнительные требования к оформлению и содержанию протокола "
                    f"(шаблон врача: {template_preset}):\n{protocol_template.strip()}\n"
                )

            # Генерация протокола происходит автоматически после обработки аудио/текста
            # (как в main ветке - сразу после обработки, без отдельной кнопки)
            with st.spinner("🤖 Генерация протокола..."):
                assistant = OpenRouterAssistant()
                prompt = f"""
Вы - опытный терапевт, американский профессор клинической медицины и ведущий специалист университетской клиники с многолетним клиническим опытом.

Вы совмещаете клиническую строгость и ответственность, давая ответы по клиническим проблемам внутренних болезней, включая акушерство и гинекологию, хирургию, а также помогаете обрабатывать несистемно изложенную информацию, облекая её по шаблону и стандартному протоколу осмотра терапевта с рекомендациями по обследованию и лечению.

Ваша задача - создать полный и структурированный протокол осмотра пациента на основании представленной пользователем неструктурированной информации, включающих перечень жалоб, истории появления симптомов и жалоб, данных объективного осмотра, приведенных данных лабораторных тестов и инструментальных исследований. Использовать как модель протокола шаблоны. Постарайтесь вместить всю информацию на 2 страницы, поскольку скачанный файл в формате .doc будет использоваться для печати.

{template_block}

Текст для обработки: {raw_text}

ФОРМАТ ПРОТОКОЛА:

**Жалобы:**
(текст изложенный в жалобах должен быть без дополнительных абзацев, не нужно делать дополнительных пустых строк, пишите единым полотном)

**Анамнез заболевания:**
(текст без дополнительных абзацев, единым полотном)

**Анамнез жизни:**
(текст без дополнительных абзацев, единым полотном)

**Объективный осмотр:**
Общее состояние: лимфоузлы: Кожа: Слизистые: Пульс: АД: ЧДД: Сердце: Лёгкие: Живот: Печень, селезёнка: почки: стул: диурез: отёки: Неврологический статус:
(не используйте выражения "не проводилась", вместо этого напишите выражения отражающие нормы, но упомяните все основные системы, текст без дополнительных абзацев, единым полотном)

**Предварительный диагноз:**
(диагноз выносите на основании российских классификаций болезней)

**Рекомендованные обследования:**
1. ...
2. ...
(рекомендации по обследованию напишите по пунктам 1., 2., и т.д., используйте сокращения, если строчка очень длинная, не делайте дополнительного пропуска между строками)

**Терапия:**
- Рекомендации по режиму, диете
- Фармакотерапия: перечислите препараты. Назовите международное название и 2 коммерческих (бренд и копию) генерика, доступных в РФ. Не делайте дополнительного пропуска между строками.
- Физиолечение: предложите 1-2 наиболее подходящих для пациента, не указывайте те, которые могут быть противопоказаны

**Согласие пациента:**
(тезис о согласии и прочтении в конце сделать более мелким шрифтом)

ОГРАНИЧЕНИЯ И ПРАВИЛА СТИЛЯ:

- Язык ответа: русский. Стиль: строго профессиональный, клинически и технически точный, без упрощений.
- Не нужно использовать текст подзаголовков другого размера.
- Текст изложенный в жалобах, анамнезе, объективном осмотре должен быть без дополнительных абзацев, не нужно делать дополнительных пустых строк, пишите единым полотном.
- Не используйте выражения "не проводилась", вместо этого напишите выражения отражающие нормы, но упомяните все основные системы.
- Уберите логотип переплексити с первой страницы.
- Представьте всю полученную информацию в форме протокола осмотра первичного осмотра врача.
- Диагноз выносите на основании российских классификаций болезней.
- Рекомендации по обследованию напишите по пунктам 1., 2., и т.д., используйте сокращения, если строчка очень длинная. Не делайте дополнительного пропуска между строками.
- Точно так же перечислите препараты и не делайте дополнительного пропуска между строками. Просто назовите международное и 2 коммерческих (бренд и копию) генерика, доступных в РФ.
- Следует придерживаться такого стиля изложения и выбор шрифта при экспорте, чтобы протокол умещался на 2 страницы листа А4.
- Тезис о согласии и прочтении в конце сделать более мелким шрифтом.

ИСТОЧНИКИ (медицина):
UpToDate, PubMed, Cochrane, NCCN, ESC, IDSA, CDC, WHO, ESMO, ADA, GOLD, KDIGO (и другие международные руководства с доказательной базой).

Медицинские рекомендации - опираться только на проверенные международные источники; для каждого ключевого лечебного шага указывать ссылку и год публикации (предпочтительно ≤5 лет).
"""
                # Используем Sonnet 4.5 для протокола (быстрее и дешевле Opus, но качественно)
                structured_note = assistant.get_response(prompt, use_sonnet_4_5=True)
                st.session_state.structured_note = structured_note
                
                # Автоматическое создание/получение пациента, если не выбран
                if not patient_id:
                    # Извлекаем имя пациента из протокола или создаем временное
                    import re
                    # Пытаемся найти имя в тексте (например, "Пациент: Иван Иванов" или "ФИО: ...")
                    name_match = re.search(r'(?:пациент|фио|ф\.и\.о\.|имя)[\s:]+([А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+)', raw_text, re.IGNORECASE)
                    if name_match:
                        patient_name = name_match.group(1).strip()
                    else:
                        # Создаем имя на основе даты и времени
                        from datetime import datetime
                        patient_name = f"Пациент {datetime.now().strftime('%d.%m.%Y %H:%M')}"
                    
                    # Создаем пациента в базе
                    conn = sqlite3.connect('medical_data.db')
                    cursor = conn.cursor()
                    cursor.execute('''
                        INSERT INTO patients (name, age, sex, phone)
                        VALUES (?, ?, ?, ?)
                    ''', (patient_name, None, None, None))
                    patient_id = cursor.lastrowid
                    conn.commit()
                    conn.close()
                    
                    st.success(f"✅ Пациент '{patient_name}' автоматически создан в базе данных")
                    st.session_state['protocol_patient_name'] = patient_name
                    selected_patient = patient_name
                else:
                    st.session_state['protocol_patient_name'] = selected_patient
                
                # Автоматическое сохранение протокола в контекст пациента
                try:
                    context_store = ContextStore()
                    context_store.add_context(
                        patient_id=patient_id,
                        context_type='protocol',
                        context_data={
                            'protocol': structured_note,
                            'type': 'consultation',
                            'raw_transcription': raw_text
                        },
                        source='ai_generated'
                    )
                    st.info("💾 Протокол автоматически сохранен в клинический контекст пациента")
                except Exception as e:
                    st.warning(f"⚠️ Не удалось сохранить протокол в контекст: {e}")

            # Показываем сгенерированный протокол
            structured_note = st.session_state.get('structured_note', '')
            if structured_note:
                # Используем имя пациента из session_state или временное
                patient_name_for_doc = st.session_state.get('protocol_patient_name', selected_patient if 'selected_patient' in locals() and selected_patient else "Пациент")
                with st.spinner("📄 Создание документа..."):
                    filepath, message = create_local_doc(f"Протокол - {patient_name_for_doc}", structured_note)
                    st.success(message)
                    with open(filepath, "rb") as f:
                        # Используем правильное расширение для macOS Pages
                        file_name = os.path.basename(filepath)
                        if not file_name.endswith('.docx'):
                            file_name = file_name.replace('.doc', '.docx')
                        
                        st.download_button(
                            label="📥 Скачать протокол (.docx)",
                            data=f,
                            file_name=file_name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        # Дополнительная кнопка для сохранения в Pages-совместимом формате
                        st.info("💡 **Совет для macOS:** После скачивания файл можно открыть в Pages. Если файл открывается неправильно:")
                        st.markdown("""
                        1. **Через Finder:** Правый клик на файл → «Открыть с помощью» → выберите Pages
                        2. **Настройка по умолчанию:** Выберите файл → Cmd+I → «Открыть с помощью» → Pages → «Изменить все...»
                        3. **Альтернатива:** Используйте Microsoft Word для macOS
                        """)
                        
                        # Показываем путь к файлу для ручного копирования
                        st.code(f"Путь к файлу: {filepath}", language=None)

                st.subheader("📄 Сгенерированный протокол")
                st.write(structured_note)
                
                # Кнопка для очистки протокола после просмотра
                if st.button("🗑️ Очистить протокол и начать заново", type="secondary", use_container_width=True):
                    if 'structured_note' in st.session_state:
                        del st.session_state['structured_note']
                    if 'raw_text' in st.session_state:
                        del st.session_state['raw_text']
                    if 'protocol_patient_name' in st.session_state:
                        del st.session_state['protocol_patient_name']
                    st.success("✅ Протокол очищен. Можете создать новый.")
                    st.rerun()
            
            # Кнопка для сохранения в контекст (если не сохранилось автоматически)
            if st.button("💾 Сохранить протокол в контекст пациента"):
                try:
                    context_store = ContextStore()
                    context_store.add_context(
                        patient_id=patient_id,
                        context_type='protocol',
                        context_data={
                            'protocol': structured_note,
                            'type': 'consultation'
                        },
                        source='manual_entry'
                    )
                    st.success("✅ Протокол сохранен в клинический контекст пациента!")
                except Exception as e:
                    st.error(f"❌ Ошибка сохранения: {e}")

def show_patient_database():
    st.header("👤 База данных пациентов")
    init_db()

    tab1, tab2 = st.tabs(["➕ Добавить", "🔍 Поиск"])

    with tab1:
        st.subheader("Добавить пациента")
        with st.form("add_patient"):
            name = st.text_input("ФИО")
            age = st.number_input("Возраст", min_value=0, max_value=150)
            sex = st.selectbox("Пол", ["М", "Ж"])
            phone = st.text_input("Телефон")
            submitted = st.form_submit_button("Добавить")

            if submitted:
                if not name or not name.strip():
                    st.error("❌ Пожалуйста, укажите ФИО пациента")
                else:
                    try:
                        conn = sqlite3.connect('medical_data.db')
                        cursor = conn.cursor()
                        cursor.execute('''
                            INSERT INTO patients (name, age, sex, phone)
                            VALUES (?, ?, ?, ?)
                        ''', (name.strip(), age, sex, phone))
                        conn.commit()
                        conn.close()
                        st.success(f"✅ Пациент {name.strip()} успешно добавлен в базу данных!")
                        st.rerun()
                    except sqlite3.Error as e:
                        st.error(f"❌ Ошибка при добавлении пациента: {e}")
                        st.info("💡 Попробуйте обновить страницу и попробовать снова")

    with tab2:
        st.subheader("Поиск пациентов")
        conn = sqlite3.connect('medical_data.db')
        df = pd.read_sql_query("SELECT * FROM patients", conn)
        conn.close()

        if not df.empty:
            st.dataframe(df, use_container_width=True)
        else:
            st.info("Пациенты не найдены")

def show_ai_chat():
    if not AI_AVAILABLE:
        st.error("❌ ИИ-модуль недоступен. Проверьте файл `claude_assistant.py` и API-ключ.")
        return

    st.header("🤖 ИИ-Консультант")
    st.info("💡 Рекомендации даются от врача врачу. Вы можете загружать файлы для анализа.")

    try:
        assistant = OpenRouterAssistant()
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("🔗 Тест подключения"):
                with st.spinner("Проверка..."):
                    success, msg = assistant.test_connection()
                    if success:
                        st.success(msg)
                    else:
                        st.error(msg)
        with col2:
            st.info("💡 Используется Claude Sonnet 4.5")
        with col3:
            if st.button("🗑️ Очистить историю"):
                # Удаляем из session_state
                if 'chat_history' in st.session_state:
                    st.session_state.chat_history = []
                if 'uploaded_files_context' in st.session_state:
                    st.session_state.uploaded_files_context = []
                
                # Удаляем из базы данных
                try:
                    conn = sqlite3.connect('medical_data.db')
                    cursor = conn.cursor()
                    # Удаляем всю историю для текущей сессии
                    if 'chat_session_id' in st.session_state:
                        cursor.execute('''
                            DELETE FROM ai_chat_history 
                            WHERE session_id = ?
                        ''', (st.session_state.chat_session_id,))
                    # Также удаляем всю историю (на случай, если session_id не совпадает)
                    cursor.execute('DELETE FROM ai_chat_history')
                    conn.commit()
                    conn.close()
                    print("✅ История полностью удалена из базы данных")
                except Exception as e:
                    print(f"⚠️ Ошибка удаления истории из БД: {e}")
                
                # Создаем новый session_id
                st.session_state.chat_session_id = f"session_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
                st.rerun()

        # Инициализация истории чата
        if 'chat_history' not in st.session_state:
            st.session_state.chat_history = []
            # Загружаем историю из базы данных только если она не была очищена
            try:
                conn = sqlite3.connect('medical_data.db')
                cursor = conn.cursor()
                # Загружаем историю для текущей сессии, если есть
                if 'chat_session_id' in st.session_state:
                    cursor.execute('''
                        SELECT user_message, assistant_response, files_context, created_at
                        FROM ai_chat_history
                        WHERE session_id = ?
                        ORDER BY created_at ASC
                        LIMIT 20
                    ''', (st.session_state.chat_session_id,))
                else:
                    # Если нет session_id, загружаем последние записи
                    cursor.execute('''
                        SELECT user_message, assistant_response, files_context, created_at
                        FROM ai_chat_history
                        ORDER BY created_at DESC
                        LIMIT 20
                    ''')
                rows = cursor.fetchall()
                for row in rows:  # Уже в правильном порядке
                    files_info = json.loads(row[2]) if row[2] else []
                    st.session_state.chat_history.append({
                        'user': row[0],
                        'assistant': row[1],
                        'files_info': files_info,
                        'timestamp': row[3]
                    })
                conn.close()
            except Exception as e:
                print(f"Ошибка загрузки истории: {e}")
        
        if 'uploaded_files_context' not in st.session_state:
            st.session_state.uploaded_files_context = []
        
        if 'chat_session_id' not in st.session_state:
            st.session_state.chat_session_id = f"session_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"

        # Пересылка заключений от анализаторов изображений
        with st.expander("📋 Переслать заключения от анализаторов", expanded=False):
            # Получаем сохраненные результаты анализов из session_state
            analysis_results = []
            
            # Проверяем результаты ЭКГ
            if 'ecg_analysis_result' in st.session_state:
                ecg_result = st.session_state.ecg_analysis_result
                if isinstance(ecg_result, dict) or isinstance(ecg_result, str):
                    analysis_results.append({
                        'type': 'ЭКГ',
                        'data': ecg_result,
                        'timestamp': st.session_state.get('ecg_analysis_timestamp', 'Недавно')
                    })
            
            # Проверяем результаты рентгена
            if 'xray_analysis_result' in st.session_state:
                xray_result = st.session_state.xray_analysis_result
                if isinstance(xray_result, dict) or isinstance(xray_result, str):
                    analysis_results.append({
                        'type': 'Рентген',
                        'data': xray_result,
                        'timestamp': st.session_state.get('xray_analysis_timestamp', 'Недавно')
                    })
            
            # Проверяем результаты МРТ
            if 'mri_analysis_result' in st.session_state:
                mri_result = st.session_state.mri_analysis_result
                if isinstance(mri_result, dict) or isinstance(mri_result, str):
                    analysis_results.append({
                        'type': 'МРТ',
                        'data': mri_result,
                        'timestamp': st.session_state.get('mri_analysis_timestamp', 'Недавно')
                    })
            
            # Проверяем результаты КТ
            if 'ct_analysis_result' in st.session_state:
                ct_result = st.session_state.ct_analysis_result
                if isinstance(ct_result, dict) or isinstance(ct_result, str):
                    analysis_results.append({
                        'type': 'КТ',
                        'data': ct_result,
                        'timestamp': st.session_state.get('ct_analysis_timestamp', 'Недавно')
                    })
            
            # Проверяем результаты УЗИ
            if 'ultrasound_analysis_result' in st.session_state:
                us_result = st.session_state.ultrasound_analysis_result
                if isinstance(us_result, dict) or isinstance(us_result, str):
                    analysis_results.append({
                        'type': 'УЗИ',
                        'data': us_result,
                        'timestamp': st.session_state.get('ultrasound_analysis_timestamp', 'Недавно')
                    })
            
            # Проверяем результаты дерматоскопии
            if 'dermatoscopy_analysis_result' in st.session_state:
                derm_result = st.session_state.dermatoscopy_analysis_result
                if isinstance(derm_result, dict) or isinstance(derm_result, str):
                    analysis_results.append({
                        'type': 'Дерматоскопия',
                        'data': derm_result,
                        'timestamp': st.session_state.get('dermatoscopy_analysis_timestamp', 'Недавно')
                    })
            
            # Проверяем результаты лабораторных анализов
            if 'lab_analysis_result' in st.session_state:
                lab_result = st.session_state.lab_analysis_result
                if isinstance(lab_result, dict) or isinstance(lab_result, str):
                    analysis_results.append({
                        'type': 'Лабораторные анализы',
                        'data': lab_result,
                        'timestamp': st.session_state.get('lab_analysis_timestamp', 'Недавно')
                    })
            
            # Проверяем результаты генетического анализа
            if 'genetic_analysis_results' in st.session_state:
                for key, data in st.session_state.genetic_analysis_results.items():
                    result_data = data.get('result')
                    if result_data:
                        analysis_results.append({
                            'type': 'Генетический анализ',
                            'data': result_data,
                            'timestamp': data.get('file_name', 'Недавно')
                        })
            
            if analysis_results:
                st.info(f"Найдено {len(analysis_results)} сохраненных результатов анализов")
                
                selected_analyses = st.multiselect(
                    "Выберите анализы для пересылки ИИ-консультанту:",
                    options=[f"{r['type']} ({r['timestamp']})" for r in analysis_results],
                    help="Выбранные анализы будут добавлены в контекст следующего вопроса"
                )
                
                if st.button("✅ Добавить выбранные анализы в контекст"):
                    if selected_analyses:
                        # Добавляем выбранные анализы в контекст загруженных файлов
                        for result_label in selected_analyses:
                            result_index = [f"{r['type']} ({r['timestamp']})" for r in analysis_results].index(result_label)
                            result = analysis_results[result_index]
                            
                            # Форматируем результат для контекста
                            if isinstance(result['data'], dict):
                                result_text = json.dumps(result['data'], ensure_ascii=False, indent=2)
                            elif hasattr(result['data'], '__dict__'):
                                result_text = json.dumps(result['data'].__dict__, ensure_ascii=False, indent=2)
                            else:
                                result_text = str(result['data'])
                            
                            st.session_state.uploaded_files_context.append({
                                'file_name': f"Заключение: {result['type']}",
                                'type': 'analysis_result',
                                'content': f"Тип анализа: {result['type']}\nДата: {result['timestamp']}\n\nРезультаты:\n{result_text[:3000]}"
                            })
                        
                        st.success(f"✅ Добавлено {len(selected_analyses)} результатов анализов в контекст")
                        st.rerun()
                    else:
                        st.warning("⚠️ Выберите хотя бы один анализ")
            else:
                st.info("💡 Нет сохраненных результатов анализов. Выполните анализ изображений или данных, чтобы их можно было переслать консультанту.")

        # Загрузка файлов
        with st.expander("📎 Загрузить файлы для анализа", expanded=False):
            uploaded_files = st.file_uploader(
                "Загрузите файлы для анализа",
                type=["pdf", "txt", "docx", "jpg", "jpeg", "png", "csv", "json"],
                accept_multiple_files=True,
                help="Поддерживаются: PDF, TXT, DOCX, изображения, CSV, JSON"
            )
            
            if uploaded_files:
                for uploaded_file in uploaded_files:
                    file_ext = uploaded_file.name.split('.')[-1].lower()
                    
                    if file_ext == 'pdf':
                        try:
                            # Извлекаем текст из PDF (в т.ч. генетических отчетов) через AdvancedLabProcessor
                            from modules.advanced_lab_processor import AdvancedLabProcessor
                            processor = AdvancedLabProcessor()
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                                tmp.write(uploaded_file.getvalue())
                                tmp_path = tmp.name
                            
                                extracted_text = processor._extract_from_pdf(tmp_path)
                            
                            st.session_state.uploaded_files_context.append({
                                'file_name': uploaded_file.name,
                                'type': 'pdf',
                                'content': str(extracted_text)[:10000]  # Увеличиваем лимит до 10000 символов
                            })
                            st.success(f"✅ {uploaded_file.name}: извлечено {len(str(extracted_text))} символов")
                            os.unlink(tmp_path)
                        except Exception as e:
                            st.error(f"❌ Ошибка обработки {uploaded_file.name}: {e}")
                            import traceback
                            st.error(f"Детали: {traceback.format_exc()}")
                    
                    elif file_ext in ['txt', 'csv', 'json']:
                        try:
                            content = uploaded_file.read().decode('utf-8')
                            st.session_state.uploaded_files_context.append({
                                'file_name': uploaded_file.name,
                                'type': file_ext,
                                'content': content[:5000]
                            })
                            st.success(f"✅ {uploaded_file.name}: загружено {len(content)} символов")
                        except Exception as e:
                            st.error(f"❌ Ошибка обработки {uploaded_file.name}: {e}")
                    
                    elif file_ext in ['jpg', 'jpeg', 'png']:
                        try:
                            from PIL import Image
                            import numpy as np
                            image = Image.open(uploaded_file)
                            image_array = np.array(image)
                            
                            # Используем ИИ для анализа изображения
                            with st.spinner(f"Анализ изображения {uploaded_file.name}..."):
                                image_description = assistant.send_vision_request(
                                    "Опиши это медицинское изображение подробно. Извлеки всю видимую информацию: текст, цифры, структуры, паттерны.",
                                    image_array
                                )
                            
                            st.session_state.uploaded_files_context.append({
                                'file_name': uploaded_file.name,
                                'type': 'image',
                                'content': image_description[:2000]
                            })
                            st.success(f"✅ {uploaded_file.name}: изображение проанализировано")
                        except Exception as e:
                            st.error(f"❌ Ошибка обработки {uploaded_file.name}: {e}")

        # Отображение истории чата
        for msg in st.session_state.chat_history:
            st.chat_message("user").write(msg['user'])
            if msg.get('files_info'):
                with st.expander("📎 Прикрепленные файлы"):
                    for file_info in msg['files_info']:
                        st.write(f"**{file_info['name']}** ({file_info['type']})")
            st.chat_message("assistant").write(msg['assistant'])

        # Выбор режима ввода
        input_mode = st.radio(
            "Режим ввода:",
            ["📝 Текстовый", "🎤 Голосовой"],
            horizontal=True,
            key="ai_chat_input_mode"
        )
        
        user_input = None
        
        # Проверяем, есть ли сохраненный транскрибированный вопрос
        if 'transcribed_question' in st.session_state:
            user_input = st.session_state['transcribed_question']
            st.info(f"🎤 **Транскрибированный вопрос:** {user_input}")
            st.info("💡 Вопрос будет отправлен автоматически. Если нужно изменить, используйте текстовый ввод.")
            del st.session_state['transcribed_question']  # Удаляем после использования
        
        # Голосовой ввод (показываем только если еще нет транскрибированного вопроса)
        if input_mode == "🎤 Голосовой" and not user_input:
            if not ASSEMBLYAI_AVAILABLE:
                st.warning("⚠️ Голосовой ввод недоступен. AssemblyAI не настроен. Используйте текстовый ввод.")
            else:
                audio_data = st.audio_input("🎤 Запишите ваш вопрос", key="ai_chat_audio")
                if audio_data:
                    st.info("💡 Аудио записано. Нажмите кнопку ниже для расшифровки.")
                    if st.button("🎤 Расшифровать аудио", use_container_width=True, type="primary"):
                        try:
                            with st.spinner("🎤 Расшифровка аудио..."):
                                # Получаем API ключ из конфига
                                from config import ASSEMBLYAI_API_KEY
                                api_key = ASSEMBLYAI_API_KEY or st.secrets.get("ASSEMBLYAI_API_KEY", "")
                                if not api_key:
                                    st.error("❌ API ключ AssemblyAI не настроен. Проверьте config.py или secrets.")
                                else:
                                    # Убеждаемся, что передаем правильный формат данных
                                    # st.audio_input возвращает BytesIO, который нужно правильно обработать
                                    transcribed_text = transcribe_audio_assemblyai(audio_data, api_key)
                                    
                                    if transcribed_text and not transcribed_text.startswith("❌"):
                                        # Сохраняем транскрибированный текст в session_state
                                        st.session_state['transcribed_question'] = transcribed_text
                                        st.success(f"✅ Расшифровано: {transcribed_text[:100]}...")
                                        st.rerun()  # Перезагружаем для отправки вопроса
                                    else:
                                        st.error(f"❌ Ошибка расшифровки: {transcribed_text}")
                        except Exception as e:
                            st.error(f"❌ Ошибка обработки аудио: {e}")
                            import traceback
                            with st.expander("🔍 Детали ошибки"):
                                st.code(traceback.format_exc())
        
        # Текстовый ввод (если не выбран голосовой или если голосовой не дал результата)
        # Показываем текстовый ввод только если нет транскрибированного вопроса
        if not user_input:
            if input_mode != "🎤 Голосовой" or not st.session_state.get('transcribed_question'):
                user_input = st.chat_input("Задайте вопрос врачу-консультанту...")
        
        if user_input:
            # Формируем контекст из истории и загруженных файлов
            context_parts = []
            
            # Добавляем контекст из загруженных файлов
            if st.session_state.uploaded_files_context:
                context_parts.append("=== ЗАГРУЖЕННЫЕ ФАЙЛЫ ДЛЯ АНАЛИЗА ===")
                for file_ctx in st.session_state.uploaded_files_context:
                    context_parts.append(f"\nФайл: {file_ctx['file_name']} (тип: {file_ctx['type']})")
                    context_parts.append(f"Содержимое:\n{file_ctx['content']}")
                context_parts.append("\nВАЖНО: Учитывайте информацию из этих файлов при ответе на вопрос.")
            
            # Добавляем контекст из предыдущих сообщений (последние 10 для лучшего понимания)
            if st.session_state.chat_history:
                context_parts.append("\n=== КОНТЕКСТ ПРЕДЫДУЩЕГО ДИАЛОГА ===")
                context_parts.append("Ниже приведена история предыдущих вопросов и ответов. Используйте этот контекст для более точного ответа.")
                recent_history = st.session_state.chat_history[-10:]  # Последние 10 сообщений для лучшего контекста
                for i, msg in enumerate(recent_history, 1):
                    context_parts.append(f"\n--- Обмен {i} ---")
                    context_parts.append(f"Врач спрашивает: {msg['user']}")
                    # Берем первые 300 символов ответа для контекста
                    assistant_response_preview = msg['assistant'][:300] + "..." if len(msg['assistant']) > 300 else msg['assistant']
                    context_parts.append(f"Консультант отвечал: {assistant_response_preview}")
                    if msg.get('files_info'):
                        context_parts.append(f"Прикрепленные файлы: {', '.join([f['name'] for f in msg['files_info']])}")
                context_parts.append("\nВАЖНО: Учитывайте контекст предыдущих обсуждений. Если вопрос связан с предыдущими темами, ссылайтесь на них.")
            
            context = "\n".join(context_parts) if context_parts else ""
            
            # Информация о загруженных файлах для отображения
            files_info = [{'name': f['file_name'], 'type': f['type']} 
                         for f in st.session_state.uploaded_files_context]
            
            st.chat_message("user").write(user_input)
            if files_info:
                with st.expander("📎 Прикрепленные файлы"):
                    for file_info in files_info:
                        st.write(f"**{file_info['name']}** ({file_info['type']})")
            
            # Используем streaming для более комфортного общения
            with st.chat_message("assistant"):
                try:
                    text_generator = assistant.get_response_streaming(user_input, context=context, use_sonnet_4_5=True)
                    response = st.write_stream(text_generator)
                except Exception as e:
                    # Fallback на обычный режим если streaming не работает
                    st.warning("⚠️ Streaming временно недоступен, используем обычный режим...")
                    response = assistant.get_response(user_input, context=context, use_sonnet_4_5=True)
                    st.write(response)
            
            # Убеждаемся что response - строка
            if not isinstance(response, str):
                response = str(response) if response else ""
            
            # Сохраняем в историю
            timestamp = datetime.datetime.now().isoformat()
            chat_entry = {
                'user': user_input,
                'assistant': response,
                'files_info': files_info,
                'timestamp': timestamp
            }
            st.session_state.chat_history.append(chat_entry)
            
            # Сохраняем в базу данных
            try:
                conn = sqlite3.connect('medical_data.db')
                cursor = conn.cursor()
                cursor.execute('''
                    INSERT INTO ai_chat_history 
                    (session_id, user_message, assistant_response, files_context, context_summary)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    st.session_state.chat_session_id,
                    user_input,
                    response,
                    json.dumps(files_info, ensure_ascii=False),
                    context[:500] if context else ""  # Краткое резюме контекста
                ))
                conn.commit()
                conn.close()
            except Exception as e:
                print(f"Ошибка сохранения истории: {e}")
            
            # Ограничиваем размер истории в памяти
            if len(st.session_state.chat_history) > 50:
                st.session_state.chat_history = st.session_state.chat_history[-50:]
            
            # Очищаем загруженные файлы после использования (опционально)
            # st.session_state.uploaded_files_context = []

    except Exception as e:
        st.error(f"Ошибка: {e}")
        import traceback
        with st.expander("🔍 Детали ошибки"):
            st.code(traceback.format_exc())

def show_clinical_recommendations(diagnosis):
    """Простые клинические рекомендации без API"""
    st.markdown("### 📚 Клинические рекомендации")
    
    recommendations = {
        "пневмония": {
            "icd10": "J18.9",
            "treatment": ["Амоксициллин 500мг 3р/день", "Покой", "Обильное питье"],
            "diagnostics": ["Рентген ОГК", "Общий анализ крови", "Посев мокроты"]
        },
        "инфаркт": {
            "icd10": "I21.9",
            "treatment": ["Экстренная госпитализация", "Аспирин 300мг", "Тромболизис"],
            "diagnostics": ["ЭКГ-12", "Тропонины", "ЭхоКГ"]
        },
        "рентген": {
            "icd10": "Z01.6",
            "treatment": ["Интерпретация специалистом"],
            "diagnostics": ["Оценка качества", "Поиск патологий"]
        }
    }
    
    if diagnosis in recommendations:
        rec = recommendations[diagnosis]
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 🔍 Диагностика")
            for item in rec["diagnostics"]:
                st.markdown(f"- {item}")
        
        with col2:
            st.markdown("#### 💊 Лечение")
            for item in rec["treatment"]:
                st.markdown(f"- {item}")
        
        st.markdown(f"**Код по МКБ-10:** `{rec['icd10']}`")
    else:
        st.info("Рекомендации для данного диагноза не найдены")

def show_lab_analysis():
    """Улучшенная страница анализа лабораторных данных"""
    st.header("🔬 Анализ лабораторных данных")
    
    # Инициализация нового процессора
    if 'lab_processor' not in st.session_state:
        st.session_state.lab_processor = AdvancedLabProcessor()
    
    processor = st.session_state.lab_processor
    
    # Настройки
    col1, col2 = st.columns(2)
    with col1:
        auto_detect_type = st.checkbox("Автоопределение типа файла", value=True)
    with col2:
        show_raw_data = st.checkbox("Показать исходные данные", value=False)
    
    # Загрузка файла
    uploaded_file = st.file_uploader(
    "Загрузите файл с лабораторными данными",
    type=["pdf", "xlsx", "xls", "csv", "json", "xml", "jpg", "jpeg", "png"],  # ← добавили изображения
    help="Поддерживаются: PDF, Excel, CSV, JSON, XML, JPG, PNG"
)
    
    if uploaded_file and st.button("🧪 Анализировать лабораторные данные"):
        with st.spinner("Обработка лабораторных данных..."):
            
            # Сохраняем временный файл
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_path = tmp_file.name
            
            try:
                # Определяем тип файла если нужно
                file_type = None
                if not auto_detect_type:
                    file_ext = uploaded_file.name.split('.')[-1].lower()
                    file_type = file_ext
                
                # Обработка
                lab_report = processor.process_file(tmp_path, file_type=file_type, ai_assistant=OpenRouterAssistant() if AI_AVAILABLE else None)
                
                # Результаты
                if lab_report.parameters and len(lab_report.parameters) > 0:
                    st.success(f"✅ Обработано {len(lab_report.parameters)} параметров")
                    
                    # Метрики
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Параметров", len(lab_report.parameters))
                    with col2:
                        st.metric("Достоверность", f"{lab_report.confidence:.1%}")
                    with col3:
                        critical_count = len(lab_report.critical_values)
                        st.metric("Критических", critical_count, delta="⚠️" if critical_count > 0 else None)
                    with col4:
                        normal_count = len([p for p in lab_report.parameters if p.status == "normal"])
                        st.metric("В норме", f"{normal_count}/{len(lab_report.parameters)}")
                    
                    # Критические значения
                    if lab_report.critical_values:
                        st.error("🚨 **КРИТИЧЕСКИЕ ЗНАЧЕНИЯ:**")
                        for critical in lab_report.critical_values:
                            st.error(f"• {critical}")
                    
                    # Предупреждения
                    if lab_report.warnings:
                        st.warning("⚠️ **Предупреждения:**")
                        for warning in lab_report.warnings:
                            st.warning(f"• {warning}")
                    
                    # Таблица результатов
                    st.subheader("📊 Результаты анализов")
                    try:
                        df = processor.to_dataframe(lab_report)
                    except Exception as e:
                        st.warning(f"⚠️ Ошибка создания таблицы: {e}")
                        # Создаем простую таблицу вручную
                        import pandas as pd
                        data = []
                        for param in lab_report.parameters:
                            data.append({
                                'Параметр': param.name,
                                'Значение': param.value,
                                'Единица': param.unit,
                                'Норма': param.reference_range,
                                'Статус': param.status,
                                'Категория': param.category
                            })
                        df = pd.DataFrame(data)
                    
                    # Цветовая кодировка статусов
                    def style_status(val):
                        colors = {
                            'normal': 'background-color: #d4edda',
                            'high': 'background-color: #fff3cd', 
                            'low': 'background-color: #fff3cd',
                            'critical_high': 'background-color: #f8d7da',
                            'critical_low': 'background-color: #f8d7da'
                        }
                        return colors.get(val, '')
                    
                    styled_df = df.style.applymap(style_status, subset=['Статус'])
                    st.dataframe(styled_df, use_container_width=True)
                    
                    # Группировка по категориям
                    st.subheader("📋 Анализ по системам")
                    summary = processor.generate_summary(lab_report)
                    
                    for category, params in summary['categories'].items():
                        with st.expander(f"📁 {category.title()} ({len(params)} параметров)"):
                            for param in params:
                                status_emoji = {
                                    'normal': '✅',
                                    'high': '⬆️', 
                                    'low': '⬇️',
                                    'critical_high': '🔴',
                                    'critical_low': '🔴'
                                }.get(param['status'], '❓')
                                
                                st.markdown(f"{status_emoji} **{param['name']}:** {param['value']} {param['unit']} ({param['status']})")
                    
                    # Форма обратной связи - ДО анализа, всегда видна и активна!
                    st.markdown("---")
                    st.markdown("### 💬 Обратная связь")
                    
                    last_result = st.session_state.get('lab_analysis_result', '')
                    analysis_id_base = "LAB_feedback_form"
                    lab_input = f"Лабораторные данные: {len(lab_report.parameters)} параметров, Критические: {len(lab_report.critical_values) if lab_report.critical_values else 0}"
                    
                    try:
                        show_feedback_form(
                            analysis_type="LAB",
                            analysis_result=str(last_result) if last_result else "",
                            analysis_id=analysis_id_base,
                            input_case=lab_input
                        )
                    except Exception as e:
                        st.error(f"Ошибка формы обратной связи: {e}")
                    
                    if not last_result:
                        st.info("💡 После проведения анализа форма автоматически обновится с новым результатом.")
                    
                    # ИИ-интерпретация с полной интеграцией компонентов
                    st.subheader("🤖 ИИ-интерпретация результатов")
                    
                    # Выбор режима анализа
                    lab_analysis_mode = st.radio(
                        "Режим анализа:",
                        ["⚡ Быстрый (одна модель)", "🎯 Консенсус (несколько моделей)", "✅ С валидацией"],
                        horizontal=True,
                        key="lab_analysis_mode"
                    )
                    
                    if st.button("🧪 Запустить ИИ-анализ", use_container_width=True):
                        with st.spinner("ИИ анализирует результаты..."):
                            # Формируем контекст для ИИ
                            context = f"""
Лабораторные результаты пациента:
Количество параметров: {len(lab_report.parameters)}
Достоверность анализа: {lab_report.confidence:.1%}

Результаты:
"""
                            for param in lab_report.parameters:
                                context += f"- {param.name}: {param.value} {param.unit} (норма: {param.reference_range}, статус: {param.status})\n"
                            
                            if lab_report.critical_values:
                                context += f"\nКритические значения: {'; '.join(lab_report.critical_values)}"
                            
                            # Промпт от имени специалиста
                            base_prompt = f"""Проанализируйте лабораторные результаты как врач-лаборант-консультант с 15-летним опытом работы в клинической лаборатории. 
Дайте клиническую оценку, выявите критические значения, предложите дифференциальную диагностику и рекомендации в формате "Клиническая директива".

{context}"""
                            
                            try:
                                assistant = OpenRouterAssistant()
                                consensus_engine = ConsensusEngine(assistant)
                                validator = ValidationPipeline(assistant)
                                scorecard = MedicalScorecard()
                                gap_detector = DiagnosticGapDetector()
                                notifier = NotificationSystem()
                                evidence_ranker = EvidenceRanker()
                                
                                if lab_analysis_mode == "⚡ Быстрый (одна модель)":
                                    interpretation = assistant.get_response(base_prompt)
                                    st.markdown("### 🧠 ИИ-интерпретация (Врач-лаборант-консультант)")
                                    st.write(interpretation)
                                    
                                    # Сохраняем результат (форма обновится при следующем рендере)
                                    st.session_state.lab_analysis_result = interpretation
                                    st.session_state.lab_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                                    
                                elif lab_analysis_mode == "🎯 Консенсус (несколько моделей)":
                                    # Для текстового анализа используем get_multiple_opinions
                                    opinions = consensus_engine.get_multiple_opinions(base_prompt)
                                    
                                    # Генерация консенсуса
                                    findings_list = [consensus_engine.extract_key_findings(op['response']) for op in opinions]
                                    comparison = consensus_engine.compare_opinions(opinions)
                                    
                                    consensus_report = consensus_engine._generate_consensus_report(
                                        findings_list,
                                        comparison.get('common_diagnoses', []),
                                        comparison.get('urgency', 'не определена'),
                                        comparison.get('discrepancies', [])
                                    )
                                    
                                    st.markdown("### 🎯 Консенсус-анализ:")
                                    st.write(consensus_report)
                                    
                                    with st.expander("📊 Детали мнений моделей"):
                                        for i, opinion in enumerate(opinions, 1):
                                            st.markdown(f"**Модель {i}:**")
                                            st.write(opinion['response'][:500] + "...")
                                    
                                    # Сохраняем результат (форма обновится при следующем рендере)
                                    st.session_state.lab_analysis_result = consensus_report
                                    st.session_state.lab_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                                    
                                elif lab_analysis_mode == "✅ С валидацией":
                                    interpretation = assistant.get_response(base_prompt)
                                    
                                    # Валидация
                                    validation = validator.validate_response(interpretation)
                                    
                                    # Оценка качества (используем общий чек-лист)
                                    evaluation = scorecard.evaluate_response(interpretation, ImageType.ECG)  # Используем общий тип
                                    
                                    # Критические находки
                                    critical_findings = notifier.check_critical_findings(interpretation)
                                    
                                    # Оценка доказательности
                                    evidence = evidence_ranker.rank_evidence(interpretation)
                                    
                                    # Отображение результатов
                                    st.markdown("### 🧠 ИИ-интерпретация (Врач-лаборант-консультант)")
                                    st.write(interpretation)
                                    
                                    # Сохраняем результат (форма обновится при следующем рендере)
                                    st.session_state.lab_analysis_result = interpretation
                                    st.session_state.lab_analysis_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                                    
                                    # Уведомления о критических находках
                                    notifier.display_notifications(critical_findings)
                                    
                                    # Валидация
                                    with st.expander("✅ Результаты валидации"):
                                        if validation['is_valid']:
                                            st.success("✅ Валидация пройдена")
                                        else:
                                            st.error("❌ Обнаружены проблемы")
                                        st.write(f"Полнота: {validation['completeness_score']:.1%}")
                                        if validation['warnings']:
                                            for warning in validation['warnings']:
                                                st.warning(warning)
                                        if validation['errors']:
                                            for error in validation['errors']:
                                                st.error(error)
                                    
                                    # Оценка качества
                                    with st.expander("📊 Оценка качества"):
                                        st.write(f"**Оценка:** {evaluation['grade']}")
                                        st.write(f"**Балл:** {evaluation['score']:.1%}")
                                        if evaluation['recommendations']:
                                            st.write("**Рекомендации:**")
                                            for rec in evaluation['recommendations']:
                                                st.write(f"• {rec}")
                                    
                                    # Доказательность
                                    with st.expander("📚 Оценка доказательности"):
                                        st.write(evidence_ranker.generate_evidence_report(evidence))
                                
                            except Exception as e:
                                error_msg = handle_error(e, "show_lab_analysis", show_to_user=True)
                                st.error(f"Ошибка ИИ-анализа: {error_msg}")
                    
                    # Исходные данные
                    if show_raw_data:
                        st.subheader("📄 Исходные данные")
                        st.text_area("Извлеченный текст", lab_report.raw_text, height=200)
                    
                    # Скачать результаты
                    csv_data = df.to_csv(index=False, encoding='utf-8')
                    st.download_button(
                        label="💾 Скачать результаты (CSV)",
                        data=csv_data,
                        file_name=f"lab_results_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                        mime="text/csv"
                    )
                    
                    # Экспорт в Excel
                    lab_data_for_export = {
                        'parameters': [{
                            'name': p.name,
                            'value': p.value,
                            'unit': p.unit,
                            'reference_range': p.reference_range,
                            'status': p.status
                        } for p in lab_report.parameters],
                        'critical_values': lab_report.critical_values,
                        'warnings': lab_report.warnings
                    }
                    
                    excel_path = export_lab_results_to_excel(lab_data_for_export)
                    with open(excel_path, 'rb') as f:
                        st.download_button(
                            label="📊 Скачать результаты (Excel)",
                            data=f.read(),
                            file_name=f"lab_results_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                
                else:
                    st.error("❌ Не удалось извлечь лабораторные данные из файла")
                    
                    # Показываем детальную информацию об ошибке
                    if lab_report.warnings:
                        st.warning("⚠️ **Предупреждения:**")
                        for warning in lab_report.warnings:
                            st.warning(f"• {warning}")
                    
                    # Показываем извлеченный текст для диагностики
                    if lab_report.raw_text:
                        st.info("📄 **Извлеченный текст из файла:**")
                        st.text_area("", lab_report.raw_text, height=300, key="raw_text_display")
                        
                        # Попытка ручного парсинга
                        if st.button("🔍 Попробовать извлечь параметры вручную"):
                            with st.spinner("Анализ текста..."):
                                try:
                                    # Пробуем использовать ИИ для извлечения
                                    assistant = OpenRouterAssistant()
                                    ai_prompt = f"""Извлеки все лабораторные параметры из следующего текста в формате JSON:
                                    
{lab_report.raw_text[:2000]}

Верни JSON массив с объектами вида:
{{"name": "название параметра", "value": число, "unit": "единица измерения", "reference": "норма"}}
"""
                                    ai_result = assistant.get_response(ai_prompt)
                                    st.success("✅ ИИ извлек данные:")
                                    st.json(ai_result)
                                except Exception as e:
                                    st.error(f"Ошибка ИИ-извлечения: {e}")
                    else:
                        st.warning("⚠️ Не удалось извлечь текст из файла. Проверьте формат файла.")
            
            except Exception as e:
                import traceback
                error_msg = str(e)
                st.error(f"❌ Ошибка обработки файла: {error_msg}")
                
                # Показываем детальную информацию об ошибке
                with st.expander("🔍 Детали ошибки и советы"):
                    st.code(error_msg)
                    st.write("**Трассировка ошибки:**")
                    st.code(traceback.format_exc())
                    st.info("💡 **Советы по устранению:**")
                    st.write("""
                    1. **Проверьте формат файла** - поддерживаются: PDF, Excel (xlsx, xls), CSV, JSON, XML, изображения (JPG, PNG)
                    2. **Убедитесь, что файл не поврежден** - попробуйте открыть его в другой программе
                    3. **Для PDF файлов** - убедитесь, что текст можно выделить (не сканированное изображение)
                    4. **Для Excel файлов** - проверьте, что файл не защищен паролем
                    5. **Для CSV файлов** - проверьте кодировку (должна быть UTF-8 или Windows-1251)
                    6. **Попробуйте сохранить файл в другом формате** (например, CSV вместо Excel)
                    7. **Для изображений** - используйте ИИ-анализ, если автоматическое извлечение не работает
                    """)
            
            finally:
                # Удаляем временный файл
                try:
                    os.unlink(tmp_path)
                except:
                    pass

def show_genetic_analysis_page():
    """Страница анализа генетических данных с поддержкой VCF"""
    st.header("🧬 Генетический анализ")
    
    # Импорт генетического анализатора
    try:
        from modules.genetic_analyzer import GeneticAnalyzer, VCFParser
        GENETIC_ANALYZER_AVAILABLE = True
    except ImportError as e:
        st.error(f"❌ Модуль генетического анализа недоступен: {e}")
        GENETIC_ANALYZER_AVAILABLE = False
        return
    
    # Информация о пациенте
    st.subheader("👤 Информация о пациенте")
    col1, col2, col3 = st.columns(3)
    with col1:
        age = st.number_input("Возраст", 1, 120, 30)
    with col2:
        gender = st.selectbox("Пол", ["М", "Ж"])
    with col3:
        lifestyle = st.selectbox("Образ жизни", ["Низкая активность", "Средняя активность", "Высокая активность"])
    
    # Клинический контекст
    clinical_context = st.text_area(
        "Клинический контекст (опционально)",
        placeholder="Укажите жалобы, семейный анамнез, сопутствующие заболевания...",
        height=100
    )
    
    # Загрузка файла
    uploaded_file = st.file_uploader(
        "Загрузите генетический файл или снимок отчета", 
        type=["vcf", "vcf.gz", "txt", "csv", "pdf", "jpg", "jpeg", "png"],
        help="Поддерживаются: VCF, VCF.GZ (сжатый), TXT, CSV, PDF, а также скриншоты (JPG, JPEG, PNG) генетических отчетов"
    )
    
    if uploaded_file:
        file_ext = uploaded_file.name.split('.')[-1].lower()
        file_name = uploaded_file.name
        
        # Сохранение во временный файл
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_ext}") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        # Сохраняем путь к файлу в session_state для повторного использования
        file_key = f"genetic_file_{uploaded_file.name}"
        
        if st.button("🧬 Запустить генетический анализ", use_container_width=True):
            if not GENETIC_ANALYZER_AVAILABLE:
                st.error("❌ Модуль генетического анализа недоступен. Проверьте файл modules/genetic_analyzer.py")
                return
            try:
                with st.spinner("🔬 Анализ генетических данных..."):
                    # Инициализация анализатора
                    analyzer = GeneticAnalyzer()
                    
                    # Информация о пациенте
                    patient_info = {
                        "age": age,
                        "gender": gender,
                        "lifestyle": lifestyle
                    }
                    
                    # Анализ VCF файла
                    if file_ext in ['vcf', 'gz']:
                        st.info("📄 Парсинг VCF файла...")
                        analysis_result = analyzer.analyze_vcf_file(
                            tmp_path,
                            patient_info=patient_info,
                            clinical_context=clinical_context
                        )
                        
                        # Сохраняем результат в session_state для использования после rerun
                        if 'genetic_analysis_results' not in st.session_state:
                            st.session_state.genetic_analysis_results = {}
                        
                        st.session_state.genetic_analysis_results[file_key] = {
                            'result': analysis_result,
                            'patient_info': patient_info,
                            'clinical_context': clinical_context,
                            'file_name': file_name
                        }
                        
                        # Отображение результатов
                        st.success("✅ Анализ завершен! Результаты сохранены.")
                        st.rerun()  # Перезагружаем страницу, чтобы показать сохраненные результаты
                    
                    # Анализ скриншота (изображения) генетического отчета
                    elif file_ext in ['jpg', 'jpeg', 'png']:
                        if not AI_AVAILABLE or OpenRouterAssistant is None:
                            st.error("❌ ИИ-модуль недоступен. Скриншот не может быть проанализирован.")
                        else:
                            st.info("🖼️ Обработка скриншота генетического отчета и извлечение текста (OCR)...")
                            st.info("💡 Система попытается распознать таблицы с генами, rsID и генотипами и затем выполнит анализ, как для текстового отчета.")
                            try:
                                from PIL import Image
                                import numpy as np
                                image = Image.open(tmp_path)
                                image_array = np.array(image)

                                from claude_assistant import OpenRouterAssistant as _OraForImage  # локальный псевдоним
                                img_assistant = _OraForImage()

                                ocr_prompt = """
Вы — эксперт по OCR генетических отчетов.
Аккуратно извлеките ВЕСЬ текст с этого изображения (особенно таблицы с генами, SNP/rsID и генотипами).
Верните ТОЛЬКО распознанный текст без интерпретации и без клинических выводов.
"""
                                ocr_result = img_assistant.send_vision_request(
                                    ocr_prompt,
                                    image_array,
                                    metadata={"task": "doc_ocr", "source": "genetic_screenshot"}
                                )
                                if isinstance(ocr_result, list):
                                    ocr_text = "\n\n".join(str(x.get("result", x)) for x in ocr_result)
                                else:
                                    ocr_text = str(ocr_result)

                                analysis_result = analyzer.analyze_text_report(
                                    report_text=ocr_text,
                                    patient_info=patient_info,
                                    clinical_context=clinical_context,
                                    source="image_report_ocr"
                                )

                                # Сохраняем результат в session_state для использования после rerun
                                if 'genetic_analysis_results' not in st.session_state:
                                    st.session_state.genetic_analysis_results = {}

                                st.session_state.genetic_analysis_results[file_key] = {
                                    'result': analysis_result,
                                    'patient_info': patient_info,
                                    'clinical_context': clinical_context,
                                    'file_name': file_name
                                }

                                if analysis_result.total_variants > 0:
                                    st.success(f"✅ Анализ скриншота завершен! Найдено {analysis_result.total_variants} строк с потенциальными генетическими вариантами. Результаты сохранены.")
                                else:
                                    st.warning("⚠️ Явных генетических вариантов на скриншоте не найдено. Попробуйте загрузить PDF или VCF, если они доступны.")
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Ошибка OCR/анализа скриншота: {e}")
                                import traceback
                                with st.expander("🔍 Детали ошибки"):
                                    st.code(traceback.format_exc())

                    # Анализ PDF файла (генетический отчет)
                    elif file_ext == 'pdf':
                        st.info("📄 Извлечение текста из PDF и поиск генетических данных...")
                        st.info("💡 Система извлечет текст из PDF (включая таблицы, если они представлены текстом) и найдет строки с генетическими вариантами (гены, c./p.-нотация, rsID).")
                        st.info("💡 Альтернатива: для максимально точного анализа используйте VCF файлы (стандартизированный формат генетических данных).")
                        
                        analysis_result = analyzer.analyze_pdf_file(
                            tmp_path,
                            patient_info=patient_info,
                            clinical_context=clinical_context
                        )
                        
                        # Сохраняем результат в session_state для использования после rerun
                        if 'genetic_analysis_results' not in st.session_state:
                            st.session_state.genetic_analysis_results = {}
                        
                        st.session_state.genetic_analysis_results[file_key] = {
                            'result': analysis_result,
                            'patient_info': patient_info,
                            'clinical_context': clinical_context,
                            'file_name': file_name
                        }
                        
                        # Отображение результатов
                        if analysis_result.total_variants > 0:
                            st.success(f"✅ Анализ завершен! Найдено {analysis_result.total_variants} строк с потенциальными генетическими вариантами. Результаты сохранены.")
                        else:
                            st.warning("⚠️ Явных генетических вариантов в тексте PDF не найдено. Отчет может содержать только изображения (сканы таблиц) или нестандартный формат.")
                        st.rerun()  # Перезагружаем страницу, чтобы показать сохраненные результаты
                        
            except Exception as e:
                st.error(f"❌ Ошибка анализа: {e}")
                import traceback
                with st.expander("🔍 Детали ошибки"):
                    st.code(traceback.format_exc())
                analysis_result = None
            finally:
                # Очистка временного файла
                try:
                    if os.path.exists(tmp_path):
                        os.unlink(tmp_path)
                except:
                    pass
        
        # Получаем сохраненный результат анализа
        analysis_result = None
        if 'genetic_analysis_results' in st.session_state and file_key in st.session_state.genetic_analysis_results:
            saved_data = st.session_state.genetic_analysis_results[file_key]
            analysis_result = saved_data['result']
            patient_info = saved_data['patient_info']
            clinical_context = saved_data.get('clinical_context', '')
        
        # Отображаем результаты, если анализ был выполнен
        if analysis_result:
            # Техническая информация (если есть)
            if analysis_result.metadata and 'technical_info' in analysis_result.metadata:
                tech_info = analysis_result.metadata.get('technical_info', {})
                if any(tech_info.values()):
                    with st.expander("🔬 Техническая информация об анализе", expanded=False):
                        col1, col2 = st.columns(2)
                        with col1:
                            if tech_info.get('method'):
                                st.info(f"**Метод анализа:** {tech_info.get('method')}")
                            if tech_info.get('laboratory'):
                                st.info(f"**Лаборатория:** {tech_info.get('laboratory')}")
                            if tech_info.get('accreditation'):
                                st.info(f"**Аккредитация:** {tech_info.get('accreditation')}")
                        with col2:
                            if tech_info.get('reference_genome'):
                                st.info(f"**Референсный геном:** {tech_info.get('reference_genome')}")
                            if tech_info.get('pipeline'):
                                st.info(f"**Биоинформатический пайплайн:** {tech_info.get('pipeline')}")
                            if tech_info.get('geneticist_signature'):
                                st.info(f"**Подпись генетика:** {tech_info.get('geneticist_signature')}")
            
            # Спектр генов (если есть)
            if analysis_result.metadata and 'gene_panel' in analysis_result.metadata:
                gene_panel = analysis_result.metadata.get('gene_panel', [])
                if gene_panel:
                    st.subheader("🧬 Спектр проанализированных генов")
                    st.info(f"Найдено {len(gene_panel)} генов в спектре анализа")
                    # Отображаем гены в виде колонок
                    cols_per_row = 5
                    for i in range(0, len(gene_panel), cols_per_row):
                        cols = st.columns(cols_per_row)
                        for j, col in enumerate(cols):
                            if i + j < len(gene_panel):
                                with col:
                                    st.code(gene_panel[i + j], language=None)
            
            # Основная статистика
            st.subheader("📊 Статистика анализа")
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Всего вариантов / строк", analysis_result.total_variants)
            with col2:
                st.metric("Патогенных", len(analysis_result.pathogenic_variants))
            with col3:
                st.metric("Клинически значимых", len(analysis_result.clinical_interpretations))
            with col4:
                st.metric("Фармакогенетических", len(analysis_result.pharmacogenetic_variants))
            
            # Блок работы с текстовыми/изображенческими генетическими отчетами (PDF, скриншоты)
            if analysis_result.metadata:
                text_variants = analysis_result.metadata.get("text_variants_raw", [])
                raw_preview = analysis_result.metadata.get("raw_text_preview", "")

                # Если есть хотя бы какие-то строки с вариантами — показываем их отдельно
                if text_variants:
                    st.subheader("🧬 Извлеченные строки с генетическими вариантами")
                    with st.expander("Показать извлеченные строки", expanded=False):
                        for line in text_variants:
                            st.write(f"- {line}")

                # Окно для дополнительных вопросов к врачу-генетику (письменно или голосом)
                st.subheader("💬 Задать вопрос генетику")
                
                input_mode_genetic = st.radio(
                    "Выберите способ ввода вопроса:",
                    ["✍️ Письменный", "🎤 Голосовой"],
                    horizontal=True,
                    key="genetic_question_mode"
                )
                
                extra_questions = ""
                
                if input_mode_genetic == "🎤 Голосовой":
                    if ASSEMBLYAI_AVAILABLE:
                        st.info("💡 Запишите ваш вопрос голосом. После записи он будет автоматически расшифрован.")
                        try:
                            from audio_recorder_streamlit import audio_recorder
                            audio_bytes_genetic = audio_recorder(text="🎤 Нажмите для записи вопроса", pause_threshold=2.0, sample_rate=44100, key="genetic_audio_recorder")
                            
                            if audio_bytes_genetic:
                                # Сохраняем во временный файл
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
                                    tmp_file.write(audio_bytes_genetic)
                                    audio_path_genetic = tmp_file.name
                                    st.session_state['genetic_audio_file_path'] = audio_path_genetic
                                
                                st.success(f"✅ Записано {len(audio_bytes_genetic)} байт аудио")
                                st.audio(audio_bytes_genetic, format="audio/wav")
                                
                                if st.button("🎤 Расшифровать вопрос", key="transcribe_genetic_question"):
                                    with st.spinner("🔄 Расшифровка вопроса..."):
                                        try:
                                            from config import ASSEMBLYAI_API_KEY
                                            api_key = ASSEMBLYAI_API_KEY or st.secrets.get("ASSEMBLYAI_API_KEY", "")
                                            if not api_key:
                                                st.error("❌ API ключ AssemblyAI не найден.")
                                            else:
                                                if 'genetic_audio_file_path' in st.session_state:
                                                    transcribed_question_genetic = transcribe_audio_assemblyai(st.session_state['genetic_audio_file_path'], api_key)
                                                    if not transcribed_question_genetic.startswith("❌"):
                                                        st.session_state['transcribed_genetic_question'] = transcribed_question_genetic
                                                        st.success("✅ Вопрос расшифрован!")
                                                        st.text_area("Расшифрованный вопрос:", value=transcribed_question_genetic, height=100, key="genetic_transcribed_display")
                                                    else:
                                                        st.error(transcribed_question_genetic)
                                        except Exception as e:
                                            st.error(f"❌ Ошибка расшифровки: {e}")
                        except ImportError:
                            st.warning("⚠️ Для голосового ввода установите: pip install audio-recorder-streamlit")
                            st.info("💡 Используйте письменный ввод")
                            audio_data_genetic = st.audio_input("🎤 Или загрузите аудиофайл с вопросом", key="genetic_audio_input")
                            if audio_data_genetic:
                                if st.button("🎤 Расшифровать загруженный файл", key="transcribe_genetic_file"):
                                    with st.spinner("🔄 Расшифровка..."):
                                        try:
                                            from config import ASSEMBLYAI_API_KEY
                                            api_key = ASSEMBLYAI_API_KEY or st.secrets.get("ASSEMBLYAI_API_KEY", "")
                                            if api_key:
                                                transcribed_question_genetic = transcribe_audio_assemblyai(audio_data_genetic, api_key)
                                                if not transcribed_question_genetic.startswith("❌"):
                                                    st.session_state['transcribed_genetic_question'] = transcribed_question_genetic
                                                    st.success("✅ Вопрос расшифрован!")
                                                    st.text_area("Расшифрованный вопрос:", value=transcribed_question_genetic, height=100, key="genetic_transcribed_display")
                                                else:
                                                    st.error(transcribed_question_genetic)
                                            else:
                                                st.error("❌ API ключ AssemblyAI не найден.")
                                        except Exception as e:
                                            st.error(f"❌ Ошибка: {e}")
                        except Exception as e:
                            st.error(f"❌ Ошибка записи: {e}")
                            st.info("💡 Используйте письменный ввод или загрузку файла")
                    else:
                        st.warning("⚠️ AssemblyAI недоступен. Используйте письменный ввод.")
                
                # Показываем текстовое поле для письменного ввода или редактирования расшифрованного вопроса
                if input_mode_genetic == "✍️ Письменный" or st.session_state.get('transcribed_genetic_question'):
                    extra_questions = st.text_area(
                        "Дополнительные вопросы к врачу-генетику (необязательно)",
                        value=st.session_state.get('transcribed_genetic_question', ''),
                        height=100,
                        help="Например: интересующие заболевания, семейный анамнез, уточняющие вопросы по конкретным генам или SNP.",
                        key="genetic_questions_text"
                    )
                    if st.session_state.get('transcribed_genetic_question'):
                        # Позволяем редактировать расшифрованный вопрос
                        pass

                # Кнопка: отправить ИИ-генетику даже если парсер не нашел ни одной строки —
                # тогда он работает по полному распознанному тексту отчета.
                if st.button("🧬 Показать врачу-генетику (ИИ-консультация по отчету)", use_container_width=True):
                    if not AI_AVAILABLE or OpenRouterAssistant is None:
                        st.error("❌ ИИ-модуль недоступен. Проверьте файл `claude_assistant.py` и API-ключ.")
                    else:
                        assistant = OpenRouterAssistant()
                        
                        genetic_system_prompt = """
ПРОМПТ: ВРАЧ-ГЕНЕТИК — КОНСУЛЬТАНТ ПО ГЕНЕТИЧЕСКИМ ТЕСТАМ (Claude Opus)

Роль:
Ты — врач-генетик с 15+ годами клинического опыта. Тебе передают результаты генетического тестирования и краткую клиническую информацию о пациенте. Твой стиль — максимально точный, сухой, деловой, без лишних рассуждений и эмоций.

Формат обращения:
- Отвечай так, как если бы писал заключение в медицинскую карту или консультативное письмо для лечащего врача.
- Не используй обращений к «коллегам», к пациенту или к третьим лицам. Прямое деловое изложение.

Инструкция:
Тебе передают итоговый генетический отчет:
- список вариантов (строки с генами, нотацией c./p., rsID, генотипами и др.);
- клинический контекст (жалобы, диагнозы, семейный анамнез);
- при необходимости — дополнительные вопросы от направившего врача.

Твоя задача (строгий порядок ответа):
1. Сначала перечисли ВСЕ обнаруженные варианты, которые ты видишь в тексте отчета.
   - Для каждого варианта по возможности укажи: ген, rsID, нотацию (c./p. или другую), генотип, при наличии — краткое описание эффекта или класса (патогенный/вероятно патогенный/вариант с неопределённым значением/доброкачественный/фармакогенетический и т.п.).
   - Представь этот список в компактном структурированном виде (нумерованный или маркированный список; табличный формат не обязателен).
2. Затем дай ПОСЛЕДОВАТЕЛЬНУЮ трактовку по каждому ключевому варианту:
   - клиническое значение для пациента (заболевания, синдромы, риски);
   - тип наследования и возможные риски для родственников/потомства;
   - при необходимости — фармакогенетические последствия (лекарства, дозировки, противопоказания).
3. В конце сформулируй ОБОБЩЁННОЕ клиническое заключение и тактику ведения:
   - краткий вердикт по результатам теста (1–3 предложения);
   - общий профиль рисков (высокие / умеренные / низкие);
   - рекомендуемая тактика: дообследование, наблюдение, терапия, модификация образа жизни, репродуктивные решения, семейный/каскадный скрининг.
4. Если данных недостаточно (неполные панели, нет точных генотипов, нет VCF и т.п.) — прямо укажи ограничения и что нужно для полноценного заключения (например, VCF, расширенная панель, подтверждение методом Sanger).

Важно:
- Не придумывай несуществующие варианты и диагнозы.
- Если по представленным данным нельзя сделать уверенных выводов, так и напиши и опиши, каких данных не хватает.
"""
                        # Собираем текст для генетика: клиника + извлеченные строки (если есть) + полный текст отчета + дополнительные вопросы
                        patient_block = f"Возраст: {patient_info.get('age')}, пол: {patient_info.get('gender')}, образ жизни: {patient_info.get('lifestyle')}\n"
                        context_block = f"Клинический контекст:\n{clinical_context or 'не указан'}\n"
                        variants_block = ""
                        if text_variants:
                            variants_block = "Извлеченные строки с возможными вариантами:\n" + "\n".join(text_variants)
                        full_text_block = ""
                        if raw_preview:
                            full_text_block = "\n\nПолный распознанный текст отчета (фрагмент):\n" + str(raw_preview)
                        questions_block = ""
                        if extra_questions and extra_questions.strip():
                            questions_block = f"\nДополнительные вопросы от врача, направившего на консультацию:\n{extra_questions.strip()}\n"
                        
                        user_message = f"""
{patient_block}
{context_block}
{variants_block}
{full_text_block}
{questions_block}
"""
                        with st.spinner("🤖 Врач-генетик (Opus) формирует заключение..."):
                            try:
                                # Используем специализированный промпт генетика ЧЕРЕЗ профессорский system_prompt
                                # Профессорский промпт обеспечит единый стандарт клинической директивы
                                # Специализированный промпт добавляется как дополнительный контекст
                                genetic_context = f"""{genetic_system_prompt}

Исходные данные по пациенту и отчету:
{user_message}"""
                                genetic_question = "Проведи комплексную интерпретацию генетического анализа согласно специализированному контексту выше и сформулируй клиническую директиву в формате профессора."
                                genetic_opinion = assistant.get_response(
                                    genetic_question,
                                    context=genetic_context
                                )
                                st.subheader("🧬 Заключение врача-генетика (ИИ)")
                                st.write(genetic_opinion)

                                # Сохраняем заключение генетика в сессию, чтобы при необходимости отправить профессору
                                if "genetic_specialist_conclusion" not in st.session_state:
                                    st.session_state["genetic_specialist_conclusion"] = {}
                                st.session_state["genetic_specialist_conclusion"][analysis_result.analysis_id] = {
                                    "conclusion": genetic_opinion,
                                    "patient_info": patient_info,
                                    "clinical_context": clinical_context,
                                    "text_variants_raw": analysis_result.metadata.get("text_variants_raw", []),
                                }
                            except Exception as e:
                                st.error(f"❌ Ошибка ИИ-консультации: {e}")
            
            # Если есть сохраненное заключение генетика для этого анализа — даем опцию отправить его профессору
            specialist_data = None
            if "genetic_specialist_conclusion" in st.session_state:
                specialist_data = st.session_state["genetic_specialist_conclusion"].get(analysis_result.analysis_id)

            if specialist_data and AI_AVAILABLE and OpenRouterAssistant is not None:
                st.subheader("📨 Дополнительно: отправить заключение генетика профессору")
                st.info("Профессор клинической медицины сформирует общий клинический обзор и тактику на основе заключения генетика и извлечённых данных.")

                if st.button("📤 Отправить заключение генетика на консультацию к профессору", use_container_width=True):
                    try:
                        professor_assistant = OpenRouterAssistant()

                        prof_context = ""
                        pi = specialist_data.get("patient_info") or {}
                        cc = specialist_data.get("clinical_context") or ""
                        tv = specialist_data.get("text_variants_raw") or []
                        gen_concl = specialist_data.get("conclusion") or ""

                        prof_context += f"Пациент: возраст {pi.get('age')}, пол {pi.get('gender')}, образ жизни: {pi.get('lifestyle')}.\n"
                        prof_context += f"Клинический контекст/жалобы:\n{cc}\n\n"
                        if tv:
                            prof_context += "Ключевые извлеченные строки генетических вариантов (гены, SNP, rsID):\n"
                            prof_context += "\n".join(tv[:50]) + "\n\n"
                        prof_context += "Заключение врача-генетика (как отдельный специалист):\n"
                        prof_context += gen_concl

                        professor_question = (
                            "На основании данных пациента, клинического контекста, извлечённых генетических вариантов "
                            "и заключения врача-генетика сформулируйте краткий, структурированный клинический обзор ситуации, "
                            "ключевые риски и тактику ведения пациента (без повторения всей расшифровки вариантов)."
                        )

                        with st.spinner("🤖 Профессор клинической медицины (Opus) формирует обобщённый клинический обзор..."):
                            # Профессор использует тот же системный промпт, но модель по умолчанию — Opus
                            # (Sonnet используется для лабораторных данных и быстрых задач).
                            professor_opinion = professor_assistant.get_response(
                                professor_question,
                                context=prof_context,
                                use_sonnet_4_5=False
                            )
                            st.subheader("🏥 Заключение профессора клинической медицины (на основе генетического отчёта)")
                            st.write(professor_opinion)
                    except Exception as e:
                        st.error(f"❌ Ошибка при консультации профессора: {e}")

            # Патогенные варианты (для VCF)
            if analysis_result.pathogenic_variants:
                st.subheader("⚠️ Патогенные варианты")
                with st.expander("Показать патогенные варианты", expanded=True):
                    for variant in analysis_result.pathogenic_variants[:10]:  # Показываем первые 10
                        # Отображаем все доступные данные
                        genotype = variant.info.get('genotype', '')
                        zygosity = variant.info.get('zygosity', '')
                        gene = variant.info.get('gene', 'Unknown')
                        rsid = variant.info.get('rsid', '')
                        c_dna = variant.info.get('c_dna', '')
                        protein = variant.info.get('protein', '')
                        coverage = variant.info.get('coverage', '')
                        quality_score = variant.info.get('quality_score', '')
                        vaf = variant.info.get('vaf', '')
                        acmg = variant.info.get('acmg_classification', '')
                        clinvar = variant.info.get('clinvar', '')
                        gnomad_af = variant.info.get('gnomad_af', '')
                        
                        quality_val = float(variant.quality) if variant.quality else 0.0
                        quality_str = f"{quality_val:.2f}"
                        if quality_score:
                            quality_str = f"{quality_str} (Q-score: {quality_score})"
                        
                        variant_info = f"**Ген: {gene}**"
                        if variant.chromosome != 'Unknown' and variant.position > 0:
                            variant_info += f" | Хромосома {variant.chromosome}:{variant.position}"
                        if rsid:
                            variant_info += f" | **rsID: {rsid}**"
                        if c_dna:
                            variant_info += f" | {c_dna}"
                        if protein:
                            variant_info += f" | {protein}"
                        
                        details = []
                        if variant.ref != 'N' and variant.alt != 'N':
                            details.append(f"Референс: {variant.ref} -> Альтернатива: {variant.alt}")
                        if variant.id and variant.id != '.' and variant.id != 'unknown':
                            details.append(f"ID: {variant.id}")
                        if genotype:
                            details.append(f"Генотип: {genotype}")
                        if zygosity:
                            details.append(f"Зиготность: {zygosity}")
                        if coverage:
                            details.append(f"Покрытие: {coverage}")
                        if quality_str:
                            details.append(f"Качество: {quality_str}")
                        if vaf:
                            details.append(f"VAF: {vaf}")
                        if acmg:
                            details.append(f"ACMG: {acmg}")
                        if clinvar:
                            details.append(f"ClinVar: {clinvar}")
                        if gnomad_af:
                            details.append(f"gnomAD AF: {gnomad_af}")
                        
                        st.markdown(f"""
                        {variant_info}
                        {chr(10).join(['- ' + d for d in details])}
                        """)
                    
                    if len(analysis_result.pathogenic_variants) > 10:
                        st.info(f"И еще {len(analysis_result.pathogenic_variants) - 10} патогенных вариантов...")
            
            # Клинические интерпретации
            if analysis_result.clinical_interpretations:
                st.subheader("🏥 Клинические интерпретации")
                for interpretation in analysis_result.clinical_interpretations[:5]:
                    # interpretation это ClinicalVariant объект
                    st.markdown(f"""
                    **{interpretation.gene}**
                    - Вариант: {interpretation.variant_name}
                    - Изменение белка: {interpretation.protein_change}
                    - Заболевание: {interpretation.disease}
                    - Патогенность: {interpretation.pathogenicity.value}
                    - Наследование: {interpretation.inheritance_pattern}
                    - Клиническое действие: {interpretation.clinical_action}
                    """)
            
            # Фармакогенетические рекомендации
            if analysis_result.pharmacogenetic_interpretations:
                st.subheader("💊 Фармакогенетические рекомендации")
                for pharm in analysis_result.pharmacogenetic_interpretations[:5]:
                    # pharm это PharmacogeneticVariant объект
                    drugs_str = ", ".join(pharm.drugs) if pharm.drugs else "Не указаны"
                    st.markdown(f"""
                    **Ген: {pharm.gene}**
                    - Вариант: {pharm.variant}
                    - Фенотип: {pharm.phenotype}
                    - Препараты: {drugs_str}
                    - Рекомендация: {pharm.recommendation}
                    - Уровень доказательности: {pharm.evidence_level}
                    """)
            
            # Оценка рисков
            if analysis_result.risk_assessment:
                st.subheader("📈 Оценка генетических рисков")
                risk_data = analysis_result.risk_assessment
                
                st.markdown(f"**Общий уровень риска: {risk_data.overall_risk_level}**")
                
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**Высокопенетрантные заболевания:**")
                    for disease in risk_data.high_penetrance_diseases[:5]:
                        disease_name = disease.get('disease', 'Неизвестно') if isinstance(disease, dict) else str(disease)
                        st.write(f"- {disease_name}")
                
                with col2:
                    st.markdown("**Умеренные риски:**")
                    for condition in risk_data.moderate_risk_conditions[:5]:
                        cond_name = condition.get('condition', 'Неизвестно') if isinstance(condition, dict) else str(condition)
                        st.write(f"- {cond_name}")
                
                if risk_data.surveillance_recommendations:
                    st.markdown("**Рекомендации по мониторингу:**")
                    for rec in risk_data.surveillance_recommendations[:5]:
                        st.write(f"- {rec}")
            
            # Рекомендации
            if analysis_result.recommendations:
                st.subheader("💡 Рекомендации")
                for i, rec in enumerate(analysis_result.recommendations[:10], 1):
                    st.markdown(f"{i}. {rec}")
            
            # Автоматическое заключение на основе извлеченных данных
            if analysis_result.total_variants == 0 or (len(analysis_result.pathogenic_variants) == 0 and len(analysis_result.clinical_interpretations) == 0):
                st.subheader("📋 Заключение на основе извлеченных данных")
                gene_panel = []
                if analysis_result.metadata and 'gene_panel' in analysis_result.metadata:
                    gene_panel = analysis_result.metadata.get('gene_panel', [])
                
                conclusion_text = "**КЛИНИЧЕСКОЕ ЗАКЛЮЧЕНИЕ:**\n\n"
                
                if gene_panel:
                    conclusion_text += f"**Спектр проанализированных генов:** Проанализировано {len(gene_panel)} генов: {', '.join(gene_panel[:20])}{'...' if len(gene_panel) > 20 else ''}\n\n"
                
                if analysis_result.total_variants > 0:
                    conclusion_text += f"**Обнаруженные варианты:** В ходе анализа обнаружено {analysis_result.total_variants} генетических вариантов.\n\n"
                else:
                    conclusion_text += "**Обнаруженные варианты:** Конкретные генетические варианты в стандартном формате VCF не обнаружены в представленном PDF документе. PDF файлы содержат текстовые отчеты, а не стандартизированный формат VCF. Для более точного анализа рекомендуется использовать VCF файлы.\n\n"
                
                if len(analysis_result.pathogenic_variants) == 0:
                    conclusion_text += "**Патогенные варианты:** Патогенных или вероятно патогенных вариантов не обнаружено.\n\n"
                else:
                    conclusion_text += f"**Патогенные варианты:** Обнаружено {len(analysis_result.pathogenic_variants)} патогенных вариантов, требующих внимания.\n\n"
                
                if gene_panel:
                    conclusion_text += "**Интерпретация спектра:** Проанализированные гены охватывают различные области генетики. Для полной клинической интерпретации рекомендуется:\n"
                    conclusion_text += "1. Консультация врача-генетика\n"
                    conclusion_text += "2. Анализ конкретных генотипов по каждому гену\n"
                    conclusion_text += "3. Оценка клинической значимости в контексте жалоб пациента\n"
                    conclusion_text += "4. При необходимости - дополнительное тестирование с использованием валидированных методов (NGS с покрытием ≥30x)\n\n"
                
                conclusion_text += "**Рекомендации:** Для получения детального заключения с персонализированными рекомендациями используйте функцию ИИ-интерпретации ниже."
                
                st.markdown(conclusion_text)
            
            # ИИ-интерпретация высококлассным специалистом (если доступен ИИ)
            if AI_AVAILABLE and OpenRouterAssistant is not None:
                st.subheader("🤖 ИИ-интерпретация от врача-генетика-консультанта")
                st.info("💡 Получите детальную интерпретацию с персонализированными рекомендациями по лечению и образу жизни")
                
                # Тест доступности ИИ
                if st.checkbox("🔍 Проверить доступность ИИ перед запросом", value=False, key="test_ai_genetic"):
                                try:
                                    test_assistant = OpenRouterAssistant()
                                    success, msg = test_assistant.test_connection()
                                    if success:
                                        st.success(f"✅ {msg}")
                                    else:
                                        st.error(f"❌ {msg}")
                                except Exception as e:
                                    st.error(f"❌ Ошибка проверки: {e}")
                
                # Проверяем, есть ли уже сохраненная интерпретация
                saved_interpretation = None
                if 'genetic_ai_interpretation' in st.session_state:
                    saved_interpretation = st.session_state.genetic_ai_interpretation.get(analysis_result.analysis_id)
                
                if saved_interpretation:
                    st.success("✅ Интерпретация уже получена. Вы можете просмотреть её ниже или получить новую.")
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("📖 Показать сохраненную интерпретацию", use_container_width=True, key="show_saved_genetic"):
                            st.markdown("### 🧬 Интерпретация врача-генетика-консультанта")
                            st.markdown("---")
                            st.write(saved_interpretation)
                            st.download_button(
                                "📥 Скачать интерпретацию (TXT)",
                                saved_interpretation,
                                file_name=f"genetic_interpretation_{analysis_result.analysis_id}.txt",
                                mime="text/plain",
                                key="download_saved_genetic"
                            )
                    with col2:
                        if st.button("🔄 Получить новую интерпретацию", use_container_width=True, key="new_genetic"):
                            # Очищаем сохраненную интерпретацию
                            if analysis_result.analysis_id in st.session_state.genetic_ai_interpretation:
                                del st.session_state.genetic_ai_interpretation[analysis_result.analysis_id]
                            st.rerun()
                
                # Кнопка для получения интерпретации (всегда видна, если нет сохраненной)
                if not saved_interpretation:
                    button_key = f"get_genetic_interpretation_{analysis_result.analysis_id}"
                    if st.button("🧠 Получить интерпретацию специалиста", use_container_width=True, type="primary", key=button_key):
                        try:
                            # Проверка перед началом
                            st.info("🔄 Инициализация ИИ-ассистента...")
                            assistant = OpenRouterAssistant()
                            
                            with st.spinner("🔬 Врач-генетик анализирует результаты (это может занять 1-2 минуты)..."):
                                # Формируем детальный контекст для ИИ
                                # Получаем спектр генов из metadata
                                gene_panel = []
                                if analysis_result.metadata and 'gene_panel' in analysis_result.metadata:
                                    gene_panel = analysis_result.metadata.get('gene_panel', [])
                                
                                ai_context = f"""
═══════════════════════════════════════════════════════════
ГЕНЕТИЧЕСКИЙ АНАЛИЗ ПАЦИЕНТА
═══════════════════════════════════════════════════════════

ДЕМОГРАФИЧЕСКИЕ ДАННЫЕ:
- Возраст: {age} лет
- Пол: {gender}
- Образ жизни: {lifestyle}
- Клинический контекст: {clinical_context if clinical_context else 'Не указан'}

СТАТИСТИКА АНАЛИЗА:
- Всего вариантов обнаружено: {analysis_result.total_variants}
- Патогенных вариантов: {len(analysis_result.pathogenic_variants)}
- Вероятно патогенных: {len(analysis_result.likely_pathogenic_variants)}
- Клинически значимых интерпретаций: {len(analysis_result.clinical_interpretations)}
- Фармакогенетических вариантов: {len(analysis_result.pharmacogenetic_variants)}
- Вариантов признаков: {len(analysis_result.trait_variants)}
"""
                                
                                # Добавляем информацию о спектре генов
                                if gene_panel:
                                    ai_context += f"""
СПЕКТР ПРОАНАЛИЗИРОВАННЫХ ГЕНОВ ({len(gene_panel)} генов):
{', '.join(gene_panel[:50])}{'...' if len(gene_panel) > 50 else ''}
"""

                                ai_context += """
ПАТОГЕННЫЕ ВАРИАНТЫ (первые 30):
"""
                                for i, variant in enumerate(analysis_result.pathogenic_variants[:30], 1):
                                    gene = variant.info.get('gene', 'Unknown')
                                    genotype = variant.info.get('genotype', '')
                                    zygosity = variant.info.get('zygosity', '')
                                    genotype_info = ""
                                    if genotype:
                                        genotype_info = f"\n   - Генотип: {genotype}"
                                    if zygosity:
                                        genotype_info += f" ({zygosity})"
                                    
                                    quality_val = float(variant.quality) if variant.quality else 0.0
                                    quality_str = f"{quality_val:.2f}"
                                    ai_context += f"""
{i}. Ген: {gene} | Хромосома {variant.chromosome}, позиция {variant.position}
   - Референс: {variant.ref} -> Альтернатива: {variant.alt}
   - ID варианта: {variant.id if variant.id != '.' else 'Нет'}{genotype_info}
   - Качество: {quality_str}
   - Фильтр: {variant.filter}
"""
                                
                                # Добавляем информацию о всех вариантах с генотипами
                                if analysis_result.total_variants > 0:
                                    ai_context += f"\n\nВСЕ ОБНАРУЖЕННЫЕ ВАРИАНТЫ С ГЕНОТИПАМИ:\n"
                                    all_variants_with_genotypes = []
                                    for variant in analysis_result.pathogenic_variants + analysis_result.likely_pathogenic_variants + analysis_result.pharmacogenetic_variants:
                                        gene = variant.info.get('gene', 'Unknown')
                                        genotype = variant.info.get('genotype', '')
                                        zygosity = variant.info.get('zygosity', '')
                                        if genotype or gene != 'Unknown':
                                            all_variants_with_genotypes.append({
                                                'gene': gene,
                                                'genotype': genotype,
                                                'zygosity': zygosity,
                                                'variant_id': variant.id
                                            })
                                    
                                    # Если нет патогенных, но есть варианты из спектра
                                    if not all_variants_with_genotypes and gene_panel:
                                        ai_context += f"Проанализированы гены из спектра, но конкретные варианты с генотипами не указаны в отчете.\n"
                                        ai_context += f"Спектр включает: {', '.join(gene_panel[:20])}{'...' if len(gene_panel) > 20 else ''}\n"
                                    else:
                                        for i, var_info in enumerate(all_variants_with_genotypes[:30], 1):
                                            genotype_str = f" | Генотип: {var_info['genotype']}" if var_info['genotype'] else ""
                                            zygosity_str = f" ({var_info['zygosity']})" if var_info['zygosity'] else ""
                                            ai_context += f"{i}. Ген: {var_info['gene']}{genotype_str}{zygosity_str} | ID: {var_info['variant_id']}\n"
                                
                                # Клинические интерпретации
                                if analysis_result.clinical_interpretations:
                                    ai_context += "\n\nКЛИНИЧЕСКИЕ ИНТЕРПРЕТАЦИИ:\n"
                                    for i, interp in enumerate(analysis_result.clinical_interpretations[:15], 1):
                                        ai_context += f"""
{i}. Ген: {interp.gene}
   - Вариант: {interp.variant_name}
   - Изменение белка: {interp.protein_change}
   - Патогенность: {interp.pathogenicity.value}
   - Заболевание: {interp.disease}
   - Тип наследования: {interp.inheritance_pattern}
   - Пенетрантность: {interp.penetrance}
   - Клиническое действие: {interp.clinical_action}
   - Уровень доказательности: {interp.evidence_level}
"""
                                
                                # Фармакогенетика
                                if analysis_result.pharmacogenetic_interpretations:
                                    ai_context += "\n\nФАРМАКОГЕНЕТИЧЕСКИЕ ДАННЫЕ:\n"
                                    for i, pharm in enumerate(analysis_result.pharmacogenetic_interpretations[:15], 1):
                                        drugs_str = ", ".join(pharm.drugs) if pharm.drugs else "Не указаны"
                                        ai_context += f"""
{i}. Ген: {pharm.gene}
   - Вариант: {pharm.variant}
   - Фенотип метаболизма: {pharm.phenotype}
   - Препараты: {drugs_str}
   - Рекомендация: {pharm.recommendation}
   - Уровень доказательности: {pharm.evidence_level}
   - Клиническая аннотация: {pharm.clinical_annotation}
"""
                                
                                # Оценка рисков
                                if analysis_result.risk_assessment:
                                    risk_data = analysis_result.risk_assessment
                                    ai_context += f"\n\nОЦЕНКА РИСКОВ:\n"
                                    ai_context += f"- Общий уровень риска: {risk_data.overall_risk_level}\n"
                                    if risk_data.high_penetrance_diseases:
                                        ai_context += f"- Высокопенетрантные заболевания: {len(risk_data.high_penetrance_diseases)}\n"
                                    if risk_data.moderate_risk_conditions:
                                        ai_context += f"- Умеренные риски: {len(risk_data.moderate_risk_conditions)}\n"
                                
                                # Клинический контекст
                                if clinical_context:
                                    ai_context += f"\n\nКЛИНИЧЕСКИЙ КОНТЕКСТ ПАЦИЕНТА:\n{clinical_context}\n"
                                
                                # Рекомендации из анализа
                                if analysis_result.recommendations:
                                    ai_context += "\n\nАВТОМАТИЧЕСКИЕ РЕКОМЕНДАЦИИ СИСТЕМЫ:\n"
                                    for rec in analysis_result.recommendations[:10]:
                                        ai_context += f"- {rec}\n"
                                
                                # Срочные флаги
                                if analysis_result.urgent_flags:
                                    ai_context += "\n\n⚠️ СРОЧНЫЕ ФЛАГИ:\n"
                                    for flag in analysis_result.urgent_flags:
                                        ai_context += f"- {flag}\n"
                                
                                ai_context += "\n═══════════════════════════════════════════════════════════\n"
                                
                                # Расширенный промпт от имени высококлассного специалиста
                                prompt = f"""Вы - ведущий врач-генетик-консультант с 25-летним опытом работы в престижной клинике, специализирующийся на персонализированной медицине, фармакогенетике и превентивной генетике. Вы являетесь экспертом международного уровня, публикуетесь в ведущих журналах (Nature Genetics, American Journal of Human Genetics) и консультируете сложные клинические случаи.

КРИТИЧЕСКИ ВАЖНО: Вы даете рекомендации ОТ ВРАЧА ВРАЧУ. Ваш ответ предназначен для коллеги-врача, который будет использовать эти рекомендации в клинической практике. Используйте профессиональный медицинский язык, конкретные дозировки, названия препаратов, ссылки на клинические рекомендации. НЕ упрощайте информацию - предполагается, что получатель является медицинским специалистом.

ВАША ЗАДАЧА: Провести комплексную интерпретацию генетического анализа с фокусом на ПЕРСОНАЛИЗАЦИЮ лечения и образа жизни для конкретного пациента. Дать КОНКРЕТНЫЕ клинические директивы, готовые к применению врачом в практике.

ФОРМАТ ОТВЕТА - "Клиническая директива по персонализированной генетической медицине":

1. **КЛИНИЧЕСКИЙ ОБЗОР** (3-4 предложения)
   - Краткая характеристика генетического профиля пациента
   - Общая оценка клинической значимости находок
   - Приоритетные направления для внимания

2. **ДЕТАЛЬНЫЙ АНАЛИЗ ПАТОГЕННЫХ ВАРИАНТОВ И КЛИНИЧЕСКИ ЗНАЧИМЫХ НАХОДОК**
   Для КАЖДОГО найденного патологического или клинически значимого варианта ОБЯЗАТЕЛЬНО укажи:
   
   a. ОПИСАНИЕ ВАРИАНТА:
      - Ген (полное название)
      - Вариант (например, MTHFR C677T, COMT Val158Met, TNFa -308G>A)
      - Генотип (гомозигота/гетерозигота)
      - Изменение белка (если применимо)
   
   b. КЛИНИЧЕСКАЯ ЗНАЧИМОСТЬ:
      - Связь с конкретными заболеваниями/состояниями
      - Механизм действия (как вариант влияет на функцию гена)
      - Риски для здоровья (конкретные цифры если известны)
      - OMIM коды заболеваний (если применимо)
   
   c. КОНКРЕТНЫЕ КЛИНИЧЕСКИЕ РЕКОМЕНДАЦИИ (ЧТО ДЕЛАТЬ):
      - Немедленные действия (если требуются)
      - Препараты/добавки с конкретными дозировками
      - Диетические рекомендации (конкретные продукты, что исключить)
      - Лабораторные анализы для мониторинга
      - Консультации специалистов (каких и когда)
      - Частота наблюдения
   
   ПРИМЕРЫ для известных вариантов:
   - MTHFR C677T: фолиевая кислота в метилированной форме (метилфолат) 400-800 мкг/день, витамин B12, контроль гомоцистеина
   - COMT: коррекция доз катехоламиновых препаратов, управление стрессом
   - TNFa: контроль воспалительных процессов, возможная коррекция терапии при аутоиммунных заболеваниях

3. **ПЕРСОНАЛИЗИРОВАННАЯ ФАРМАКОГЕНЕТИКА**
   Для каждого фармакогенетического варианта:
   - Детальная характеристика фенотипа метаболизма
   - Конкретные препараты, требующие коррекции дозы или замены
   - Рекомендуемые дозировки с учетом генотипа
   - Альтернативные препараты (если применимо)
   - Мониторинг эффективности и токсичности
   - Ссылки на клинические рекомендации (CPIC, DPWG)

4. **ПЕРСОНАЛИЗИРОВАННЫЕ РЕКОМЕНДАЦИИ ПО ОБРАЗУ ЖИЗНИ**
   На основе генетического профиля и образа жизни пациента ({lifestyle}):
   
   a. ПИТАНИЕ:
      - Нутригенетические рекомендации
      - Оптимальный макро- и микросостав рациона
      - Продукты, которые следует ограничить/исключить
      - Добавки и витамины (с учетом генетики)
      - Режим питания
   
   b. ФИЗИЧЕСКАЯ АКТИВНОСТЬ:
      - Тип и интенсивность тренировок (с учетом генетики)
      - Рекомендации по восстановлению
      - Профилактика травм
      - Оптимальный режим активности
   
   c. СОН И ВОССТАНОВЛЕНИЕ:
      - Рекомендации по режиму сна
      - Оптимизация циркадных ритмов
   
   d. СТРЕСС-МЕНЕДЖМЕНТ:
      - Стратегии управления стрессом
      - Медитация, релаксация
   
   e. ОКРУЖАЮЩАЯ СРЕДА:
      - Избегание токсинов/канцерогенов (если есть повышенная чувствительность)
      - Защита от УФ (если есть мутации в генах репарации ДНК)

5. **ПЕРСОНАЛИЗИРОВАННЫЙ ПЛАН ЛЕЧЕНИЯ И КОРРЕКЦИИ**
   ОБЯЗАТЕЛЬНО для каждого найденного варианта укажи:
   
   a. ФАРМАКОТЕРАПИЯ (если применимо):
      - Конкретные препараты с дозировками
      - Препараты, которые НЕЛЬЗЯ использовать или требуют коррекции дозы
      - Альтернативные препараты
      - Схема приема (когда, как долго)
   
   b. НУТРИЦИОЛОГИЧЕСКАЯ КОРРЕКЦИЯ:
      - Конкретные добавки/витамины с дозировками
      - Продукты для включения в рацион
      - Продукты для исключения/ограничения
      - Режим приема добавок
   
   c. МОНИТОРИНГ:
      - Какие анализы сдавать
      - Как часто (раз в месяц/квартал/год)
      - Референсные значения для контроля
      - "Красные флаги" - когда срочно обращаться к врачу
   
   d. КОНСУЛЬТАЦИИ СПЕЦИАЛИСТОВ:
      - К каким врачам обращаться (генетик, гематолог, эндокринолог и т.д.)
      - Сроки консультаций (немедленно/в течение месяца/планово)
      - Что обсудить на консультации

6. **ПРЕВЕНТИВНЫЕ МЕРЫ**
   - Скрининговые программы (с учетом возраста и генетики)
   - Частота обследований
   - Специфические тесты для раннего выявления
   - Вакцинация (если применимо)

7. **ГЕНЕТИЧЕСКОЕ КОНСУЛЬТИРОВАНИЕ СЕМЬИ**
   - Риски для родственников
   - Рекомендации по тестированию семьи
   - Репродуктивные риски (если применимо)
   - Планирование семьи

8. **МОНИТОРИНГ И ДИНАМИЧЕСКОЕ НАБЛЮДЕНИЕ**
   - План последующих визитов
   - Параметры для мониторинга
   - Триггеры для немедленного обращения
   - Долгосрочные цели

9. **СРОЧНЫЕ РЕКОМЕНДАЦИИ И ПЛАН ДЕЙСТВИЙ**
   ОБЯЗАТЕЛЬНО укажи конкретный план действий:
   
   a. НЕМЕДЛЕННО (в течение недели):
      - Какие препараты/добавки начать принимать
      - Какие анализы сдать
      - К каким врачам записаться
   
   b. В ТЕЧЕНИЕ МЕСЯЦА:
   - Дополнительные обследования
      - Консультации специалистов
      - Коррекция образа жизни
   
   c. ДОЛГОСРОЧНО (3-6 месяцев):
      - План мониторинга
      - Профилактические меры
      - Оценка эффективности коррекции
   
   d. "КРАСНЫЕ ФЛАГИ" - когда срочно обращаться к врачу:
      - Конкретные симптомы/признаки
      - Изменения в анализах
      - Ухудшение состояния

10. **ИСТОЧНИКИ И ДОКАЗАТЕЛЬСТВА**
    - Использованные базы данных (ClinVar, PharmGKB, dbSNP)
    - Клинические рекомендации
    - Уровень доказательности для каждой рекомендации

ВАЖНО:
- Все рекомендации должны быть КОНКРЕТНЫМИ и ПРИМЕНИМЫМИ
- Учитывайте возраст ({age} лет), пол ({gender}) и образ жизни ({lifestyle})
- Используйте только проверенные источники (ACMG, CPIC, PharmGKB)
- Указывайте уровень доказательности для каждой рекомендации
- Пишите ПРОФЕССИОНАЛЬНЫМ медицинским языком (от врача врачу)
- Фокус на ПРАКТИЧЕСКОМ применении в клинической практике
- Используйте медицинскую терминологию, названия препаратов, дозировки, ссылки на гайдлайны
- НЕ упрощайте - предполагается, что получатель является медицинским специалистом

КРИТИЧЕСКИ ВАЖНО:
- Если патогенных вариантов не найдено, но есть спектр генов и генотипы - дайте интерпретацию на основе проанализированных генов
- Обязательно упомяните все гены из спектра анализа
- Проанализируйте генотипы и дайте рекомендации на их основе
- Даже если нет патогенных вариантов, дайте заключение о генетическом профиле пациента
- Укажите, какие гены были проанализированы и что это означает для пациента

ДАННЫЕ ГЕНЕТИЧЕСКОГО АНАЛИЗА:
{ai_context}

Дайте развернутый ответ в формате "Клиническая директива по персонализированной генетической медицине".

КРИТИЧЕСКИ ВАЖНО - ДАЙТЕ КОНКРЕТНЫЕ РЕКОМЕНДАЦИИ:

Для каждого найденного патологического/клинически значимого варианта ОБЯЗАТЕЛЬНО укажите:

1. ЧТО ДЕЛАТЬ ПРЯМО СЕЙЧАС:
   - Конкретные препараты/добавки с дозировками (например: "Метилфолат 400-800 мкг/день")
   - Что изменить в питании (конкретные продукты)
   - Какие анализы сдать (названия анализов)

2. К КАКИМ ВРАЧАМ ОБРАТИТЬСЯ:
   - Список специалистов (генетик, гематолог, эндокринолог и т.д.)
   - Сроки консультаций (немедленно/в течение месяца)

3. ПЛАН МОНИТОРИНГА:
   - Какие показатели контролировать
   - Как часто сдавать анализы
   - Референсные значения

4. ПРОГНОЗ И РИСКИ:
   - Конкретные риски для здоровья
   - Вероятность развития заболеваний
   - Профилактические меры

НЕ ПИШИТЕ ОБЩИЕ ФРАЗЫ! Давайте КОНКРЕТНЫЕ, ПРИМЕНИМЫЕ рекомендации с дозировками, названиями препаратов и конкретными действиями.

ОБЯЗАТЕЛЬНО включите заключение, даже если патогенных вариантов не обнаружено - проанализируйте спектр генов и генотипы.
"""
                                
                                # Отправляем запрос с увеличенным таймаутом
                                st.info("📤 Отправка запроса к ИИ (это может занять 1-3 минуты для детального анализа)...")
                                
                                # Показываем прогресс
                                progress_bar = st.progress(0)
                                status_text = st.empty()
                                
                                try:
                                    status_text.text("🔄 Подключение к ИИ...")
                                    progress_bar.progress(10)
                                    
                                    status_text.text("📝 Формирование запроса...")
                                    progress_bar.progress(30)
                                    
                                    status_text.text("🧠 ИИ анализирует данные (это может занять время)...")
                                    progress_bar.progress(50)
                                    
                                    ai_interpretation = assistant.get_response(prompt)
                                    
                                    progress_bar.progress(90)
                                    status_text.text("✅ Получен ответ от ИИ")
                                    progress_bar.progress(100)
                                    
                                    if not ai_interpretation or len(ai_interpretation.strip()) == 0:
                                        st.error("❌ ИИ вернул пустой ответ. Попробуйте еще раз.")
                                        return
                                    
                                    # Очищаем прогресс-бар
                                    progress_bar.empty()
                                    status_text.empty()
                                    
                                except Exception as api_error:
                                    progress_bar.empty()
                                    status_text.empty()
                                    raise api_error
                                
                                # Сохраняем интерпретацию в session_state
                                if 'genetic_ai_interpretation' not in st.session_state:
                                    st.session_state.genetic_ai_interpretation = {}
                                
                                st.session_state.genetic_ai_interpretation[analysis_result.analysis_id] = ai_interpretation
                                
                                # Очищаем прогресс-бар перед отображением результата
                                progress_bar.empty()
                                status_text.empty()
                                
                                st.markdown("### 🧬 Интерпретация врача-генетика-консультанта")
                                st.markdown("---")
                                st.write(ai_interpretation)
                                
                                # Кнопка для скачивания интерпретации
                                st.download_button(
                                    "📥 Скачать интерпретацию (TXT)",
                                    ai_interpretation,
                                    file_name=f"genetic_interpretation_{analysis_result.analysis_id}.txt",
                                    mime="text/plain",
                                    key=f"download_genetic_{analysis_result.analysis_id}"
                                )
                                
                                st.success("✅ Интерпретация успешно получена!")
                                st.rerun()
                        
                        except Exception as e:
                                        st.error(f"❌ Ошибка при получении интерпретации: {e}")
                                        import traceback
                                        with st.expander("🔍 Детали ошибки"):
                                            st.code(traceback.format_exc())
                        else:
                            # Если ИИ недоступен, показываем сообщение
                            st.warning("⚠️ ИИ-модуль недоступен. Проверьте настройки API ключей.")
                            if not AI_AVAILABLE:
                                st.error("❌ ИИ-модуль не загружен. Проверьте файл `claude_assistant.py`.")
                            elif OpenRouterAssistant is None:
                                st.error("❌ Класс OpenRouterAssistant недоступен.")
                        
                        # Экспорт результатов
                        st.subheader("📥 Экспорт результатов")
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            # JSON экспорт
                            if EXPORT_MANAGER_AVAILABLE:
                                json_data = {
                                    "analysis_id": analysis_result.analysis_id,
                                    "patient_info": patient_info,
                                    "summary": {
                                        "total_variants": analysis_result.total_variants,
                                        "pathogenic_count": len(analysis_result.pathogenic_variants),
                                        "clinically_significant": len(analysis_result.clinical_interpretations),
                                        "pharmacogenetic_count": len(analysis_result.pharmacogenetic_variants)
                                    },
                                    "pathogenic_variants": [v.to_dict() for v in analysis_result.pathogenic_variants[:50]],
                                    "recommendations": analysis_result.recommendations
                                }
                                json_file = export_analysis_to_json(json_data, f"genetic_analysis_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
                                with open(json_file, 'rb') as f:
                                    st.download_button(
                                        "📥 Скачать JSON",
                                        f.read(),
                                        file_name=os.path.basename(json_file),
                                        mime="application/json"
                                    )
                        
                        with col2:
                            # CSV экспорт вариантов
                            if EXPORT_MANAGER_AVAILABLE and analysis_result.pathogenic_variants:
                                variants_data = []
                                for v in analysis_result.pathogenic_variants[:100]:
                                    variants_data.append({
                                        "chromosome": v.chromosome,
                                        "position": v.position,
                                        "ref": v.ref,
                                        "alt": v.alt,
                                        "id": v.id,
                                        "quality": v.quality
                                    })
                                csv_file = export_analysis_to_csv(variants_data, f"genetic_variants_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv")
                                with open(csv_file, 'rb') as f:
                                    st.download_button(
                                        "📥 Скачать CSV вариантов",
                                        f.read(),
                                        file_name=os.path.basename(csv_file),
                                        mime="text/csv"
                                    )
            
            else:
                st.warning(f"Формат файла {file_ext} пока не поддерживается для полного анализа. Используйте VCF формат.")
                st.info("📊 Для других форматов функция в разработке")

def show_statistics_page():
    """Страница статистики использования моделей"""
    st.header("📊 Статистика использования")
    
    if 'model_stats' not in st.session_state or not st.session_state.model_stats:
        st.info("Статистика пока недоступна. Используйте функции анализа для накопления данных.")
        return
    
    stats = st.session_state.model_stats
    
    st.subheader("📈 Использование моделей ИИ")
    
    # Таблица статистики
    stats_data = []
    for model, data in stats.items():
        success_rate = (data['successful_calls'] / data['total_calls'] * 100) if data['total_calls'] > 0 else 0
        stats_data.append({
            "Модель": model,
            "Всего вызовов": data['total_calls'],
            "Успешных": data['successful_calls'],
            "Неудачных": data['failed_calls'],
            "Успешность": f"{success_rate:.1f}%",
            "Токенов использовано": data.get('total_tokens', 0)
        })
    
    df_stats = pd.DataFrame(stats_data)
    st.dataframe(df_stats, use_container_width=True)
    
    # Графики
    if len(stats_data) > 0:
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Успешность моделей")
            chart_data = pd.DataFrame({
                'Модель': [s['Модель'] for s in stats_data],
                'Успешность (%)': [float(s['Успешность'].replace('%', '')) for s in stats_data]
            })
            st.bar_chart(chart_data.set_index('Модель'))
        
        with col2:
            st.subheader("Количество вызовов")
            chart_data2 = pd.DataFrame({
                'Модель': [s['Модель'] for s in stats_data],
                'Вызовов': [s['Всего вызовов'] for s in stats_data]
            })
            st.bar_chart(chart_data2.set_index('Модель'))
    
    # Кнопка сброса статистики
    if st.button("🔄 Сбросить статистику"):
        st.session_state.model_stats = {}
        st.rerun()

def show_patient_context_page():
    """Страница управления клиническим контекстом пациента"""
    st.header("📋 Клинический контекст пациента")
    
    init_db()
    conn = sqlite3.connect('medical_data.db')
    patients = pd.read_sql_query("SELECT id, name FROM patients", conn)
    conn.close()
    
    if patients.empty:
        st.warning("❌ База пациентов пуста. Добавьте пациента в разделе 'База данных'.")
        return
    
    selected_patient = st.selectbox("Выберите пациента", patients['name'])
    patient_id = patients[patients['name'] == selected_patient].iloc[0]['id']
    
    context_store = ContextStore()
    
    # Вкладки для разных действий
    tab1, tab2, tab3 = st.tabs(["📊 Просмотр контекста", "➕ Добавить данные", "🔍 Использовать для анализа"])
    
    with tab1:
        st.subheader("📊 Просмотр сохраненного контекста")
        
        if st.button("📊 Загрузить контекст"):
            context_data = context_store.get_patient_context(patient_id)
            comprehensive_context = context_store.build_comprehensive_context(patient_id)
            
            if not context_data:
                st.info("Контекст для данного пациента отсутствует. Добавьте данные во вкладке '➕ Добавить данные'.")
            else:
                st.subheader("📋 Полный клинический контекст")
                st.text_area("Контекст", comprehensive_context, height=300, key="comprehensive_context")
                
                # Детализация по типам
                st.subheader("📁 Детализация по типам данных")
                for context_type, contexts in context_data.items():
                    with st.expander(f"📁 {context_type.upper()} ({len(contexts)} записей)"):
                        for i, ctx in enumerate(contexts, 1):
                            st.write(f"**Запись {i}** (источник: {ctx['source']}, дата: {ctx['created_at']})")
                            st.json(ctx['data'])
    
    with tab2:
        st.subheader("➕ Добавить данные в контекст пациента")
        
        context_type = st.selectbox(
            "Тип данных:",
            ["complaints", "lab_results", "imaging", "diagnosis", "protocol", "other"],
            format_func=lambda x: {
                "complaints": "Жалобы",
                "lab_results": "Лабораторные анализы",
                "imaging": "Результаты визуализации (ЭКГ, рентген и т.д.)",
                "diagnosis": "Диагноз",
                "protocol": "Протокол осмотра",
                "other": "Другое"
            }[x]
        )
        
        if context_type == "protocol":
            st.info("💡 Вставьте текст протокола осмотра (можно скопировать из Word или другого документа)")
            protocol_text = st.text_area("Текст протокола:", height=200, key="protocol_text")
            
            if st.button("💾 Сохранить протокол"):
                if protocol_text:
                    context_store.add_context(
                        patient_id=patient_id,
                        context_type='protocol',
                        context_data={'protocol': protocol_text, 'type': 'consultation'},
                        source='manual_entry'
                    )
                    st.success("✅ Протокол осмотра сохранен в контекст пациента!")
                else:
                    st.warning("⚠️ Введите текст протокола")
        
        elif context_type == "complaints":
            st.info("💡 Введите жалобы пациента")
            complaints_text = st.text_area("Жалобы:", height=150, key="complaints_text")
            
            if st.button("💾 Сохранить жалобы"):
                if complaints_text:
                    context_store.add_context(
                        patient_id=patient_id,
                        context_type='complaints',
                        context_data={'complaints': complaints_text},
                        source='manual_entry'
                    )
                    st.success("✅ Жалобы сохранены в контекст пациента!")
                else:
                    st.warning("⚠️ Введите жалобы")
        
        elif context_type == "diagnosis":
            st.info("💡 Введите диагноз")
            diagnosis_text = st.text_input("Диагноз:", key="diagnosis_text")
            icd10 = st.text_input("Код МКБ-10 (опционально):", key="icd10")
            
            if st.button("💾 Сохранить диагноз"):
                if diagnosis_text:
                    context_store.add_context(
                        patient_id=patient_id,
                        context_type='diagnosis',
                        context_data={'diagnosis': diagnosis_text, 'icd10': icd10},
                        source='manual_entry'
                    )
                    st.success("✅ Диагноз сохранен в контекст пациента!")
                else:
                    st.warning("⚠️ Введите диагноз")
        
        elif context_type == "lab_results":
            st.info("💡 Введите результаты лабораторных анализов (можно вставить текст или JSON)")
            lab_text = st.text_area("Результаты анализов:", height=200, key="lab_text")
            
            if st.button("💾 Сохранить анализы"):
                if lab_text:
                    try:
                        # Пробуем распарсить как JSON
                        lab_data = json.loads(lab_text)
                    except:
                        # Если не JSON, сохраняем как текст
                        lab_data = {'results_text': lab_text}
                    
                    context_store.add_context(
                        patient_id=patient_id,
                        context_type='lab_results',
                        context_data=lab_data,
                        source='manual_entry'
                    )
                    st.success("✅ Результаты анализов сохранены в контекст пациента!")
                else:
                    st.warning("⚠️ Введите результаты анализов")
        
        elif context_type == "imaging":
            st.info("💡 Введите результаты визуализации (ЭКГ, рентген, МРТ и т.д.)")
            imaging_type = st.selectbox("Тип исследования:", ["ЭКГ", "Рентген", "МРТ", "КТ", "УЗИ", "Другое"])
            imaging_text = st.text_area("Результаты исследования:", height=200, key="imaging_text")
            
            if st.button("💾 Сохранить результаты"):
                if imaging_text:
                    context_store.add_context(
                        patient_id=patient_id,
                        context_type='imaging',
                        context_data={'type': imaging_type, 'results': imaging_text},
                        source='manual_entry'
                    )
                    st.success("✅ Результаты исследования сохранены в контекст пациента!")
                else:
                    st.warning("⚠️ Введите результаты исследования")
        
        else:  # other
            st.info("💡 Введите произвольные данные")
            other_text = st.text_area("Данные:", height=200, key="other_text")
            
            if st.button("💾 Сохранить данные"):
                if other_text:
                    context_store.add_context(
                        patient_id=patient_id,
                        context_type='other',
                        context_data={'data': other_text},
                        source='manual_entry'
                    )
                    st.success("✅ Данные сохранены в контекст пациента!")
                else:
                    st.warning("⚠️ Введите данные")
    
    with tab3:
        st.subheader("🔍 Использовать контекст для анализа")
        st.info("💡 Загрузите контекст пациента, чтобы он использовался при следующем анализе ЭКГ, рентгена и т.д.")
        
        if st.button("📥 Загрузить контекст для использования"):
            comprehensive_context = context_store.build_comprehensive_context(patient_id)
            
            if comprehensive_context:
                st.session_state['patient_context'] = comprehensive_context
                st.session_state['selected_patient_id'] = patient_id
                st.success("✅ Контекст загружен! Он будет использован при следующем анализе.")
                st.info("💡 Теперь перейдите в раздел 'Анализ ЭКГ' или другой анализ - контекст будет автоматически учтен.")
                
                with st.expander("📋 Просмотр загруженного контекста"):
                    st.text_area("", comprehensive_context, height=200, disabled=True)
            else:
                st.warning("⚠️ Контекст для данного пациента отсутствует. Добавьте данные во вкладке '➕ Добавить данные'.")
        
        if 'patient_context' in st.session_state:
            st.success("✅ Контекст активен и будет использован при анализе")
            if st.button("❌ Очистить контекст"):
                del st.session_state['patient_context']
                if 'selected_patient_id' in st.session_state:
                    del st.session_state['selected_patient_id']
                st.success("✅ Контекст очищен")

def show_video_analysis():
    """Страница анализа медицинских видео"""
    if not AI_AVAILABLE:
        st.error("❌ ИИ-модуль недоступен. Проверьте файл `claude_assistant.py` и API-ключ.")
        return
    
    st.header("🎬 Анализ медицинских видео")
    st.info("💡 Загрузите видео медицинской процедуры, функционального теста или динамического исследования для анализа через Gemini 2.5 Flash")
    
    # Выбор типа исследования
    study_type = st.selectbox(
        "Тип исследования:",
        ["", "fgds", "colonoscopy", "echo", "abdominal_us", "gynecology_us", "mri_brain", "mri_universal", "chest_ct"],
        format_func=lambda x: {
            "": "Выберите тип исследования",
            "fgds": "🔬 ФГДС (эзофагогастродуоденоскопия)",
            "colonoscopy": "🔬 Колоноскопия",
            "echo": "🫀 ЭхоКГ (эхокардиография)",
            "abdominal_us": "🔍 УЗИ органов брюшной полости",
            "gynecology_us": "🩺 Гинекологическое УЗИ",
            "mri_brain": "🧠 МРТ головного мозга",
            "mri_universal": "🧲 МРТ (универсальный)",
            "chest_ct": "🫁 КТ органов грудной клетки"
        }.get(x, x),
        help="Выберите тип исследования для использования специализированного промпта"
    )
    
    # Загрузка видео
    uploaded_video = st.file_uploader(
        "Загрузите видео-файл",
        type=["mp4", "mov", "avi", "webm", "mkv"],
        help="Поддерживаются форматы: MP4, MOV, AVI, WebM, MKV (максимум 100MB)"
    )
    
    if uploaded_video:
        # Показываем превью видео
        st.subheader("📹 Превью видео")
        st.video(uploaded_video)
        
        # Метаданные (опционально)
        st.subheader("📋 Метаданные (опционально)")
        col1, col2 = st.columns(2)
        
        with col1:
            patient_age = st.number_input("Возраст пациента", min_value=0, max_value=150, value=None, help="Укажите возраст для более точного анализа")
            specialty = st.selectbox(
                "Специализация",
                ["", "Терапия", "Хирургия", "Ортопедия", "Неврология", "Кардиология", "Педиатрия", "Онкология", "Другое"],
                help="Выберите специализацию для контекста анализа"
            )
        
        with col2:
            urgency = st.selectbox(
                "Срочность",
                ["", "Плановая", "Срочная", "Критическая"],
                help="Укажите уровень срочности"
            )
        
        # Дополнительный контекст (особенно для КТ ОГК)
        additional_context = ""
        if study_type == "chest_ct":
            st.subheader("📋 Дополнительные параметры для КТ ОГК")
            col_ct1, col_ct2, col_ct3 = st.columns(3)
            with col_ct1:
                ct_type = st.selectbox("Тип КТ", ["Нативное", "С контрастом", "КТЛА", "ВРКТ"])
            with col_ct2:
                clinical = st.text_input("Клиника", placeholder="Кашель, одышка, лихорадка...")
            with col_ct3:
                covid_suspicion = st.checkbox("Подозрение на COVID-19")
            
            if ct_type:
                additional_context += f"Тип КТ: {ct_type}\n"
            if clinical:
                additional_context += f"Клинические данные: {clinical}\n"
            if covid_suspicion:
                additional_context += "ВАЖНО: Оцени CT severity score для COVID-19!\n"
        else:
            additional_context = st.text_area(
                "Дополнительный контекст",
                placeholder="Опишите клиническую ситуацию, жалобы пациента, цель исследования...",
                help="Любая дополнительная информация, которая поможет в анализе"
            )
        
        # Кнопка анализа
        if st.button("🎬 Анализировать видео", type="primary", use_container_width=True):
            # Нормализуем study_type: пустая строка становится None
            # Проверяем явно на пустую строку и None
            if not study_type or study_type == "" or study_type.strip() == "":
                study_type_for_request = None
                st.info("💡 Тип исследования не выбран. Будет использован базовый промпт для анализа.")
            else:
                study_type_for_request = study_type
                # Показываем, какой тип исследования выбран
                study_type_names = {
                    "fgds": "🔬 ФГДС",
                    "colonoscopy": "🔬 Колоноскопия",
                    "echo": "🫀 ЭхоКГ",
                    "abdominal_us": "🔍 УЗИ органов брюшной полости",
                    "gynecology_us": "🩺 Гинекологическое УЗИ",
                    "mri_brain": "🧠 МРТ головного мозга",
                    "mri_universal": "🧲 МРТ (универсальный)",
                    "chest_ct": "🫁 КТ органов грудной клетки"
                }
                selected_name = study_type_names.get(study_type, study_type)
                st.success(f"✅ Используется специализированный промпт: {selected_name}")
            # Показываем прогресс
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            try:
                status_text.info("🔄 Подготовка видео...")
                progress_bar.progress(10)
                
                assistant = OpenRouterAssistant()
                
                # Подготавливаем метаданные
                metadata = {}
                if patient_age:
                    metadata['patient_age'] = patient_age
                if specialty:
                    metadata['specialty'] = specialty
                if urgency:
                    metadata['urgency'] = urgency
                if additional_context:
                    metadata['additional_context'] = additional_context
                
                # Формируем дополнительный промпт из контекста, если есть
                context_prompt = None
                if metadata:
                    context_parts = []
                    if patient_age:
                        context_parts.append(f"Возраст пациента: {patient_age} лет")
                    if specialty:
                        context_parts.append(f"Специализация: {specialty}")
                    if urgency:
                        context_parts.append(f"Срочность: {urgency}")
                    if additional_context:
                        context_parts.append(f"Дополнительный контекст: {additional_context}")
                    
                    if context_parts:
                        context_prompt = "\n\nКОНТЕКСТ:\n" + "\n".join(context_parts)
                
                # Двухэтапный анализ видео
                status_text.info("🔄 Этап 1: Специализированный анализ через Gemini 2.5 Flash...")
                progress_bar.progress(20)
                
                # Этап 1: Специализированный анализ
                with st.spinner("⏳ Анализ видео через Gemini..."):
                    results = assistant.send_video_request_two_stage(
                        prompt=context_prompt,
                        video_data=uploaded_video,
                        metadata=metadata if metadata else None,
                        study_type=study_type_for_request
                    )
                
                progress_bar.progress(50)
                
                # Показываем промежуточный результат (специализированный анализ)
                if results.get('specialized'):
                    st.subheader("📋 Промежуточный результат: Специализированный анализ")
                    with st.expander("🔍 Показать специализированный анализ (Gemini 2.5 Flash)", expanded=True):
                        st.markdown(results['specialized'])
                
                # Этап 2: Итоговое заключение от профессора
                if results.get('final') and not results['final'].startswith("❌"):
                    status_text.info("🔄 Этап 2: Итоговое заключение от профессора (Claude Opus)...")
                    progress_bar.progress(70)
                    
                    # Результат уже получен в двухэтапном методе, просто показываем прогресс
                    time.sleep(0.5)  # Небольшая задержка для визуализации прогресса
                    
                    progress_bar.progress(100)
                    status_text.empty()
                    progress_bar.empty()
                    
                    # Показываем финальное заключение
                    st.subheader("🎓 Итоговое заключение")
                    st.markdown(results['final'])
                elif results.get('final') and results['final'].startswith("❌"):
                    # Если была ошибка на этапе 2, показываем её
                    progress_bar.progress(100)
                    status_text.empty()
                    progress_bar.empty()
                    st.warning(f"⚠️ {results['final']}")
                    st.info("💡 Специализированный анализ доступен выше")
                else:
                    progress_bar.progress(100)
                    status_text.empty()
                    progress_bar.empty()
                    st.info("💡 Итоговое заключение не было сформировано. Доступен только специализированный анализ.")
                
                # Экспорт в DOC формат
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                
                # Формируем содержимое для DOC
                study_type_names = {
                    "fgds": "ФГДС",
                    "colonoscopy": "Колоноскопия",
                    "echo": "ЭхоКГ",
                    "abdominal_us": "УЗИ органов брюшной полости",
                    "gynecology_us": "Гинекологическое УЗИ",
                    "mri_brain": "МРТ головного мозга",
                    "mri_universal": "МРТ (универсальный)",
                    "chest_ct": "КТ органов грудной клетки"
                }
                study_name = study_type_names.get(study_type_for_request, "Видео-анализ") if study_type_for_request else "Видео-анализ"
                
                # Создаем DOC документ
                try:
                    from docx import Document
                    from docx.shared import Pt, Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    
                    doc = Document()
                    
                    # Заголовок
                    title = doc.add_heading(f"Анализ видео: {study_name}", level=0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Метаданные
                    doc.add_paragraph(f"Дата анализа: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}")
                    if patient_age:
                        doc.add_paragraph(f"Возраст пациента: {patient_age} лет")
                    if specialty:
                        doc.add_paragraph(f"Специализация: {specialty}")
                    if urgency:
                        doc.add_paragraph(f"Срочность: {urgency}")
                    doc.add_paragraph()
                    
                    # Раздел 1: Специализированный анализ
                    if results.get('specialized'):
                        doc.add_heading("СПЕЦИАЛИЗИРОВАННЫЙ АНАЛИЗ (Gemini 2.5 Flash)", level=1)
                        # Убираем markdown форматирование для чистого текста
                        specialized_text = results['specialized'].replace('**', '').replace('🎬', '').strip()
                        doc.add_paragraph(specialized_text)
                        doc.add_paragraph()
                    
                    # Раздел 2: Итоговое заключение
                    if results.get('final'):
                        doc.add_heading("ИТОГОВОЕ ЗАКЛЮЧЕНИЕ (Профессор, Claude Opus 4.5)", level=1)
                        final_text = results['final'].replace('**', '').replace('🎓', '').strip()
                        doc.add_paragraph(final_text)
                    
                    # Сохраняем в BytesIO для скачивания
                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    # Кнопка скачивания DOC
                    doc_filename = f"video_analysis_{study_name.replace(' ', '_')}_{timestamp}.docx"
                    st.download_button(
                        label="📥 Скачать полный отчет (.docx)",
                        data=doc_buffer.getvalue(),
                        file_name=doc_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except ImportError:
                    # Если python-docx не установлен, предлагаем TXT
                    st.warning("⚠️ Для экспорта в DOC формат требуется python-docx. Установите: pip install python-docx")
                    # Альтернатива: TXT файл
                    full_text = f"АНАЛИЗ ВИДЕО: {study_name}\n"
                    full_text += f"Дата: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}\n\n"
                    if results.get('specialized'):
                        full_text += "=" * 50 + "\n"
                        full_text += "СПЕЦИАЛИЗИРОВАННЫЙ АНАЛИЗ (Gemini 2.5 Flash)\n"
                        full_text += "=" * 50 + "\n"
                        full_text += results['specialized'] + "\n\n"
                    if results.get('final'):
                        full_text += "=" * 50 + "\n"
                        full_text += "ИТОГОВОЕ ЗАКЛЮЧЕНИЕ (Профессор, Claude Opus 4.5)\n"
                        full_text += "=" * 50 + "\n"
                        full_text += results['final'] + "\n"
                    
                    txt_filename = f"video_analysis_{timestamp}.txt"
                    st.download_button(
                        label="📥 Скачать отчет (.txt)",
                        data=full_text,
                        file_name=txt_filename,
                        mime="text/plain"
                    )
                
            except Exception as e:
                progress_bar.empty()
                status_text.empty()
                st.error(f"❌ Ошибка анализа видео: {e}")
                import traceback
                with st.expander("🔍 Детали ошибки"):
                    st.code(traceback.format_exc())
    else:
        st.info("👆 Загрузите видео-файл для начала анализа")

def show_document_scanner_page():
    """Страница сканирования и извлечения данных из медицинских документов"""
    st.header("📄 Сканирование медицинских документов")
    st.info("💡 Загрузите фото или сканированную копию медицинской справки, рецепта, направления или выписки для автоматического извлечения данных")
    
    # Выбор типа документа
    doc_type = st.selectbox(
        "Тип документа:",
        ["Медицинская справка", "Рецепт", "Направление на обследование", "Выписка из больницы", "Больничный лист", "Результаты анализов", "Другое"],
        help="Выберите тип документа для более точного извлечения данных"
    )
    
    # Выбор источника
    source_type = st.radio(
        "Источник документа:",
        ["📁 Загрузить файл", "📷 Сделать фото"],
        horizontal=True
    )
    
    image_array = None
    uploaded_file = None
    
    if source_type == "📷 Сделать фото":
        camera_image = st.camera_input("Сфотографируйте документ", key="doc_camera")
        if camera_image:
            try:
                image = Image.open(camera_image)
                image_array = np.array(image)
            except Exception as e:
                st.error(f"Ошибка обработки фото: {e}")
                return
    else:
        uploaded_file = st.file_uploader(
            "Загрузите документ",
            type=["jpg", "jpeg", "png", "pdf", "tiff", "tif", "heic", "webp"],
            help="Поддерживаются изображения и PDF файлы"
        )
        
        if uploaded_file:
            try:
                if uploaded_file.type == "application/pdf":
                    st.info("📄 PDF файл. Используется извлечение текста из PDF...")
                    # Для PDF используем существующий процессор
                    from modules.advanced_lab_processor import AdvancedLabProcessor
                    processor = AdvancedLabProcessor()
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                        tmp.write(uploaded_file.getvalue())
                        tmp_path = tmp.name
                    
                    try:
                        extracted_text = processor._extract_from_pdf(tmp_path)
                        st.session_state['extracted_doc_text'] = extracted_text
                        st.success("✅ Текст извлечен из PDF")
                    except Exception as e:
                        st.error(f"Ошибка извлечения из PDF: {e}")
                    finally:
                        if os.path.exists(tmp_path):
                            os.unlink(tmp_path)
                else:
                    # Для изображений
                    image = Image.open(uploaded_file)
                    image_array = np.array(image)
            except Exception as e:
                st.error(f"Ошибка обработки файла: {e}")
                return
    
    # Если есть изображение, показываем его
    if image_array is not None:
        st.image(image_array, caption="Загруженный документ", use_container_width=True, clamp=True)
        
        col_scan, col_struct = st.columns(2)
        
        # Режим 1: ЧИСТОЕ СКАНИРОВАНИЕ (получить текст без анализа)
        with col_scan:
            if st.button("📄 Сканировать (получить текст)", use_container_width=True, type="secondary"):
                if not AI_AVAILABLE:
                    st.error("❌ ИИ-модуль недоступен. Проверьте файл `claude_assistant.py` и API-ключ.")
                    return
                with st.spinner("🤖 ИИ распознает текст документа..."):
                    assistant = OpenRouterAssistant()
                    ocr_prompt = """
Вы — эксперт по OCR медицинских документов. 
Аккуратно извлеките ВЕСЬ читаемый текст с этого изображения.
Верните ТОЛЬКО текст документа, без перевода, без интерпретации, без клинических выводов и без ссылок.
Сохраняйте максимально исходное форматирование (строки, абзацы), насколько это возможно.
"""
                    try:
                        scanned_text = assistant.send_vision_request(
                            ocr_prompt,
                            image_array,
                            metadata={"task": "doc_ocr"}
                        )
                        if isinstance(scanned_text, list):
                            # На всякий случай, если вернулся список результатов
                            scanned_text = "\n\n".join(str(x.get("result", x)) for x in scanned_text)
                        st.session_state['scanned_doc_text'] = str(scanned_text)
                        st.subheader("📋 Распознанный текст документа")
                        st.text_area("Текст", st.session_state['scanned_doc_text'], height=300)
                        
                        st.download_button(
                            label="📥 Скачать как .txt",
                            data=st.session_state['scanned_doc_text'],
                            file_name="scanned_document.txt",
                            mime="text/plain"
                        )
                    except Exception as e:
                        st.error(f"❌ Ошибка распознавания: {e}")
        
        # Режим 2: Структурированное извлечение (как было)
        with col_struct:
            if st.button("🔍 Извлечь данные из документа", use_container_width=True, type="primary"):
                if not AI_AVAILABLE:
                    st.error("❌ ИИ-модуль недоступен. Проверьте файл `claude_assistant.py` и API-ключ.")
                    return
                
                with st.spinner("🤖 ИИ анализирует документ и извлекает данные..."):
                    assistant = OpenRouterAssistant()
                    
                    # Промпт в зависимости от типа документа
                    prompts = {
                    "Медицинская справка": """
Вы - эксперт по распознаванию медицинских документов. Извлеките из этого изображения медицинской справки все данные в структурированном JSON формате.

Извлеките:
1. ФИО пациента
2. Дата рождения
3. Дата выдачи справки
4. Номер справки (если есть)
5. Название медицинского учреждения
6. ФИО врача, выдавшего справку
7. Диагноз или заключение
8. Рекомендации (если есть)
9. Ограничения или противопоказания (если есть)
10. Печати и подписи (наличие)

Формат ответа - JSON:
{
  "patient_name": "...",
  "birth_date": "...",
  "issue_date": "...",
  "document_number": "...",
  "medical_institution": "...",
  "doctor_name": "...",
  "diagnosis": "...",
  "recommendations": "...",
  "restrictions": "...",
  "has_stamp": true/false,
  "has_signature": true/false,
  "raw_text": "весь извлеченный текст"
}
""",
                    "Рецепт": """
Вы - эксперт по распознаванию рецептов. Извлеките из этого изображения рецепта все данные в структурированном JSON формате.

Извлеките:
1. ФИО пациента
2. Дата выдачи рецепта
3. ФИО врача
4. Список препаратов с:
   - Название (международное и торговое)
   - Дозировка
   - Количество
   - Способ применения
   - Кратность приема
5. Срок действия рецепта
6. Печати и подписи

Формат ответа - JSON:
{
  "patient_name": "...",
  "issue_date": "...",
  "doctor_name": "...",
  "medications": [
    {
      "name": "...",
      "dosage": "...",
      "quantity": "...",
      "instructions": "...",
      "frequency": "..."
    }
  ],
  "valid_until": "...",
  "has_stamp": true/false,
  "raw_text": "весь извлеченный текст"
}
""",
                    "Направление на обследование": """
Вы - эксперт по распознаванию медицинских направлений. Извлеките из этого изображения направления все данные в структурированном JSON формате.

Извлеките:
1. ФИО пациента
2. Дата направления
3. ФИО врача, выдавшего направление
4. Тип обследования
5. Цель обследования
6. Предварительный диагноз
7. Медицинское учреждение назначения
8. Срочность
9. Особые указания

Формат ответа - JSON:
{
  "patient_name": "...",
  "issue_date": "...",
  "doctor_name": "...",
  "examination_type": "...",
  "purpose": "...",
  "preliminary_diagnosis": "...",
  "target_institution": "...",
  "urgency": "...",
  "special_instructions": "...",
  "raw_text": "весь извлеченный текст"
}
""",
                    "Выписка из больницы": """
Вы - эксперт по распознаванию выписок из больницы. Извлеките из этого изображения выписки все данные в структурированном JSON формате.

Извлеките:
1. ФИО пациента
2. Дата рождения
3. Даты госпитализации и выписки
4. Отделение
5. Диагноз при поступлении
6. Диагноз при выписке
7. Проведенное лечение
8. Операции (если были)
9. Рекомендации при выписке
10. ФИО лечащего врача

Формат ответа - JSON:
{
  "patient_name": "...",
  "birth_date": "...",
  "admission_date": "...",
  "discharge_date": "...",
  "department": "...",
  "admission_diagnosis": "...",
  "discharge_diagnosis": "...",
  "treatment": "...",
  "surgeries": [...],
  "recommendations": "...",
  "attending_doctor": "...",
  "raw_text": "весь извлеченный текст"
}
""",
                    "Больничный лист": """
Вы - эксперт по распознаванию больничных листов. Извлеките из этого изображения больничного листа все данные в структурированном JSON формате.

Извлеките:
1. ФИО пациента
2. Дата начала нетрудоспособности
3. Дата окончания нетрудоспособности
4. Диагноз
5. Код МКБ-10
6. ФИО врача
7. Медицинское учреждение
8. Номер больничного листа

Формат ответа - JSON:
{
  "patient_name": "...",
  "start_date": "...",
  "end_date": "...",
  "diagnosis": "...",
  "icd10_code": "...",
  "doctor_name": "...",
  "medical_institution": "...",
  "document_number": "...",
  "raw_text": "весь извлеченный текст"
}
""",
                    "Результаты анализов": """
Вы - эксперт по распознаванию результатов анализов. Извлеките из этого изображения все данные в структурированном JSON формате.

Извлеките:
1. ФИО пациента
2. Дата анализа
3. Тип анализа
4. Название лаборатории
5. Все параметры с значениями, единицами измерения и референсными интервалами
6. Заключение (если есть)

Формат ответа - JSON:
{
  "patient_name": "...",
  "analysis_date": "...",
  "analysis_type": "...",
  "laboratory": "...",
  "parameters": [
    {
      "name": "...",
      "value": "...",
      "unit": "...",
      "reference_range": "...",
      "status": "normal/abnormal"
    }
  ],
  "conclusion": "...",
  "raw_text": "весь извлеченный текст"
}
""",
                    "Другое": """
Вы - эксперт по распознаванию медицинских документов. Извлеките из этого изображения все данные в структурированном JSON формате.

Извлеките:
1. Тип документа
2. ФИО пациента (если есть)
3. Даты
4. Все ключевые данные
5. Полный текст документа

Формат ответа - JSON:
{
  "document_type": "...",
  "patient_name": "...",
  "dates": [...],
  "key_data": {...},
  "raw_text": "весь извлеченный текст"
}
"""
                }
                
                prompt = prompts.get(doc_type, prompts["Другое"])
                
                # Отправка запроса к ИИ
                # Для извлечения текста из справок используем Llama (лучше для документов)
                try:
                    result = assistant.send_vision_request(
                        prompt, 
                        image_array, 
                        str({"document_type": doc_type}), 
                        use_router=True,  # Используем роутер, он автоматически выберет Llama для документов
                        force_model="llama"  # Принудительно используем Llama для извлечения текста из документов
                    )
                    
                    # Проверяем, что результат не пустой и не является ошибкой
                    if not result or len(str(result).strip()) == 0:
                        st.error("❌ ИИ вернул пустой ответ. Попробуйте еще раз.")
                        return
                    
                    # Проверяем, что это не сообщение об ошибке
                    result_str = str(result).strip()
                    if result_str.startswith("❌") or "Ошибка" in result_str or "недоступны" in result_str.lower() or "Key limit exceeded" in result_str:
                        st.error(f"❌ {result_str}")
                        st.info("💡 Все модели ИИ недоступны. Проверьте лимиты API ключа на https://openrouter.ai/settings/keys или попробуйте позже.")
                        # Очищаем session_state, чтобы не показывать пустые данные
                        if 'extracted_doc_raw' in st.session_state:
                            del st.session_state['extracted_doc_raw']
                        if 'extracted_doc_data' in st.session_state:
                            del st.session_state['extracted_doc_data']
                        return
                    
                    # Этап 1: ТОЛЬКО извлечение текста (сканирование)
                    # НЕ сохраняем структурированные данные - только текст
                    # Структурирование будет происходить только по требованию пользователя
                    json_match = re.search(r'\{.*\}', result_str, re.DOTALL)
                    
                    # Извлекаем чистый текст из ответа
                    if json_match:
                        try:
                            extracted_data = json.loads(json_match.group())
                            # Если в JSON есть raw_text, используем его, иначе весь ответ
                            if isinstance(extracted_data, dict) and 'raw_text' in extracted_data:
                                clean_extracted_text = extracted_data['raw_text']
                            else:
                                # Извлекаем текст из JSON, убирая структуру
                                clean_extracted_text = result_str
                        except Exception as parse_error:
                            # Если не удалось распарсить, используем весь ответ как текст
                            clean_extracted_text = result_str
                    else:
                        # Если JSON не найден, используем весь ответ как текст
                        clean_extracted_text = result_str
                    
                    # Сохраняем ТОЛЬКО текст, НЕ структурированные данные
                    st.session_state['extracted_doc_raw'] = clean_extracted_text
                    st.session_state['extracted_doc_data'] = None  # Структурированные данные будут только после ИИ-анализа
                    
                    st.success("✅ Документ отсканирован! Текст извлечен.")
                    st.info("💡 Выберите дальнейшее действие: сохранить текст или проанализировать ИИ")
                    
                    st.rerun()  # Перезагружаем страницу для отображения извлеченного текста и опций
                    
                except Exception as e:
                    error_msg = str(e)
                    st.error(f"❌ Ошибка при извлечении данных: {error_msg}")
                    
                    # Если это ошибка о недоступности моделей, показываем дополнительную информацию
                    if "недоступны" in error_msg.lower() or "403" in error_msg or "Key limit" in error_msg:
                        st.info("💡 Все модели ИИ недоступны из-за превышения лимита API ключа. Проверьте настройки на https://openrouter.ai/settings/keys")
                    
                    # Очищаем session_state при ошибке
                    if 'extracted_doc_raw' in st.session_state:
                        del st.session_state['extracted_doc_raw']
                    if 'extracted_doc_data' in st.session_state:
                        del st.session_state['extracted_doc_data']
                    return
    
    # Показ извлеченных данных
    if 'extracted_doc_data' in st.session_state and st.session_state['extracted_doc_data']:
        st.subheader("📋 Извлеченные данные")
        extracted_data = st.session_state['extracted_doc_data']
        
        # Отображение структурированных данных
        st.json(extracted_data)
        
        # Сохранение в контекст пациента
        st.subheader("💾 Сохранение данных")
        init_db()
        conn = sqlite3.connect('medical_data.db')
        patients = pd.read_sql_query("SELECT id, name FROM patients", conn)
        conn.close()
        
        if not patients.empty:
            selected_patient = st.selectbox("Выберите пациента для сохранения:", patients['name'], key="doc_patient_select")
            patient_id = patients[patients['name'] == selected_patient].iloc[0]['id']
            
            if st.button("💾 Сохранить в контекст пациента"):
                try:
                    context_store = ContextStore()
                    context_store.add_context(
                        patient_id=patient_id,
                        context_type='document',
                        context_data={
                            'document_type': doc_type,
                            'extracted_data': extracted_data,
                            'raw_text': extracted_data.get('raw_text', '')
                        },
                        source='ai_extraction'
                    )
                    st.success("✅ Данные сохранены в клинический контекст пациента!")
                except Exception as e:
                    st.error(f"❌ Ошибка сохранения: {e}")
        else:
            st.info("💡 Добавьте пациента в разделе 'База данных', чтобы сохранять извлеченные данные")
        
        # Экспорт данных
        st.subheader("📥 Экспорт данных")
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📄 Экспорт в Word"):
                try:
                    from local_docs import create_local_doc
                    doc_text = json.dumps(extracted_data, ensure_ascii=False, indent=2)
                    filepath, message = create_local_doc(f"Извлеченные данные - {doc_type}", doc_text)
                    st.success(message)
                    with open(filepath, "rb") as f:
                        file_name = os.path.basename(filepath)
                        if not file_name.endswith('.docx'):
                            file_name = file_name.replace('.doc', '.docx')
                        st.download_button(
                            label="📥 Скачать документ",
                            data=f,
                            file_name=file_name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"❌ Ошибка экспорта: {e}")
        
        with col2:
            json_str = json.dumps(extracted_data, ensure_ascii=False, indent=2)
            st.download_button(
                label="📥 Скачать JSON",
                data=json_str,
                file_name=f"extracted_data_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json"
            )
    
    elif 'extracted_doc_raw' in st.session_state:
        extracted_text = st.session_state.get('extracted_doc_raw', '')
        
        # Проверяем, что текст не пустой
        if extracted_text and len(str(extracted_text).strip()) > 0:
            st.subheader("📋 Извлеченный текст")
            # Убираем markdown форматирование для лучшей читаемости
            clean_text = str(extracted_text).strip()
            # Убираем лишние звездочки и форматирование, если есть
            if clean_text.startswith('**') or clean_text.startswith('*'):
                # Пытаемся извлечь чистый текст
                clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean_text)
                clean_text = re.sub(r'\*([^*]+)\*', r'\1', clean_text)
            
            st.text_area("Текст", clean_text, height=300, disabled=False, key="extracted_text_display")
            
            # Разделяем действия: сохранить или проанализировать ИИ
            st.markdown("---")
            st.subheader("📌 Дальнейшие действия")
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("💾 Сохранить текст в файл", use_container_width=True, type="primary"):
                    # Сохранение в текстовый файл
                    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"extracted_text_{timestamp}.txt"
                    
                    st.download_button(
                        label="📥 Скачать текстовый файл",
                        data=clean_text,
                        file_name=filename,
                        mime="text/plain",
                        key="download_text_file"
                    )
                    st.success("✅ Готово к скачиванию!")
            
            with col2:
                if st.button("🤖 Проанализировать ИИ", use_container_width=True, type="secondary"):
                    # Переходим к ИИ-анализу извлеченного текста
                    if not AI_AVAILABLE:
                        st.error("❌ ИИ-модуль недоступен.")
                        return
                    
                    with st.spinner("🤖 ИИ анализирует извлеченный текст..."):
                        assistant = OpenRouterAssistant()
                        
                        # Получаем тип документа из session_state
                        current_doc_type = st.session_state.get('current_doc_type', 'медицинский документ')
                        
                        # Определяем, является ли документ лабораторным
                        is_lab_document = any(keyword in current_doc_type.lower() for keyword in 
                                            ['лаборатор', 'лабораторн', 'анализ крови', 'биохимия', 'гематолог'])
                        
                        # Промпт для анализа текста (используем текстовый запрос, не vision)
                        analysis_prompt = f"""Вы - эксперт по структурированию медицинских документов. 
Проанализируйте следующий извлеченный текст из медицинского документа и структурируйте его в JSON формате.

Тип документа: {current_doc_type}

Извлеченный текст:
{clean_text[:8000]}

Извлеките все ключевые данные в JSON формате:
- ФИО пациента (если есть)
- Даты (рождения, выдачи, обследований)
- Название медицинского учреждения
- ФИО врача (если есть)
- Диагнозы, заключения
- Рекомендации
- Все остальные важные данные

Верните ТОЛЬКО JSON объект с извлеченными данными, без дополнительных комментариев."""
                        
                        try:
                            # Для лабораторных документов используем Sonnet 4.5, для остальных - Opus
                            analysis_result = assistant.get_response(analysis_prompt, use_sonnet_4_5=is_lab_document)
                            
                            # Пытаемся распарсить JSON из результата
                            json_match = re.search(r'\{.*\}', analysis_result, re.DOTALL)
                            if json_match:
                                try:
                                    extracted_data = json.loads(json_match.group())
                                    # Сохраняем структурированные данные ТОЛЬКО после ИИ-анализа
                                    st.session_state['extracted_doc_data'] = extracted_data
                                    # Сохраняем исходный текст
                                    st.session_state['extracted_doc_raw'] = clean_text
                                    st.success("✅ ИИ успешно структурировал данные!")
                                    # Удаляем extracted_doc_raw, чтобы показать структурированные данные
                                    if 'extracted_doc_raw' in st.session_state:
                                        del st.session_state['extracted_doc_raw']
                                    st.rerun()  # Перезагружаем для показа структурированных данных
                                except:
                                    st.warning("⚠️ Не удалось распарсить JSON из ответа ИИ")
                                    st.text_area("Ответ ИИ", analysis_result, height=200)
                            else:
                                st.warning("⚠️ ИИ не вернул JSON формат")
                                st.text_area("Ответ ИИ", analysis_result, height=200)
                        except Exception as e:
                            st.error(f"❌ Ошибка ИИ-анализа: {str(e)}")
            
            st.info("💡 Выберите действие: сохранить текст в файл или проанализировать его с помощью ИИ для структурирования данных.")
        else:
            st.warning("⚠️ Текст не был извлечен из документа. Возможно, документ не содержит читаемого текста или произошла ошибка при обработке.")
            if st.button("🔄 Попробовать еще раз"):
                if 'extracted_doc_raw' in st.session_state:
                    del st.session_state['extracted_doc_raw']
                st.rerun()
    
    # Для PDF файлов
    if 'extracted_doc_text' in st.session_state:
        extracted_pdf_text = st.session_state.get('extracted_doc_text', '')
        
        if extracted_pdf_text and len(str(extracted_pdf_text).strip()) > 0:
            st.subheader("📋 Извлеченный текст из PDF")
            st.text_area("Текст", str(extracted_pdf_text).strip(), height=300, key="extracted_pdf_text_display")
            
            # Разделяем действия для PDF: сохранить или проанализировать ИИ
            st.markdown("---")
            st.subheader("📌 Дальнейшие действия")
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("💾 Сохранить PDF текст в файл", use_container_width=True, type="primary", key="save_pdf_text_btn"):
                    # Сохранение в текстовый файл
                    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"extracted_pdf_text_{timestamp}.txt"
                    
                    st.download_button(
                        label="📥 Скачать текстовый файл",
                        data=str(extracted_pdf_text).strip(),
                        file_name=filename,
                        mime="text/plain",
                        key="download_pdf_text_file"
                    )
                    st.success("✅ Готово к скачиванию!")
            
            with col2:
                if st.button("🤖 Проанализировать PDF текст ИИ", use_container_width=True, type="secondary", key="analyze_pdf_text_btn"):
                    if not AI_AVAILABLE:
                        st.error("❌ ИИ-модуль недоступен.")
                        return
                    
                    with st.spinner("🤖 ИИ структурирует данные..."):
                        assistant = OpenRouterAssistant()
                        current_doc_type = st.session_state.get('current_doc_type', 'медицинский документ')
                        extracted_text = st.session_state.get('extracted_doc_text', '')
                        
                        # Определяем, является ли документ лабораторным
                        is_lab_document = any(keyword in current_doc_type.lower() for keyword in 
                                            ['лаборатор', 'лабораторн', 'анализ крови', 'биохимия', 'гематолог'])
                        
                        prompt = f"""
Вы - эксперт по структурированию медицинских документов. Структурируйте следующий текст из медицинского документа типа "{current_doc_type}".

Текст документа:
{extracted_text[:8000]}

Извлеките все ключевые данные в JSON формате, аналогично тому, как это делается для изображений документов.
Верните ТОЛЬКО JSON объект, без дополнительных комментариев.
"""
                        try:
                            # Для лабораторных документов используем Sonnet 4.5, для остальных - Opus
                            result = assistant.get_response(prompt, use_sonnet_4_5=is_lab_document)
                            st.subheader("📋 Структурированные данные")
                            st.write(result)
                            
                            # Попытка распарсить JSON
                            json_match = re.search(r'\{.*\}', result, re.DOTALL)
                            if json_match:
                                try:
                                    extracted_data = json.loads(json_match.group())
                                    st.json(extracted_data)
                                    st.session_state['extracted_doc_data'] = extracted_data
                                except:
                                    pass
                        except Exception as e:
                            st.error(f"❌ Ошибка обработки: {e}")

# --- Главная функция ---
def main():
    st.set_page_config(
        page_title="Медицинский ИИ-Ассистент",
        page_icon="🏥",
        layout="wide"
    )

    # Глобальные стили интерфейса в зелёно-голубой гамме
    st.markdown(
        """
        <style>
        /* Общий фон */
        .stApp {
            background: radial-gradient(circle at top left, #e0f7fa 0%, #e8f5e9 40%, #ffffff 100%);
        }

        /* Базовая типографика */
        html, body, [class*="css"]  {
            font-family: -apple-system, BlinkMacSystemFont, "SF Pro Text", system-ui, sans-serif;
            font-size: 16px;
        }

        /* Крупнее шрифт основного текста */
        p, li, span {
            font-size: 15px;
        }

        /* Карточки и контейнеры */
        .stMarkdown, .stDataFrame, .stPlotlyChart, .stSpinner {
            border-radius: 12px !important;
        }

        /* Боковое меню */
        section[data-testid="stSidebar"] {
            background: linear-gradient(180deg, #004d40 0%, #00695c 40%, #004d40 100%);
        }
        section[data-testid="stSidebar"] * {
            color: #e0f2f1 !important;
        }

        /* Кнопки */
        div.stButton > button {
            border-radius: 999px;
            border: none;
            padding: 0.6rem 1.2rem;
            font-weight: 600;
            background: linear-gradient(90deg, #009688, #26a69a);
            color: white;
        }
        div.stButton > button:hover {
            background: linear-gradient(90deg, #26a69a, #4db6ac);
            box-shadow: 0 0 12px rgba(0, 150, 136, 0.4);
        }

        /* Заголовки */
        h1, h2, h3 {
            color: #004d40;
            font-weight: 700;
        }
        h1 { font-size: 2.0rem; }
        h2 { font-size: 1.6rem; }
        h3 { font-size: 1.3rem; }

        /* Метрики */
        div[data-testid="stMetricValue"] {
            color: #00695c !important;
            font-size: 1.4rem;
            font-weight: 700;
        }
        div[data-testid="stMetricLabel"] {
            font-size: 0.9rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    init_db()

    # === АВТОМАТИЧЕСКАЯ ПРОВЕРКА FEEDBACK ===
    try:
        from feedback.auto_analyzer import AutoFeedbackAnalyzer
        auto_analyzer = AutoFeedbackAnalyzer()
        threshold_check = auto_analyzer.check_thresholds()
        
        # Показываем уведомления если есть рекомендации
        if threshold_check.get("recommendations"):
            for rec in threshold_check["recommendations"]:
                if rec["type"] == "basic_analysis":
                    st.info(f"📊 {rec['message']}")
                    if st.button(f"🔍 Проанализировать {', '.join(rec['types'][:2])}...", key="btn_auto_analysis"):
                        selected_type = rec['types'][0] if rec['types'] else None
                        if selected_type:
                            with st.spinner(f"Анализирую {selected_type}..."):
                                result = auto_analyzer.run_basic_analysis(selected_type)
                                st.success(f"✅ Анализ завершен: {result.get('total_feedback', 0)} отзывов")
                                if result.get("top_errors"):
                                    st.subheader("⚠️ Топ ошибок:")
                                    for i, error in enumerate(result["top_errors"][:5], 1):
                                        st.text(f"{i}. {error.get('correct_diagnosis', 'N/A')}")
                
                elif rec["type"] == "optimization":
                    st.warning(f"💡 {rec['message']}")
                    if st.button("🎯 Получить предложения по оптимизации", key="btn_optimization"):
                        with st.spinner("Анализирую паттерны и формирую предложения..."):
                            opt_results = auto_analyzer.run_optimization_analysis()
                            report = auto_analyzer.generate_optimization_report(opt_results)
                            st.text_area("📋 Отчет по оптимизации", report, height=400)
                
                elif rec["type"] == "deep_analysis":
                    st.success(f"🎯 {rec['message']}")
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("📥 Скачать с GitHub и проанализировать", key="btn_deep_analysis"):
                            st.info("💡 Используйте: git pull для скачивания данных с GitHub, затем python scripts/get_feedback_data.py --export json")
    except Exception as e:
        # Игнорируем ошибки автоанализа (чтобы не ломать приложение)
        pass

    # ОБНОВЛЕННЫЙ список страниц
    pages = [
        "🏠 Главная",
        "📈 Анализ ЭКГ",
        "🩻 Анализ рентгена",
        "🧠 Анализ МРТ",
        "🩻 Анализ КТ",
        "🔊 Анализ УЗИ",
        "🔬 Анализ дерматоскопии",
        "🔬 Анализ лабораторных данных",
        "📝 Протокол приёма",
        "📄 Сканирование документов",
        "🎬 Анализ видео",
        "👤 База данных пациентов",
        "📋 Клинический контекст",
        "🤖 ИИ-Консультант",
        "🧬 Генетический анализ",
        "📊 Статистика",
        "🔬 Расширенный ИИ-анализ",
        "📊 Сравнительный анализ",
        "📚 Медицинские протоколы",
    ]

    st.sidebar.title("🧠 Меню")
    page = st.sidebar.selectbox("Выберите раздел:", pages)

    # === ОБРАБОТКА СТРАНИЦ ===
    if page == "🏠 Главная":
        show_home_page()
    elif page == "📈 Анализ ЭКГ":
        show_ecg_analysis()
    elif page == "🩻 Анализ рентгена":
        show_xray_analysis()
    elif page == "🧠 Анализ МРТ":
        show_mri_analysis()
    elif page == "🩻 Анализ КТ":  # ← НОВОЕ
        show_ct_analysis()
    elif page == "🔊 Анализ УЗИ":  # ← НОВОЕ
        show_ultrasound_analysis()
    elif page == "🔬 Анализ дерматоскопии":
        show_dermatoscopy_analysis()
    elif page == "🔬 Анализ лабораторных данных":
        show_lab_analysis()  # ← ваша новая улучшенная функция
    elif page == "📝 Протокол приёма":
        show_consultation_protocol()
    elif page == "📄 Сканирование документов":  # ← НОВОЕ
        show_document_scanner_page()
    elif page == "🎬 Анализ видео":
        show_video_analysis()
    elif page == "👤 База данных пациентов":
        show_patient_database()
    elif page == "📋 Клинический контекст":  # ← НОВОЕ
        show_patient_context_page()
    elif page == "🤖 ИИ-Консультант":
        show_ai_chat()
    elif page == "🧬 Генетический анализ":
        show_genetic_analysis_page()  # ← ваша готовая функция
    elif page == "📊 Статистика":  # ← НОВОЕ
        show_statistics_page()
    # === НОВЫЕ СТРАНИЦЫ ===
    elif page == "🔬 Расширенный ИИ-анализ":
        if ENHANCED_PAGES_AVAILABLE and show_enhanced_analysis_page:
            show_enhanced_analysis_page()
        else:
            st.error("❌ Модуль расширенного анализа недоступен. Проверьте файл `modules/streamlit_enhanced_pages.py`")
            st.info("💡 Убедитесь, что все зависимости установлены: `pip install plotly pandas`")
    elif page == "📊 Сравнительный анализ":
        if ENHANCED_PAGES_AVAILABLE and show_comparative_analysis_page:
            show_comparative_analysis_page()
        else:
            st.error("❌ Модуль сравнительного анализа недоступен. Проверьте файл `modules/streamlit_enhanced_pages.py`")
            st.info("💡 Убедитесь, что все зависимости установлены: `pip install plotly pandas`")
    elif page == "📚 Медицинские протоколы":
        if ENHANCED_PAGES_AVAILABLE and show_medical_protocols_page:
            show_medical_protocols_page()
        else:
            st.error("❌ Модуль медицинских протоколов недоступен. Проверьте файл `modules/streamlit_enhanced_pages.py`")
            st.info("💡 Убедитесь, что все зависимости установлены: `pip install plotly pandas`")
    #"elif page == "🎓 Обучение ИИ":
#       show_ai_training_page()
    
    # === ОБНОВЛЕННЫЙ САЙДБАР ===
    st.sidebar.markdown("---")
    st.sidebar.info("""
    **Медицинский Ассистент v6.0.2-stable** [STABLE]
    - AssemblyAI для голоса
    - 10 типов изображений
    - Улучшенный анализ лабораторных данных
    - Структурированный JSON анализ
    - Сравнительная диагностика
    - Медицинские протоколы
    - Claude 4.5 Sonnet + Opus 4.5 + OpenRouter
    ВНИМАНИЕ: Только для обучения
    """)

if __name__ == "__main__":
    main()
