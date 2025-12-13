#до квена 
"""
Интеграция улучшенного ИИ-анализатора с Streamlit
Новые страницы для расширенного функционала
"""

import streamlit as st
import numpy as np
from PIL import Image
import pandas as pd
import json
import plotly.graph_objects as go
import plotly.express as px
from typing import Dict, List, Optional, Any
import io
import base64
import requests
import time
import datetime
try:
    from .medical_ai_analyzer import EnhancedMedicalAIAnalyzer, ImageType, AnalysisResult
except ImportError:
    try:
        from modules.medical_ai_analyzer import EnhancedMedicalAIAnalyzer, ImageType, AnalysisResult
    except ImportError:
        # Fallback - используем базовый класс из claude_assistant
        EnhancedMedicalAIAnalyzer = None
        ImageType = None
        AnalysisResult = None


def ensure_string_for_download(data: Any) -> str:
    """Безопасное преобразование данных в строку для download_button
    
    Args:
        data: Данные любого типа для преобразования в строку
        
    Returns:
        str: Строковое представление данных
    """
    if isinstance(data, tuple):
        # Если это кортеж, берем первый элемент или преобразуем весь кортеж
        if len(data) > 0:
            return str(data[0])
        return str(data)
    elif isinstance(data, (list, dict)):
        # Если это список или словарь, преобразуем в JSON строку
        return json.dumps(data, ensure_ascii=False, indent=2)
    elif not isinstance(data, str):
        # Любой другой тип преобразуем в строку
        return str(data)
    return data


def show_enhanced_analysis_page():
    """Страница расширенного ИИ-анализа"""
    st.header("🔬 Расширенный ИИ-Анализ")
    
    # Проверка доступности анализатора
    if EnhancedMedicalAIAnalyzer is None or ImageType is None:
        st.error("❌ Модуль EnhancedMedicalAIAnalyzer недоступен")
        st.info("💡 Убедитесь, что файл `modules/medical_ai_analyzer.py` существует и правильно настроен")
        return
    
    # Инициализация анализатора
    if 'enhanced_analyzer' not in st.session_state:
        try:
            from config import OPENROUTER_API_KEY
            api_key = OPENROUTER_API_KEY
        except:
            # Получаем ключ из config или secrets
            try:
                from config import OPENROUTER_API_KEY
                api_key = OPENROUTER_API_KEY
            except ImportError:
                api_key = st.secrets.get("api_keys", {}).get("OPENROUTER_API_KEY") or st.secrets.get("OPENROUTER_API_KEY")
        
        try:
            st.session_state.enhanced_analyzer = EnhancedMedicalAIAnalyzer(api_key)
        except Exception as e:
            st.error(f"❌ Ошибка инициализации анализатора: {e}")
            return
    
    analyzer = st.session_state.enhanced_analyzer
    
    # Настройки анализа
    col1, col2, col3 = st.columns(3)
    
    with col1:
        preprocessing = st.checkbox("Предобработка изображения", value=True)
        batch_mode = st.checkbox("Пакетный режим", value=False)
    
    with col2:
        confidence_threshold = st.slider("Порог достоверности", 0.0, 1.0, 0.7, 0.1)
        show_metadata = st.checkbox("Показать метаданные", value=False)
    
    with col3:
        st.info("💡 Анализ выполняется автоматически для любого типа медицинского изображения")
    
    # Загрузка файлов
    if batch_mode:
        uploaded_files = st.file_uploader(
            "Загрузите медицинские изображения",
            type=["jpg", "jpeg", "png", "dcm", "tiff"],
            accept_multiple_files=True
        )
    else:
        uploaded_file = st.file_uploader(
            "Загрузите медицинское изображение",
            type=["jpg", "jpeg", "png", "dcm", "tiff"]
        )
        uploaded_files = [uploaded_file] if uploaded_file else []
    
    # Дополнительный контекст
    additional_context = st.text_area(
        "Дополнительная клиническая информация",
        placeholder="Введите анамнез, жалобы пациента, предварительный диагноз..."
    )
    
    if uploaded_files and st.button("🚀 Запустить расширенный анализ"):
        
        # Обработка изображений
        images_data = []
        
        for uploaded_file in uploaded_files:
            if uploaded_file is not None:
                try:
                    # Загрузка изображения
                    image = Image.open(uploaded_file)
                    if image.mode != 'RGB' and image.mode != 'L':
                        image = image.convert('RGB')
                    
                    image_array = np.array(image)
                    
                    # Всегда используем универсальный анализ без определения типа
                    image_type = None
                    
                    images_data.append((image_array, image_type, uploaded_file.name))
                    
                except Exception as e:
                    st.error(f"Ошибка обработки файла {uploaded_file.name}: {e}")
        
        if images_data:
            # Прогресс-бар
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            results = []
            
            for i, (image_array, image_type, filename) in enumerate(images_data):
                status_text.text(f"Анализ {filename}...")
                progress_bar.progress((i + 1) / len(images_data))
                
                try:
                    result = analyzer.analyze_image(
                        image_array, 
                        image_type, 
                        additional_context
                    )
                    result.filename = filename
                    results.append(result)
                    
                except Exception as e:
                    st.error(f"Ошибка анализа {filename}: {e}")
            
            progress_bar.empty()
            status_text.empty()
            
            # Сохраняем результаты в session_state
            st.session_state['enhanced_analysis_results'] = results
            st.session_state['enhanced_analysis_timestamp'] = datetime.datetime.now().isoformat()
            
            # Отображение результатов
            if results:
                st.success(f"✅ Анализ завершен! Обработано изображений: {len(results)}")
                
                # Сводная статистика
                show_analysis_summary(results, confidence_threshold)
                
                # Детальные результаты
                for result in results:
                    show_detailed_analysis_result(result, show_metadata)
    
    # Показываем сохраненные результаты и кнопку генерации протокола
    if 'enhanced_analysis_results' in st.session_state and st.session_state['enhanced_analysis_results']:
        results = st.session_state['enhanced_analysis_results']
        
        st.markdown("---")
        st.subheader("📄 Генерация медицинского протокола")
        
        # Показываем сохраненный протокол если есть
        protocol_key = f'generated_report_{len(results)}'
        docx_key = f'{protocol_key}_docx'
        
        if protocol_key in st.session_state:
            saved_report = st.session_state[protocol_key]
            if saved_report:
                st.info("💡 Отображается ранее сгенерированный протокол. Нажмите кнопку ниже, чтобы сгенерировать новый.")
                st.text_area("📋 Медицинский протокол (текстовый просмотр)", saved_report, height=400, key=f"saved_protocol_text_{len(results)}")
                
                # Кнопки скачивания
                col1, col2 = st.columns(2)
                
                with col1:
                    # Скачать Word документ (если есть)
                    if docx_key in st.session_state:
                        st.download_button(
                            label="📄 Скачать протокол Word (.docx)",
                            data=st.session_state[docx_key],
                            file_name=f"medical_protocol_{len(results)}_images.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key=f"download_docx_report_{len(results)}"
                        )
                    else:
                        st.info("💡 Сгенерируйте протокол, чтобы получить Word документ")
                
                with col2:
                    # Скачать текстовый файл
                    st.download_button(
                        label="📝 Скачать протокол (.txt)",
                        data=saved_report,
                        file_name=f"medical_protocol_{len(results)}_images.txt",
                        mime="text/plain",
                        use_container_width=True,
                        key=f"download_txt_report_{len(results)}"
                    )
                
                st.markdown("---")
        
        if st.button("📄 Сгенерировать медицинский протокол", use_container_width=True, type="primary", key="generate_protocol_button"):
            try:
                with st.spinner("🔄 Генерирую протокол..."):
                    if 'enhanced_analyzer' not in st.session_state:
                        st.error("❌ Анализатор не найден. Выполните анализ заново.")
                    else:
                        analyzer = st.session_state.enhanced_analyzer
                        report_raw = analyzer.generate_report(results)
                        report = ensure_string_for_download(report_raw)
                        
                        if report and len(report.strip()) > 0:
                            # Сохраняем текстовый отчет в session_state
                            st.session_state[protocol_key] = report
                            
                            # Генерируем Word документ
                            try:
                                from docx import Document
                                from docx.shared import Pt, Inches
                                from docx.enum.text import WD_ALIGN_PARAGRAPH
                                import io
                                
                                doc = Document()
                                
                                # Настройка стилей
                                style = doc.styles['Normal']
                                font = style.font
                                font.name = 'Times New Roman'
                                font.size = Pt(12)
                                
                                # Заголовок
                                title = doc.add_heading('МЕДИЦИНСКОЕ ЗАКЛЮЧЕНИЕ', 0)
                                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                # Дата
                                date_para = doc.add_paragraph(f"Дата: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}")
                                date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                doc.add_paragraph()
                                
                                # Разделитель
                                doc.add_paragraph('─' * 80)
                                
                                # Разбиваем текст протокола на строки и форматируем
                                lines = report.split('\n')
                                for line in lines:
                                    line = line.strip()
                                    if not line:
                                        doc.add_paragraph()
                                        continue
                                    
                                    # Заголовки (начинаются с заглавных букв и содержат только заглавные или цифры)
                                    if line.isupper() or (len(line) < 60 and line.isupper()):
                                        doc.add_heading(line, level=1)
                                    elif line.startswith('='):
                                        doc.add_paragraph('─' * 80)
                                    elif line.startswith('-'):
                                        doc.add_paragraph(line, style='List Bullet')
                                    else:
                                        para = doc.add_paragraph(line)
                                
                                # Сохраняем Word документ в байты
                                doc_bytes = io.BytesIO()
                                doc.save(doc_bytes)
                                doc_bytes.seek(0)
                                
                                # Сохраняем Word документ в session_state
                                docx_bytes = doc_bytes.getvalue()
                                st.session_state[f'{protocol_key}_docx'] = docx_bytes
                                st.success("✅ Протокол успешно сгенерирован!")
                                
                                # Сразу показываем кнопки скачивания
                                col1, col2 = st.columns(2)
                                
                                with col1:
                                    st.download_button(
                                        label="📄 Скачать протокол Word (.docx)",
                                        data=docx_bytes,
                                        file_name=f"medical_protocol_{len(results)}_images.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        use_container_width=True,
                                        key=f"download_docx_immediate_{len(results)}"
                                    )
                                
                                with col2:
                                    st.download_button(
                                        label="📝 Скачать протокол TXT (.txt)",
                                        data=report,
                                        file_name=f"medical_protocol_{len(results)}_images.txt",
                                        mime="text/plain",
                                        use_container_width=True,
                                        key=f"download_txt_immediate_{len(results)}"
                                    )
                                
                                st.rerun()  # Обновляем страницу чтобы показать протокол
                                
                            except ImportError:
                                st.warning("⚠️ python-docx не установлен. Протокол доступен только в текстовом формате.")
                                st.info("💡 Установите: pip install python-docx")
                                st.success("✅ Текстовый протокол успешно сгенерирован!")
                                st.rerun()
                        else:
                            st.error("❌ Ошибка: Протокол пуст. Проверьте результаты анализа.")
            except AttributeError as e:
                st.error(f"❌ Ошибка: Метод generate_report не найден: {e}")
                st.info("💡 Убедитесь, что используется правильная версия EnhancedMedicalAIAnalyzer")
            except Exception as e:
                st.error(f"❌ Ошибка генерации протокола: {e}")
                import traceback
                with st.expander("🔍 Детали ошибки"):
                    st.code(traceback.format_exc())


def show_analysis_summary(results: List[AnalysisResult], confidence_threshold: float):
    """Показывает сводную статистику анализа"""
    
    st.subheader("📊 Сводная статистика анализа")
    
    # Метрики
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        total_images = len(results)
        st.metric("Всего изображений", total_images)
    
    with col2:
        high_confidence = len([r for r in results if r.confidence >= confidence_threshold])
        st.metric("Высокая достоверность", f"{high_confidence}/{total_images}")
    
    with col3:
        urgent_cases = len([r for r in results if r.urgent_flags])
        st.metric("Срочные случаи", urgent_cases, delta="⚠️" if urgent_cases > 0 else None)
    
    with col4:
        avg_confidence = np.mean([r.confidence for r in results])
        st.metric("Средняя достоверность", f"{avg_confidence:.1%}")
    
    # Распределение по типам изображений
    col1, col2 = st.columns(2)
    
    with col1:
        type_counts = {}
        for result in results:
            type_name = result.image_type.value
            type_counts[type_name] = type_counts.get(type_name, 0) + 1
        
        if type_counts:
            fig_pie = px.pie(
                values=list(type_counts.values()),
                names=list(type_counts.keys()),
                title="Распределение по типам изображений"
            )
            st.plotly_chart(fig_pie, use_container_width=True)
    
    with col2:
        # График достоверности
        confidence_data = pd.DataFrame({
            'Изображение': [getattr(r, 'filename', f'Изображение {i+1}') for i, r in enumerate(results)],
            'Достоверность': [r.confidence for r in results],
            'Тип': [r.image_type.value for r in results]
        })
        
        fig_bar = px.bar(
            confidence_data,
            x='Изображение',
            y='Достоверность',
            color='Тип',
            title="Достоверность анализа по изображениям"
        )
        fig_bar.add_hline(y=confidence_threshold, line_dash="dash", line_color="red", 
                         annotation_text="Порог достоверности")
        st.plotly_chart(fig_bar, use_container_width=True)


def show_detailed_analysis_result(result: AnalysisResult, show_metadata: bool = False):
    """Показывает детальный результат анализа"""
    
    filename = getattr(result, 'filename', 'Изображение')
    
    with st.expander(f"🔍 Детальный анализ: {filename}", expanded=True):
        
        # Основная информация
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            confidence_color = "green" if result.confidence > 0.8 else "orange" if result.confidence > 0.6 else "red"
            st.markdown(f"**Достоверность:** :{confidence_color}[{result.confidence:.1%}]")
        with col2:
            if result.urgent_flags:
                st.error(f"⚠️ Срочно: {len(result.urgent_flags)} предупреждений")
            else:
                st.success("✅ Плановое наблюдение")
        with col3:
            # Информация о модели
            if hasattr(result, 'model_name') and result.model_name:
                st.info(f"🤖 **Модель:** {result.model_name}")
        with col4:
            # Информация о токенах
            if hasattr(result, 'tokens_used') and result.tokens_used > 0:
                st.metric("📊 Токенов", result.tokens_used)
        
        # Дополнительная информация о модели и токенах
        if hasattr(result, 'model_name') and result.model_name and hasattr(result, 'tokens_used') and result.tokens_used > 0:
            st.caption(f"🤖 Анализ выполнен моделью: **{result.model_name}** | 📊 Использовано токенов: **{result.tokens_used}**")
        
        # Структурированные находки
        if result.structured_findings:
            findings = result.structured_findings
            
            # Техническая оценка
            if "technical_assessment" in findings:
                st.subheader("🔧 Техническая оценка")
                tech = findings["technical_assessment"]
                
                col1, col2 = st.columns(2)
                with col1:
                    quality = tech.get("quality", "не определено")
                    quality_color = {"отличное": "green", "хорошее": "green", 
                                   "удовлетворительное": "orange", "плохое": "red"}.get(quality, "gray")
                    st.markdown(f"**Качество:** :{quality_color}[{quality}]")
                
                with col2:
                    artifacts = tech.get("artifacts", [])
                    if artifacts:
                        st.warning(f"Артефакты: {', '.join(artifacts)}")
                    else:
                        st.success("Артефакты не обнаружены")
            
            # Клинические находки
            if "clinical_findings" in findings:
                st.subheader("🏥 Клинические находки")
                clinical = findings["clinical_findings"]
                
                # Нормальные структуры
                normal = clinical.get("normal_structures", [])
                if normal:
                    st.success(f"**Нормальные структуры:** {', '.join(normal)}")
                
                # Патологические находки
                pathological = clinical.get("pathological_findings", [])
                if pathological:
                    st.warning("**Патологические изменения:**")
                    for finding in pathological:
                        with st.container():
                            st.markdown(f"• **{finding.get('finding', 'Находка')}**")
                            if finding.get('location'):
                                st.markdown(f"  📍 Локализация: {finding['location']}")
                            if finding.get('severity'):
                                st.markdown(f"  📊 Выраженность: {finding['severity']}")
                            if finding.get('description'):
                                st.markdown(f"  📝 Описание: {finding['description']}")
                else:
                    st.success("Патологических изменений не выявлено")
            
            # Диагноз
            if "diagnosis" in findings:
                st.subheader("🎯 Диагноз")
                diagnosis = findings["diagnosis"]
                
                primary = diagnosis.get("primary_diagnosis", "Не определен")
                st.markdown(f"**Основной диагноз:** {primary}")
                
                differential = diagnosis.get("differential_diagnosis", [])
                if differential:
                    st.markdown("**Дифференциальная диагностика:**")
                    for diff_diag in differential:
                        st.markdown(f"• {diff_diag}")
                
                icd10 = diagnosis.get("icd10_codes", [])
                if icd10:
                    st.info(f"**Коды МКБ-10:** {', '.join(icd10)}")
            
            # Рекомендации
            if "recommendations" in findings:
                st.subheader("📋 Рекомендации")
                recommendations = findings["recommendations"]
                
                urgent = recommendations.get("urgent_actions", [])
                if urgent:
                    st.error("**⚠️ Срочные действия:**")
                    for action in urgent:
                        st.markdown(f"• {action}")
                
                follow_up = recommendations.get("follow_up", [])
                if follow_up:
                    st.info("**📅 План наблюдения:**")
                    for plan in follow_up:
                        st.markdown(f"• {plan}")
                
                additional = recommendations.get("additional_studies", [])
                if additional:
                    st.info("**🔬 Дополнительные исследования:**")
                    for study in additional:
                        st.markdown(f"• {study}")
            
            # Оценка риска
            if "risk_assessment" in findings:
                st.subheader("⚡ Оценка риска")
                risk = findings["risk_assessment"]
                
                urgency = risk.get("urgency_level", "планово")
                urgency_color = {"экстренно": "red", "срочно": "orange", "планово": "green"}.get(urgency, "gray")
                st.markdown(f"**Уровень срочности:** :{urgency_color}[{urgency}]")
                
                risk_factors = risk.get("risk_factors", [])
                if risk_factors:
                    st.warning(f"**Факторы риска:** {', '.join(risk_factors)}")
                
                prognosis = risk.get("prognosis", "")
                if prognosis:
                    st.info(f"**Прогноз:** {prognosis}")
        
        # Клиническая интерпретация (читаемый форматированный текст)
        # УБРАНО: Не показываем clinical_interpretation, так как вся информация уже показана выше
        # в структурированном виде (находки, диагноз, рекомендации)
        
        # НЕ показываем structured_findings как JSON - они уже показаны выше в структурированном виде
        # Убираем дублирование JSON внизу
        
        # Метаданные
        if show_metadata and hasattr(result, 'metadata') and result.metadata:
            st.subheader("🔍 Метаданные изображения")
            
            metadata_df = pd.DataFrame([
                {"Параметр": k, "Значение": str(v)} 
                for k, v in result.metadata.items()
            ])
            st.dataframe(metadata_df, use_container_width=True)


def show_comparative_analysis_page():
    """Страница сравнительного анализа"""
    st.header("📊 Сравнительный анализ изображений")
    
    # Проверка доступности анализатора
    if EnhancedMedicalAIAnalyzer is None or ImageType is None:
        st.error("❌ Модуль EnhancedMedicalAIAnalyzer недоступен")
        st.info("💡 Убедитесь, что файл `modules/medical_ai_analyzer.py` существует и правильно настроен")
        return
    
    st.info("💡 Загрузите несколько изображений одного типа для сравнения динамики или разных проекций")
    
    # Настройки сравнения
    comparison_type = st.selectbox(
        "Тип сравнения",
        ["Временная динамика", "Разные проекции", "До/после лечения", "Межпациентное сравнение"]
    )
    
    # Дополнительные настройки
    col1, col2 = st.columns(2)
    with col1:
        force_same_type = st.checkbox("Принудительно одинаковый тип", value=True, 
                                     help="Все изображения будут анализироваться как один тип")
    with col2:
        show_debug_info = st.checkbox("Показать отладочную информацию", value=False)
    
    # Загрузка изображений для сравнения
    uploaded_files = st.file_uploader(
        "Загрузите изображения для сравнения",
        type=["jpg", "jpeg", "png", "dcm"],
        accept_multiple_files=True,
        help="Оптимально 2-4 изображения одного типа"
    )
    
    if uploaded_files and len(uploaded_files) >= 2:
        
        st.success(f"✅ Загружено {len(uploaded_files)} изображений")
        
        # Предварительный просмотр загруженных изображений
        st.subheader("📸 Предварительный просмотр")
        preview_cols = st.columns(min(len(uploaded_files), 4))
        
        for i, uploaded_file in enumerate(uploaded_files):
            with preview_cols[i % len(preview_cols)]:
                try:
                    image = Image.open(uploaded_file)
                    st.image(image, caption=uploaded_file.name, use_container_width=True)
                    st.caption(f"Размер: {image.size[0]}×{image.size[1]}")
                except Exception as e:
                    st.error(f"Ошибка загрузки {uploaded_file.name}: {e}")
        
        # Проверяем, есть ли сохраненные результаты анализа
        analysis_key = f"comparative_analysis_{len(uploaded_files)}_{comparison_type}"
        saved_results = st.session_state.get('comparative_analysis_results', {}).get(analysis_key)
        saved_images = st.session_state.get('comparative_analysis_images', {}).get(analysis_key)
        
        # Если есть сохраненные результаты, используем их
        if saved_results and saved_images:
            results = saved_results
            images = saved_images
            st.info(f"💡 Используются сохраненные результаты анализа ({len(results)} изображений)")
            
            # Кнопка для повторного анализа
            if st.button("🔄 Выполнить анализ заново", key="rerun_analysis"):
                # Очищаем сохраненные результаты
                if 'comparative_analysis_results' in st.session_state:
                    if analysis_key in st.session_state.comparative_analysis_results:
                        del st.session_state.comparative_analysis_results[analysis_key]
                if 'comparative_analysis_images' in st.session_state:
                    if analysis_key in st.session_state.comparative_analysis_images:
                        del st.session_state.comparative_analysis_images[analysis_key]
                st.rerun()
        else:
            # Выполняем новый анализ только если нет сохраненных результатов
            if st.button("🔄 Выполнить сравнительный анализ", key="run_analysis"):
                
                # Проверка доступности анализатора
                if EnhancedMedicalAIAnalyzer is None:
                    st.error("❌ Модуль EnhancedMedicalAIAnalyzer недоступен")
                    return
                
                # Инициализация анализатора
                if 'enhanced_analyzer' not in st.session_state:
                    # Получаем ключ из config или secrets
                    try:
                        from config import OPENROUTER_API_KEY
                        api_key = OPENROUTER_API_KEY
                    except ImportError:
                        api_key = st.secrets.get("api_keys", {}).get("OPENROUTER_API_KEY") or st.secrets.get("OPENROUTER_API_KEY")
                    
                    try:
                        st.session_state.enhanced_analyzer = EnhancedMedicalAIAnalyzer(api_key)
                    except Exception as e:
                        st.error(f"❌ Ошибка инициализации анализатора: {e}")
                        return
                
                analyzer = st.session_state.enhanced_analyzer
                
                # Анализ каждого изображения
                results = []
                images = []
                
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                for i, uploaded_file in enumerate(uploaded_files):
                    status_text.text(f"Анализ изображения {i+1}/{len(uploaded_files)}: {uploaded_file.name}")
                    progress_bar.progress((i + 1) / len(uploaded_files))
                    
                    try:
                        # Загрузка и обработка изображения
                        image = Image.open(uploaded_file)
                        if image.mode not in ['RGB', 'L']:
                            image = image.convert('RGB')
                        
                        image_array = np.array(image)
                        images.append(image_array)
                        
                        # Универсальный анализ без определения типа
                        result = analyzer.analyze_image(
                            image_array,
                            None,  # Не определяем тип - универсальный анализ
                            additional_context=f"Сравнительный анализ ({comparison_type}), изображение {i+1} из {len(uploaded_files)}"
                        )
                        result.filename = uploaded_file.name
                        results.append(result)
                        
                    except Exception as e:
                        st.error(f"Ошибка обработки {uploaded_file.name}: {e}")
                        import traceback
                        st.error(f"Детали ошибки: {traceback.format_exc()}")
                        continue
                
                progress_bar.empty()
                status_text.empty()
                
                # Сохраняем результаты в session_state
                if 'comparative_analysis_results' not in st.session_state:
                    st.session_state.comparative_analysis_results = {}
                if 'comparative_analysis_images' not in st.session_state:
                    st.session_state.comparative_analysis_images = {}
                
                st.session_state.comparative_analysis_results[analysis_key] = results
                st.session_state.comparative_analysis_images[analysis_key] = images
                
                st.success(f"✅ Анализ завершен! Обработано изображений: {len(results)}")
                st.rerun()  # Перезагружаем страницу для отображения результатов
                return
        
        # Отображаем результаты, если они есть
        if saved_results and saved_images:
            results = saved_results
            images = saved_images
            
            if results:
                st.success(f"✅ Анализ завершен! Обработано изображений: {len(results)}")
                
                # Отображение результатов анализа
                st.subheader("🖼️ Результаты анализа")
                
                # Создаем адаптивную сетку для изображений
                num_cols = min(len(results), 3)
                cols = st.columns(num_cols)
                
                for i, (image_array, result) in enumerate(zip(images, results)):
                    with cols[i % num_cols]:
                        # Отображаем изображение
                        st.image(image_array, caption=result.filename, use_container_width=True)
                        
                        # Метрики анализа
                        st.metric("Достоверность", f"{result.confidence:.1%}")
                        
                        # Информация о модели и токенах
                        if hasattr(result, 'model_name') and result.model_name:
                            st.caption(f"🤖 {result.model_name}")
                        if hasattr(result, 'tokens_used') and result.tokens_used > 0:
                            st.caption(f"📊 Токенов: {result.tokens_used}")
                        
                        # Размер изображения для отладки
                        if show_debug_info:
                            st.caption(f"Размер: {image_array.shape}")
                
                # Сравнительная таблица
                st.subheader("📋 Сравнительная таблица")
                
                comparison_data = []
                for result in results:
                    findings = result.structured_findings
                    
                    comparison_data.append({
                        "Файл": result.filename,
                        "Достоверность": f"{result.confidence:.1%}",
                        "Основной диагноз": findings.get("diagnosis", {}).get("primary_diagnosis", "Не определен"),
                        "Качество": findings.get("technical_assessment", {}).get("quality", "Не оценено"),
                        "Срочность": findings.get("risk_assessment", {}).get("urgency_level", "планово"),
                        "Патология": "Да" if findings.get("clinical_findings", {}).get("pathological_findings") else "Нет"
                    })
                
                comparison_df = pd.DataFrame(comparison_data)
                st.dataframe(comparison_df, use_container_width=True)
                
                # Анализ динамики (если применимо)
                if comparison_type == "Временная динамика":
                    st.subheader("📈 Анализ динамики")
                    
                    # График изменения достоверности
                    confidence_trend = [r.confidence for r in results]
                    
                    fig = go.Figure()
                    fig.add_trace(go.Scatter(
                        x=list(range(1, len(confidence_trend) + 1)),
                        y=confidence_trend,
                        mode='lines+markers',
                        name='Достоверность анализа',
                        line=dict(color='blue', width=3),
                        marker=dict(size=8)
                    ))
                    
                    fig.update_layout(
                        title="Динамика достоверности анализа",
                        xaxis_title="Номер исследования",
                        yaxis_title="Достоверность",
                        yaxis=dict(range=[0, 1]),
                        height=400
                    )
                    
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Анализ изменений в диагнозах
                    diagnoses = [r.structured_findings.get("diagnosis", {}).get("primary_diagnosis", "Не определен") for r in results]
                    unique_diagnoses = len(set(diagnoses))
                    
                    if unique_diagnoses == 1:
                        st.success("✅ Диагноз стабилен во всех исследованиях")
                    else:
                        st.warning(f"⚠️ Обнаружены изменения в диагнозах ({unique_diagnoses} различных)")
                
                # ИИ-заключение по сравнению
                st.markdown("---")
                st.subheader("🤖 ИИ-заключение по сравнительному анализу")
                
                # Проверяем, есть ли сохраненное заключение
                saved_conclusion_key = f"{comparison_type}_{len(results)}"
                saved_conclusion_raw = st.session_state.get('comparative_analysis_result', {}).get(saved_conclusion_key, '')
                
                # Преобразуем в строку если это не строка (может быть tuple или другой тип)
                if saved_conclusion_raw:
                    saved_conclusion = ensure_string_for_download(saved_conclusion_raw)
                    
                    st.info("💡 Отображается сохраненное заключение. Нажмите кнопку ниже, чтобы сгенерировать новое.")
                    st.markdown("### 📋 Сравнительное заключение")
                    st.markdown(saved_conclusion)
                    st.markdown("---")
                    st.download_button(
                        label="💾 Скачать заключение",
                        data=saved_conclusion,
                        file_name=f"comparative_analysis_{comparison_type}_{len(results)}_images.txt",
                        mime="text/plain",
                        use_container_width=True,
                        key=f"download_saved_conclusion_{saved_conclusion_key}"
                    )
                    st.markdown("---")
                
                # Автоматически генерируем заключение или показываем кнопку
                if st.button("📝 Сгенерировать сравнительное заключение", use_container_width=True, type="primary", key="generate_conclusion"):
                    
                    # Получаем анализатор из session_state
                    if 'enhanced_analyzer' not in st.session_state:
                        st.error("❌ Анализатор не инициализирован. Выполните анализ изображений сначала.")
                    else:
                        analyzer = st.session_state.enhanced_analyzer
                        
                    # Формируем промпт для сравнительного анализа (работает для любого количества изображений)
                    comparison_prompt = f"""
Вы - опытный врач-диагност. Проведите детальный {'сравнительный' if len(results) > 1 else 'детальный'} анализ {len(results)} медицинских {'изображений' if len(results) > 1 else 'изображения'}.
Тип сравнения: {comparison_type}

Результаты анализа каждого изображения:
"""
                    
                    for i, result in enumerate(results, 1):
                        comparison_prompt += f"""
Изображение {i} ({result.filename}):
- Достоверность анализа: {result.confidence:.1%}
- Основные находки: {json.dumps(result.structured_findings, ensure_ascii=False, indent=2)}

"""
                    
                    if len(results) == 1:
                        comparison_prompt += f"""
Предоставьте ДЕТАЛЬНОЕ клиническое заключение, включающее:

1. ТЕХНИЧЕСКАЯ ОЦЕНКА:
   - Качество изображения
   - Технические параметры
   - Ограничения исследования

2. ДЕТАЛЬНЫЕ КЛИНИЧЕСКИЕ НАХОДКИ:
   - Все видимые анатомические структуры
   - Патологические изменения (если есть)
   - Локализация и выраженность изменений
   - Измерения и количественные параметры

3. ДИАГНОСТИЧЕСКАЯ ОЦЕНКА:
   - Основной диагноз с обоснованием
   - Дифференциальная диагностика
   - Вероятность диагноза

4. РЕКОМЕНДАЦИИ:
   - Срочные действия (если необходимы)
   - Дополнительные исследования
   - Тактика ведения пациента
   - План наблюдения

5. ПРОГНОЗ И РИСКИ:
   - Оценка тяжести состояния
   - Факторы риска
   - Прогноз

ВАЖНО: Дайте максимально подробный и детальный анализ. Не ограничивайтесь общими фразами - опишите все видимые структуры, изменения и дайте конкретные рекомендации.
"""
                    else:
                        comparison_prompt += f"""
Предоставьте детальное сравнительное заключение, включающее:

1. ТЕХНИЧЕСКОЕ СРАВНЕНИЕ:
   - Качество изображений
   - Сопоставимость исследований
   - Технические ограничения

2. КЛИНИЧЕСКИЕ НАХОДКИ:
   - Сравнение выявленных изменений между всеми изображениями
   - Динамика патологического процесса
   - Стабильные и изменившиеся параметры
   - Количественные изменения (если применимо)

3. ДИАГНОСТИЧЕСКАЯ ОЦЕНКА:
   - Подтверждение или изменение диагноза
   - Прогрессирование/регрессия заболевания
   - Эффективность лечения (если применимо)
   - Сравнение диагнозов по каждому изображению

4. РЕКОМЕНДАЦИИ:
   - Клинические выводы на основе сравнения
   - Необходимость дополнительных исследований
   - Тактика ведения пациента
   - План динамического наблюдения

5. ПРОГНОЗ:
   - Оценка динамики на основе всех изображений
   - Риски и перспективы
   - Прогноз течения заболевания

ВАЖНО: Сравните ВСЕ изображения детально. Опишите изменения между каждым изображением, динамику процесса, количественные и качественные изменения.
"""
                    
                    comparison_prompt += "\n\nОтвет структурируйте четко по разделам на русском языке."
                    
                    try:
                        # Используем streaming для сравнительного анализа
                        st.markdown("### 📋 Сравнительное заключение")
                        with st.spinner("🤖 Генерирую сравнительное заключение (Opus 4.5)..."):
                            text_generator = analyzer._send_ai_request_streaming(
                                comparison_prompt, 
                                images[0],  # Используем первое изображение как базовое
                                {"comparison_type": comparison_type, "images_count": len(results)}
                            )
                            
                            # Отображаем streaming результат
                            comparative_analysis_raw = st.write_stream(text_generator)
                            
                            # Преобразуем в строку если нужно (st.write_stream может вернуть разные типы)
                            comparative_analysis = ensure_string_for_download(comparative_analysis_raw)
                            
                            # Проверяем, что результат не пустой
                            if not comparative_analysis or len(comparative_analysis.strip()) == 0:
                                st.warning("⚠️ Получен пустой ответ. Пробую обычный режим...")
                                raise ValueError("Пустой ответ от streaming")
                        
                        # Сохраняем результат в session_state для возможности скачать
                        if comparative_analysis and len(comparative_analysis.strip()) > 0:
                            if 'comparative_analysis_result' not in st.session_state:
                                st.session_state.comparative_analysis_result = {}
                            st.session_state.comparative_analysis_result[f"{comparison_type}_{len(results)}"] = comparative_analysis
                            
                            # Возможность скачать заключение
                            st.markdown("---")
                            st.download_button(
                                label="💾 Скачать заключение",
                                data=comparative_analysis,
                                file_name=f"comparative_analysis_{comparison_type}_{len(results)}_images.txt",
                                mime="text/plain",
                                use_container_width=True,
                                key=f"download_streaming_conclusion_{comparison_type}_{len(results)}"
                            )
                        else:
                            st.error("❌ Не удалось получить заключение. Попробуйте еще раз.")
                        
                    except Exception as e:
                        st.error(f"❌ Ошибка генерации сравнительного анализа: {e}")
                        # Fallback на обычный режим
                        try:
                            st.warning("⚠️ Streaming недоступен, используем обычный режим...")
                            # _send_ai_request возвращает tuple: (content, model_name, tokens_used)
                            comparative_analysis_raw, model_name, tokens_used = analyzer._send_ai_request(
                                comparison_prompt, 
                                images[0],
                                {"comparison_type": comparison_type, "images_count": len(results)}
                            )
                            
                            # Преобразуем в строку
                            comparative_analysis = ensure_string_for_download(comparative_analysis_raw)
                            st.markdown(comparative_analysis)
                            
                            # Сохраняем результат
                            if 'comparative_analysis_result' not in st.session_state:
                                st.session_state.comparative_analysis_result = {}
                            st.session_state.comparative_analysis_result[f"{comparison_type}_{len(results)}"] = comparative_analysis
                            
                            st.download_button(
                                label="💾 Скачать заключение",
                                data=comparative_analysis,
                                file_name=f"comparative_analysis_{comparison_type}_{len(results)}_images.txt",
                                mime="text/plain",
                                use_container_width=True,
                                key=f"download_fallback_conclusion_{comparison_type}_{len(results)}"
                            )
                        except Exception as e2:
                            st.error(f"❌ Критическая ошибка: {e2}")
            else:
                st.error("❌ Не удалось обработать ни одного изображения")
    
    elif uploaded_files and len(uploaded_files) == 1:
        st.warning("⚠️ Для сравнительного анализа необходимо загрузить минимум 2 изображения")
    
    elif not uploaded_files:
        st.info("📤 Загрузите изображения для начала сравнительного анализа")


def show_ai_training_page():
    """Страница для обучения и калибровки ИИ"""
    st.header("🎓 Обучение и калибровка ИИ")
    
    st.warning("⚠️ Эта функция находится в разработке")
    
    st.info("""
    **Планируемый функционал:**
    
    🎯 **Калибровка моделей:**
    - Настройка уверенности для разных типов изображений
    - Обучение на специфических случаях вашей клиники
    
    📊 **Статистика производительности:**
    - Метрики точности по типам исследований  
    - Сравнение с экспертными заключениями
    
    🔧 **Настройка промптов:**
    - Кастомизация запросов к ИИ
    - Добавление специфических медицинских протоколов
    
    💾 **База знаний:**
    - Загрузка клинических рекомендаций
    - Интеграция с медицинскими стандартами
    """)
    
    # Временный функционал
    st.subheader("📈 Статистика текущего использования")
    
    if 'analysis_history' not in st.session_state:
        st.session_state.analysis_history = []
    
    if st.session_state.analysis_history:
        df = pd.DataFrame(st.session_state.analysis_history)
        st.dataframe(df)
    else:
        st.info("История анализов пуста")


def search_protocols_gemini(query: str, specialty: str = "") -> Dict:
    """
    Поиск актуальных медицинских протоколов через Gemini 2.5 Flash (бесплатно через OpenRouter)
    
    Args:
        query: Поисковый запрос
        specialty: Специальность для уточнения поиска
    
    Returns:
        Dict с результатами поиска и ссылками
    """
    try:
        # Получаем API ключ OpenRouter (уже используется в проекте)
        try:
            from config import OPENROUTER_API_KEY
            api_key = OPENROUTER_API_KEY
        except (ImportError, AttributeError):
            api_key = st.secrets.get("OPENROUTER_API_KEY") or st.secrets.get("api_keys", {}).get("OPENROUTER_API_KEY")
        
        if not api_key:
            return {
                "error": "API ключ OpenRouter не найден. Добавьте OPENROUTER_API_KEY в secrets.toml",
                "results": []
            }
        
        # Формируем промпт для поиска протоколов
        search_prompt = f"""Найди актуальные медицинские протоколы и клинические рекомендации по теме: {query}
Специальность: {specialty}

Предоставь структурированный ответ на русском языке:

1. НАЗВАНИЯ ПРОТОКОЛОВ/РЕКОМЕНДАЦИЙ:
   - Перечисли найденные протоколы с названиями

2. КРАТКОЕ ОПИСАНИЕ:
   - Для каждого протокола дай краткое описание

3. КЛЮЧЕВЫЕ ДИАГНОСТИЧЕСКИЕ КРИТЕРИИ:
   - Основные критерии диагностики
   - Лабораторные и инструментальные методы

4. ПРОТОКОЛЫ ЛЕЧЕНИЯ:
   - Основные принципы лечения
   - Рекомендуемые препараты и дозировки
   - Хирургические методы (если применимо)
   - Немедикаментозная терапия
   - Длительность лечения и критерии эффективности

5. ИСТОЧНИКИ:
   - Укажи названия источников (например: "Клинические рекомендации Минздрава РФ по...", "Рекомендации ESC по...")
   - Укажи общие направления поиска (например: "Искать в PubMed по ключевым словам: ...")
   - НЕ указывай конкретные PubMed ID, DOI или URL, если не уверен в их точности
   - Если знаешь точные проверенные ссылки - укажи их в формате:
     * PubMed: https://pubmed.ncbi.nlm.nih.gov/XXXXXXX (только если уверен)
     * DOI: https://doi.org/10.XXXX/XXXXX (только если уверен)
     * URL: полный рабочий адрес (только если уверен)

КРИТИЧЕСКИ ВАЖНО: 
- НЕ выдумывай и НЕ генерируй ссылки, которые не проверены
- Лучше указать только название источника и направление поиска, чем неверную ссылку
- Используй актуальные источники (2020-2024 годы)
- Если найдешь российские клинические рекомендации, укажи их в первую очередь."""
        
        # Используем Gemini 2.5 Flash через OpenRouter
        url = "https://openrouter.ai/api/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://github.com/vasiliys961/medical-assistant1",
            "X-Title": "Medical Protocol Search"
        }
        
        payload = {
            "model": "google/gemini-2.5-flash",
            "messages": [
                {
                    "role": "system",
                    "content": "Ты помощник врача. Ищешь актуальные медицинские протоколы и клинические рекомендации. Всегда предоставляй структурированную информацию с ссылками на источники."
                },
                {
                    "role": "user",
                    "content": search_prompt
                }
            ],
            "max_tokens": 3000,
            "temperature": 0.3
        }
        
        print(f"🔍 [GEMINI 2.5 FLASH] Ищу протоколы: {query} ({specialty})")
        start_time = time.time()
        response = requests.post(url, headers=headers, json=payload, timeout=60)
        latency = time.time() - start_time
        
        if response.status_code == 200:
            data = response.json()
            content = data.get("choices", [{}])[0].get("message", {}).get("content", "")
            tokens_used = data.get("usage", {}).get("total_tokens", 0)
            
            print(f"✅ [GEMINI 2.5 FLASH] Найдено протоколов. Токенов: {tokens_used}, Время: {latency:.2f}с")
            
            return {
                "success": True,
                "content": content,
                "tokens_used": tokens_used,
                "model": "Gemini 2.5 Flash"
            }
        elif response.status_code == 402:
            return {
                "error": "Недостаточно средств на OpenRouter. Пополните баланс.",
                "results": []
            }
        else:
            return {
                "error": f"Ошибка API: {response.status_code} - {response.text[:200]}",
                "results": []
            }
                
    except requests.exceptions.Timeout:
        return {
            "error": "Таймаут запроса. Попробуйте еще раз.",
            "results": []
        }
    except Exception as e:
        return {
            "error": f"Ошибка поиска: {str(e)}",
            "results": []
        }


def show_medical_protocols_page():
    """Страница медицинских протоколов и стандартов"""
    st.header("📚 Медицинские протоколы и стандарты")
    
    # Категории протоколов
    protocol_category = st.selectbox(
        "Выберите категорию",
        [
            "Кардиология",
            "Пульмонология",
            "Ревматология",
            "Гастроэнтерология",
            "Эндокринология",
            "Неврология",
            "Нефрология",
            "Гематология"
        ]
    )
    
    # Предопределенные протоколы
    protocols = {
        "Кардиология": {
            "Острый коронарный синдром (ОКС)": {
                "описание": "Протокол диагностики и ведения острого коронарного синдрома",
                "ключевые_точки": [
                    "Элевация ST > 1 мм в двух смежных отведениях (STEMI)",
                    "Депрессия ST > 0.5 мм или инверсия T (NSTEMI)",
                    "Тропонин положительный",
                    "ЭКГ в динамике каждые 15-30 минут",
                    "Антитромботическая терапия (аспирин, клопидогрел)",
                    "Статины, бета-блокаторы, ИАПФ"
                ],
                "код_мкб": ["I21", "I20.0"],
                "срочность": "экстренно"
            },
            "Артериальная гипертензия": {
                "описание": "Диагностика и лечение артериальной гипертензии",
                "ключевые_точки": [
                    "АД ≥140/90 мм рт.ст. при повторных измерениях",
                    "Суточное мониторирование АД (СМАД)",
                    "Оценка факторов риска и поражения органов-мишеней",
                    "Немедикаментозная терапия (диета, физическая активность)",
                    "Медикаментозная терапия: ИАПФ, БРА, БКК, диуретики"
                ],
                "код_мкб": ["I10", "I11", "I12", "I13", "I15"],
                "срочность": "планово"
            },
            "Хроническая сердечная недостаточность": {
                "описание": "Диагностика и лечение ХСН",
                "ключевые_точки": [
                    "Клинические признаки: одышка, отеки, утомляемость",
                    "ЭхоКГ: ФВ ЛЖ < 40% (систолическая дисфункция)",
                    "BNP/NT-proBNP для диагностики",
                    "Ингибиторы АПФ/БРА, бета-блокаторы, антагонисты альдостерона",
                    "Диуретики при застойных явлениях"
                ],
                "код_мкб": ["I50"],
                "срочность": "срочно"
            },
            "Фибрилляция предсердий": {
                "описание": "Диагностика и лечение фибрилляции предсердий",
                "ключевые_точки": [
                    "ЭКГ: отсутствие P волн, нерегулярный QRS",
                    "Оценка риска тромбоэмболий (CHA2DS2-VASc)",
                    "Антикоагулянтная терапия (варфарин, DOAC)",
                    "Контроль ЧСС (бета-блокаторы, верапамил, дигоксин)",
                    "Кардиоверсия при необходимости"
                ],
                "код_мкб": ["I48"],
                "срочность": "срочно"
            }
        },
        
        "Пульмонология": {
            "Внебольничная пневмония": {
                "описание": "Диагностика и лечение внебольничной пневмонии",
                "ключевые_точки": [
                    "Клинические признаки: кашель, лихорадка, одышка",
                    "Рентгенография ОГК: инфильтративные изменения",
                    "ОАК: лейкоцитоз, сдвиг влево",
                    "СРБ, прокальцитонин для оценки тяжести",
                    "Антибактериальная терапия (амоксициллин/клавуланат, макролиды)",
                    "Оценка по шкале CURB-65/CRB-65"
                ],
                "код_мкб": ["J13", "J14", "J15", "J16", "J18"],
                "срочность": "срочно"
            },
            "ХОБЛ (обострение)": {
                "описание": "Диагностика и лечение обострения ХОБЛ",
                "ключевые_точки": [
                    "Усиление одышки, кашля, увеличение объема мокроты",
                    "Спирометрия: ОФВ1 < 80% от должного",
                    "Рентгенография для исключения пневмонии, пневмоторакса",
                    "Бронходилататоры (сальбутамол, ипратропий)",
                    "Системные ГКС при тяжелом обострении",
                    "Антибиотики при признаках бактериальной инфекции"
                ],
                "код_мкб": ["J44"],
                "срочность": "срочно"
            },
            "Бронхиальная астма": {
                "описание": "Диагностика и лечение бронхиальной астмы",
                "ключевые_точки": [
                    "Обратимая бронхообструкция (спирография с бронхолитиком)",
                    "Клинические признаки: одышка, свистящие хрипы, кашель",
                    "Оценка контроля астмы (ACT, GINA)",
                    "Ингаляционные ГКС + бета-2-агонисты длительного действия",
                    "Короткодействующие бета-2-агонисты для купирования"
                ],
                "код_мкб": ["J45", "J46"],
                "срочность": "срочно"
            },
            "Тромбоэмболия легочной артерии (ТЭЛА)": {
                "описание": "Диагностика и лечение ТЭЛА",
                "ключевые_точки": [
                    "Клинические признаки: одышка, боль в груди, кровохарканье",
                    "D-димер для скрининга",
                    "КТ-ангиография легочных артерий (золотой стандарт)",
                    "Оценка по шкале Уэллса",
                    "Антикоагулянтная терапия (гепарин, DOAC)",
                    "Тромболизис при массивной ТЭЛА"
                ],
                "код_мкб": ["I26"],
                "срочность": "экстренно"
            }
        },
        
        "Ревматология": {
            "Ревматоидный артрит": {
                "описание": "Диагностика и лечение ревматоидного артрита",
                "ключевые_точки": [
                    "Клинические признаки: симметричный полиартрит, утренняя скованность",
                    "РФ, АЦЦП (антитела к цитруллинированному пептиду)",
                    "Рентгенография: эрозии, сужение суставных щелей",
                    "Оценка активности по DAS28",
                    "БМАРП (метотрексат, сульфасалазин, лефлуномид)",
                    "ГИБП при неэффективности БМАРП"
                ],
                "код_мкб": ["M05", "M06"],
                "срочность": "планово"
            },
            "Системная красная волчанка (СКВ)": {
                "описание": "Диагностика и лечение СКВ",
                "ключевые_точки": [
                    "Критерии ACR/EULAR 2019: кожные проявления, артрит, серозиты",
                    "АНА, анти-dsDNA, анти-Sm антитела",
                    "ОАК: лейкопения, тромбоцитопения, анемия",
                    "Оценка активности по SLEDAI",
                    "ГКС, гидроксихлорохин, иммуносупрессанты",
                    "Биологические препараты при тяжелом течении"
                ],
                "код_мкб": ["M32"],
                "срочность": "срочно"
            },
            "Остеоартроз": {
                "описание": "Диагностика и лечение остеоартроза",
                "ключевые_точки": [
                    "Клинические признаки: боль, крепитация, ограничение движений",
                    "Рентгенография: сужение суставной щели, остеофиты",
                    "Оценка по шкале WOMAC, VAS",
                    "НПВП для купирования боли",
                    "Хондропротекторы (хондроитин, глюкозамин)",
                    "Внутрисуставные инъекции ГКС, гиалуроновой кислоты"
                ],
                "код_мкб": ["M15", "M16", "M17", "M18", "M19"],
                "срочность": "планово"
            },
            "Подагра": {
                "описание": "Диагностика и лечение подагры",
                "ключевые_точки": [
                    "Острый артрит (часто I плюснефаланговый сустав)",
                    "Гиперурикемия (мочевая кислота > 420 мкмоль/л)",
                    "Идентификация кристаллов уратов в синовиальной жидкости",
                    "Купирование острого приступа: НПВП, колхицин, ГКС",
                    "Уратснижающая терапия (аллопуринол, фебуксостат)",
                    "Диета: ограничение пуринов, алкоголя"
                ],
                "код_мкб": ["M10"],
                "срочность": "срочно"
            }
        },
        
        "Гастроэнтерология": {
            "Язвенная болезнь желудка и ДПК": {
                "описание": "Диагностика и лечение язвенной болезни",
                "ключевые_точки": [
                    "Клинические признаки: боль в эпигастрии, изжога",
                    "ЭГДС: визуализация язвенного дефекта",
                    "Тест на H. pylori (дыхательный, кал, биопсия)",
                    "ИПП (омепразол, пантопразол) 4-8 недель",
                    "Эрадикация H. pylori (тройная/четверная схема)",
                    "Исключение НПВП при возможности"
                ],
                "код_мкб": ["K25", "K26", "K27", "K28"],
                "срочность": "срочно"
            },
            "Гастроэзофагеальная рефлюксная болезнь (ГЭРБ)": {
                "описание": "Диагностика и лечение ГЭРБ",
                "ключевые_точки": [
                    "Клинические признаки: изжога, регургитация",
                    "ЭГДС: эзофагит, пищевод Барретта",
                    "pH-метрия пищевода (при необходимости)",
                    "ИПП (омепразол, эзомепразол) 4-8 недель",
                    "Антациды, альгинаты для симптоматического лечения",
                    "Модификация образа жизни: диета, снижение веса"
                ],
                "код_мкб": ["K21"],
                "срочность": "планово"
            },
            "Воспалительные заболевания кишечника (ВЗК)": {
                "описание": "Диагностика и лечение ВЗК (болезнь Крона, язвенный колит)",
                "ключевые_точки": [
                    "Клинические признаки: диарея, боль в животе, кровь в стуле",
                    "Колоноскопия с биопсией",
                    "Кальпротектин в кале, СРБ, ОАК",
                    "Оценка активности по индексам (Mayo, CDAI)",
                    "5-АСК препараты, ГКС, иммуносупрессанты",
                    "Биологические препараты (инфликсимаб, адалимумаб)"
                ],
                "код_мкб": ["K50", "K51"],
                "срочность": "срочно"
            },
            "Острый панкреатит": {
                "описание": "Диагностика и лечение острого панкреатита",
                "ключевые_точки": [
                    "Клинические признаки: боль в эпигастрии, тошнота, рвота",
                    "Амилаза, липаза крови (повышение в 3 раза)",
                    "КТ брюшной полости: отек поджелудочной железы",
                    "Оценка тяжести по шкале Ranson, APACHE II",
                    "Голод, инфузионная терапия, обезболивание",
                    "Антибиотики при некротизирующем панкреатите"
                ],
                "код_мкб": ["K85"],
                "срочность": "экстренно"
            }
        },
        
        "Эндокринология": {
            "Сахарный диабет 2 типа": {
                "описание": "Диагностика и лечение СД 2 типа",
                "ключевые_точки": [
                    "Глюкоза плазмы натощак ≥ 7.0 ммоль/л или HbA1c ≥ 6.5%",
                    "ОГТТ при сомнительных результатах",
                    "Оценка осложнений: ретинопатия, нефропатия, нейропатия",
                    "Метформин как препарат первой линии",
                    "Ингибиторы ДПП-4, агонисты ГПП-1, ингибиторы SGLT2",
                    "Инсулинотерапия при неэффективности пероральных препаратов"
                ],
                "код_мкб": ["E11"],
                "срочность": "планово"
            },
            "Гипотиреоз": {
                "описание": "Диагностика и лечение гипотиреоза",
                "ключевые_точки": [
                    "Клинические признаки: утомляемость, увеличение веса, зябкость",
                    "ТТГ повышен, свободный T4 снижен",
                    "Антитела к ТПО, ТГ (при аутоиммунном тиреоидите)",
                    "Заместительная терапия левотироксином",
                    "Контроль ТТГ каждые 6-12 недель до нормализации",
                    "Поддержание ТТГ в целевом диапазоне (0.5-2.5 мЕд/л)"
                ],
                "код_мкб": ["E03", "E03.9"],
                "срочность": "планово"
            },
            "Гипертиреоз": {
                "описание": "Диагностика и лечение гипертиреоза",
                "ключевые_точки": [
                    "Клинические признаки: тахикардия, тремор, снижение веса",
                    "ТТГ снижен, свободный T4/T3 повышены",
                    "Антитела к рецептору ТТГ (при болезни Грейвса)",
                    "Тиреостатические препараты (тиамазол, пропилтиоурацил)",
                    "Бета-блокаторы для симптоматического лечения",
                    "Радиойодтерапия или хирургическое лечение при необходимости"
                ],
                "код_мкб": ["E05"],
                "срочность": "срочно"
            },
            "Метаболический синдром": {
                "описание": "Диагностика и лечение метаболического синдрома",
                "ключевые_точки": [
                    "Критерии: абдоминальное ожирение, АГ, дислипидемия, гипергликемия",
                    "Окружность талии: > 94 см (мужчины), > 80 см (женщины)",
                    "Липидограмма: ТГ ≥ 1.7, ЛПВП < 1.0 (мужчины) / < 1.3 (женщины)",
                    "Модификация образа жизни: диета, физическая активность",
                    "Статины при дислипидемии, антигипертензивные препараты",
                    "Метформин при нарушении толерантности к глюкозе"
                ],
                "код_мкб": ["E88.9"],
                "срочность": "планово"
            }
        },
        
        "Неврология": {
            "Острое нарушение мозгового кровообращения (ОНМК)": {
                "описание": "Диагностика и лечение ОНМК (инсульт)",
                "ключевые_точки": [
                    "Клинические признаки: внезапная слабость, нарушение речи, асимметрия лица",
                    "КТ головного мозга без контраста (исключение геморрагии)",
                    "Оценка по шкале NIHSS",
                    "Тромболитическая терапия в течение 4.5 часов от начала",
                    "Антиагреганты (аспирин), статины",
                    "Контроль АД, глюкозы, температуры"
                ],
                "код_мкб": ["I63", "I64", "I61", "I62"],
                "срочность": "экстренно"
            },
            "Эпилепсия": {
                "описание": "Диагностика и лечение эпилепсии",
                "ключевые_точки": [
                    "Клинические признаки: судорожные приступы, абсансы",
                    "ЭЭГ: эпилептиформная активность",
                    "МРТ головного мозга (исключение структурных изменений)",
                    "Антиэпилептические препараты (вальпроаты, карбамазепин, леветирацетам)",
                    "Монотерапия предпочтительна, комбинация при резистентности",
                    "Контроль уровня препаратов в крови"
                ],
                "код_мкб": ["G40"],
                "срочность": "срочно"
            },
            "Мигрень": {
                "описание": "Диагностика и лечение мигрени",
                "ключевые_точки": [
                    "Критерии: односторонняя пульсирующая головная боль, фото/фонофобия",
                    "Длительность 4-72 часа",
                    "Оценка частоты приступов",
                    "Купирование: НПВП, триптаны",
                    "Профилактика: бета-блокаторы, антиконвульсанты, антидепрессанты",
                    "Исключение триггеров: стресс, нерегулярный сон, продукты"
                ],
                "код_мкб": ["G43"],
                "срочность": "планово"
            },
            "Болезнь Паркинсона": {
                "описание": "Диагностика и лечение болезни Паркинсона",
                "ключевые_точки": [
                    "Клинические признаки: тремор покоя, ригидность, брадикинезия",
                    "Оценка по шкале UPDRS",
                    "МРТ для исключения вторичного паркинсонизма",
                    "Леводопа/карбидопа как основной препарат",
                    "Агонисты дофаминовых рецепторов, ингибиторы МАО-Б",
                    "Физическая реабилитация, логопедия"
                ],
                "код_мкб": ["G20"],
                "срочность": "планово"
            }
        },
        
        "Нефрология": {
            "Острое повреждение почек (ОПП)": {
                "описание": "Диагностика и лечение ОПП",
                "ключевые_точки": [
                    "Повышение креатинина ≥ 0.3 мг/дл или ≥ 1.5x от исходного",
                    "Снижение диуреза < 0.5 мл/кг/ч в течение 6 часов",
                    "Оценка по критериям KDIGO",
                    "Исключение преренальных и постренальных причин",
                    "Коррекция гиповолемии, отмена нефротоксичных препаратов",
                    "Заместительная почечная терапия при необходимости"
                ],
                "код_мкб": ["N17"],
                "срочность": "экстренно"
            },
            "Хроническая болезнь почек (ХБП)": {
                "описание": "Диагностика и лечение ХБП",
                "ключевые_точки": [
                    "Снижение СКФ < 60 мл/мин/1.73 м² или альбуминурия ≥ 30 мг/г",
                    "Стадирование по СКФ (G1-G5) и альбуминурии (A1-A3)",
                    "Оценка причины: СД, АГ, гломерулонефрит",
                    "Контроль АД (цель < 130/80), ИАПФ/БРА",
                    "Коррекция анемии, фосфорно-кальциевого обмена",
                    "Подготовка к заместительной почечной терапии при СКФ < 15"
                ],
                "код_мкб": ["N18"],
                "срочность": "планово"
            },
            "Гломерулонефрит": {
                "описание": "Диагностика и лечение гломерулонефрита",
                "ключевые_точки": [
                    "Протеинурия, гематурия, отеки, АГ",
                    "Биопсия почки для верификации типа",
                    "ОАМ: протеинурия, эритроцитурия, цилиндры",
                    "Иммунологическое обследование: АНЦА, анти-GBM, комплемент",
                    "ГКС, цитостатики при необходимости",
                    "ИАПФ/БРА для снижения протеинурии"
                ],
                "код_мкб": ["N00", "N01", "N02", "N03", "N04", "N05"],
                "срочность": "срочно"
            },
            "Инфекция мочевыводящих путей (ИМП)": {
                "описание": "Диагностика и лечение ИМП",
                "ключевые_точки": [
                    "Клинические признаки: дизурия, частое мочеиспускание, боль",
                    "ОАМ: лейкоцитурия, бактериурия",
                    "Посев мочи с определением чувствительности",
                    "Антибактериальная терапия (фосфомицин, нитрофурантоин, фторхинолоны)",
                    "Длительность терапии: 3-7 дней (неосложненная), 7-14 дней (осложненная)",
                    "Профилактика рецидивов при частых ИМП"
                ],
                "код_мкб": ["N30", "N39.0"],
                "срочность": "срочно"
            }
        },
        
        "Гематология": {
            "Железодефицитная анемия": {
                "описание": "Диагностика и лечение железодефицитной анемии",
                "ключевые_точки": [
                    "ОАК: снижение Hb, MCV, MCH, MCHC",
                    "Сывороточное железо снижено, ферритин < 15 нг/мл",
                    "ОЖСС повышена, коэффициент насыщения трансферрина < 15%",
                    "Поиск причины кровопотери (ЖКТ, гинекология)",
                    "Препараты железа (перорально или внутривенно)",
                    "Контроль ОАК через 2-4 недели, продолжение до нормализации ферритина"
                ],
                "код_мкб": ["D50"],
                "срочность": "планово"
            },
            "В12-дефицитная анемия": {
                "описание": "Диагностика и лечение В12-дефицитной анемии",
                "ключевые_точки": [
                    "ОАК: макроцитарная анемия, гиперсегментация нейтрофилов",
                    "В12 сыворотки < 200 пг/мл",
                    "Гомоцистеин, метилмалоновая кислота повышены",
                    "Антитела к внутреннему фактору Касла, париетальным клеткам",
                    "Заместительная терапия цианокобаламином (внутримышечно)",
                    "Поиск причины: атрофический гастрит, резекция желудка, веганство"
                ],
                "код_мкб": ["D51"],
                "срочность": "планово"
            },
            "Тромбоцитопения": {
                "описание": "Диагностика и лечение тромбоцитопении",
                "ключевые_точки": [
                    "Тромбоциты < 150×10⁹/л",
                    "Оценка клинических проявлений: кровоточивость, петехии",
                    "Исключение вторичных причин: лекарства, инфекции, ДВС",
                    "Иммунная тромбоцитопеническая пурпура (ИТП): антитела к тромбоцитам",
                    "ГКС, внутривенный иммуноглобулин при ИТП",
                    "Спленэктомия при рефрактерной ИТП"
                ],
                "код_мкб": ["D69"],
                "срочность": "срочно"
            },
            "Венозные тромбозы": {
                "описание": "Диагностика и лечение венозных тромбозов",
                "ключевые_точки": [
                    "Клинические признаки: отек, боль, покраснение конечности",
                    "D-димер для скрининга",
                    "УЗДГ вен (компрессионный тест)",
                    "Оценка по шкале Уэллса",
                    "Антикоагулянтная терапия (гепарин, DOAC)",
                    "Длительность: 3 месяца (провоцированный), длительно (непровоцированный)"
                ],
                "код_мкб": ["I80", "I82"],
                "срочность": "экстренно"
            }
        }
    }
    
    # Поиск актуальных протоколов через Gemini 2.5 Flash (бесплатно)
    st.markdown("---")
    st.subheader("🔍 Поиск актуальных протоколов")
    st.info("💡 Поиск выполняется через Gemini 2.5 Flash (бесплатно через OpenRouter)")
    
    search_query = st.text_input(
        "Введите запрос для поиска протоколов",
        placeholder=f"Например: протокол лечения {protocol_category.lower()}",
        key="protocol_search_query"
    )
    
    if st.button("🔍 Найти актуальные протоколы", use_container_width=True, type="primary", key="search_protocols"):
        if search_query:
            with st.spinner("🔍 Ищу актуальные протоколы через Gemini 2.5 Flash (бесплатно)..."):
                result = search_protocols_gemini(search_query, protocol_category)
                
                if result.get("success"):
                    # Gemini возвращает структурированный текст
                    st.markdown("### 📋 Найденные протоколы")
                    st.markdown(result.get("content", ""))
                    
                    if result.get("tokens_used"):
                        st.caption(f"📊 Использовано токенов: {result.get('tokens_used')}")
                    
                    st.caption(f"🤖 Поиск выполнен через {result.get('model', 'Gemini 2.5 Flash')} (бесплатно)")
                else:
                    error_msg = result.get("error", "Неизвестная ошибка")
                    st.error(f"❌ {error_msg}")
                    if "API ключ" in error_msg:
                        st.info("💡 Для использования поиска добавьте OPENROUTER_API_KEY в `.streamlit/secrets.toml`")
        else:
            st.warning("⚠️ Введите запрос для поиска")
    
    st.markdown("---")
    
    # Существующие протоколы
    st.subheader("📚 Сохраненные протоколы")
    
    if protocol_category in protocols:
        selected_protocols = protocols[protocol_category]
        
        for protocol_name, protocol_data in selected_protocols.items():
            with st.expander(f"📋 {protocol_name}", expanded=False):
                
                st.markdown(f"**Описание:** {protocol_data['описание']}")
                
                urgency_color = {
                    "экстренно": "red",
                    "срочно": "orange", 
                    "планово": "green"
                }.get(protocol_data['срочность'], "gray")
                
                st.markdown(f"**Срочность:** :{urgency_color}[{protocol_data['срочность']}]")
                
                st.markdown("**Ключевые диагностические критерии:**")
                for point in protocol_data['ключевые_точки']:
                    st.markdown(f"• {point}")
                
                st.markdown(f"**Коды МКБ-10:** {', '.join(protocol_data['код_мкб'])}")
                
                # Кнопка для поиска актуальных протоколов по конкретному протоколу
                if st.button(f"🔍 Найти актуальные протоколы: {protocol_name}", key=f"search_{protocol_name}"):
                    with st.spinner("Ищу актуальные протоколы через Gemini 2.5 Flash..."):
                        search_result = search_protocols_gemini(
                            f"{protocol_name} {protocol_data['описание']}", 
                            protocol_category
                        )
                        if search_result.get("success"):
                            st.markdown("### 📋 Актуальные протоколы:")
                            st.markdown(search_result.get("content", ""))
                            if search_result.get("tokens_used"):
                                st.caption(f"📊 Использовано токенов: {search_result.get('tokens_used')}")
                        else:
                            st.error(f"Ошибка поиска: {search_result.get('error', 'Неизвестная ошибка')}")
    
    # Кастомные протоколы
    st.subheader("➕ Добавить собственный протокол")
    
    with st.form("custom_protocol"):
        custom_name = st.text_input("Название протокола")
        custom_description = st.text_area("Описание")
        custom_criteria = st.text_area("Диагностические критерии (по одному в строке)")
        custom_icd = st.text_input("Коды МКБ-10 (через запятую)")
        custom_urgency = st.selectbox("Уровень срочности", ["планово", "срочно", "экстренно"])
        
        if st.form_submit_button("💾 Сохранить протокол"):
            if custom_name and custom_description:
                st.success(f"✅ Протокол '{custom_name}' сохранен!")
            else:
                st.error("❌ Заполните обязательные поля")


# Интеграция с основным приложением
def integrate_with_main_app():
    """Функция для интеграции с основным Streamlit приложением"""
    
    enhanced_pages = [
        "🔬 Расширенный ИИ-анализ",
        "📊 Сравнительный анализ", 
        "🎓 Обучение ИИ",
        "📚 Медицинские протоколы"
    ]
    
    page_functions = {
        "🔬 Расширенный ИИ-анализ": show_enhanced_analysis_page,
        "📊 Сравнительный анализ": show_comparative_analysis_page,
        "🎓 Обучение ИИ": show_ai_training_page,
        "📚 Медицинские протоколы": show_medical_protocols_page
    }
    
    return enhanced_pages, page_functions


if __name__ == "__main__":
    st.set_page_config(page_title="Enhanced Medical AI", layout="wide")
    
    test_page = st.sidebar.selectbox("Выберите тестовую страницу", [
        "Расширенный анализ",
        "Сравнительный анализ", 
        "Обучение ИИ",
        "Медицинские протоколы"
    ])
    
    if test_page == "Расширенный анализ":
        show_enhanced_analysis_page()
    elif test_page == "Сравнительный анализ":
        show_comparative_analysis_page()
    elif test_page == "Обучение ИИ":
        show_ai_training_page()
    elif test_page == "Медицинские протоколы":
        show_medical_protocols_page()