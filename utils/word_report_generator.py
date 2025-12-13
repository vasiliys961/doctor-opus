"""
Утилита для генерации стандартизованных Word-документов с заключениями медицинских анализов
"""
import datetime
from typing import Optional
import io

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def generate_word_report(
    analysis_type: str,
    conclusion_text: str,
    patient_info: Optional[str] = None,
    metadata: Optional[dict] = None,
    timestamp: Optional[str] = None
) -> Optional[bytes]:
    """
    Генерирует стандартизованный Word-документ с заключением медицинского анализа
    
    Args:
        analysis_type: Тип анализа (ECG, XRAY, MRI, CT, ULTRASOUND, DERMATOSCOPY, LAB, GENETIC)
        conclusion_text: Текст заключения
        patient_info: Информация о пациенте (опционально)
        metadata: Дополнительные метаданные (опционально)
        timestamp: Временная метка анализа (опционально)
    
    Returns:
        bytes: Байты Word-документа или None если docx недоступен
    """
    if not DOCX_AVAILABLE:
        return None
    
    # Создаем документ
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Заголовок документа
    title = doc.add_heading('МЕДИЦИНСКОЕ ЗАКЛЮЧЕНИЕ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Тип анализа
    analysis_type_names = {
        'ECG': 'Электрокардиография (ЭКГ)',
        'XRAY': 'Рентгенография',
        'MRI': 'Магнитно-резонансная томография (МРТ)',
        'CT': 'Компьютерная томография (КТ)',
        'ULTRASOUND': 'Ультразвуковое исследование (УЗИ)',
        'DERMATOSCOPY': 'Дерматоскопия',
        'LAB': 'Лабораторные исследования',
        'GENETIC': 'Генетический анализ'
    }
    
    analysis_name = analysis_type_names.get(analysis_type.upper(), analysis_type)
    doc.add_heading(analysis_name, level=1)
    
    # Информация о пациенте
    if patient_info:
        p = doc.add_paragraph()
        p.add_run('Информация о пациенте: ').bold = True
        p.add_run(patient_info)
    
    # Временная метка
    if timestamp:
        p = doc.add_paragraph()
        p.add_run('Дата и время анализа: ').bold = True
        p.add_run(timestamp)
    else:
        p = doc.add_paragraph()
        p.add_run('Дата и время анализа: ').bold = True
        p.add_run(datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
    
    # Разделитель
    doc.add_paragraph('─' * 50)
    
    # Заключение
    doc.add_heading('ЗАКЛЮЧЕНИЕ', level=1)
    
    # Разбиваем текст на параграфы для лучшего форматирования
    paragraphs = conclusion_text.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            # Обрабатываем заголовки (начинаются с цифр или **)
            if para_text.strip().startswith('**') or para_text.strip()[0].isdigit():
                # Убираем markdown форматирование
                clean_text = para_text.replace('**', '').strip()
                if clean_text:
                    doc.add_heading(clean_text, level=2)
            else:
                # Обычный параграф
                para = doc.add_paragraph()
                # Убираем markdown форматирование
                clean_text = para_text.replace('**', '').replace('*', '').strip()
                if clean_text:
                    para.add_run(clean_text)
    
    # Метаданные (если есть)
    if metadata:
        doc.add_paragraph('─' * 50)
        doc.add_heading('Дополнительная информация', level=2)
        for key, value in metadata.items():
            p = doc.add_paragraph()
            p.add_run(f'{key}: ').bold = True
            p.add_run(str(value))
    
    # Подпись
    doc.add_paragraph('─' * 50)
    p = doc.add_paragraph()
    p.add_run('Документ сгенерирован автоматически системой медицинского анализа с использованием ИИ')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.italic = True
    
    # Сохраняем в байты
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()


def get_word_report_filename(analysis_type: str, timestamp: Optional[str] = None) -> str:
    """
    Генерирует имя файла для Word-документа
    
    Args:
        analysis_type: Тип анализа
        timestamp: Временная метка (опционально)
    
    Returns:
        str: Имя файла
    """
    type_names = {
        'ECG': 'ECG',
        'XRAY': 'XRay',
        'MRI': 'MRI',
        'CT': 'CT',
        'ULTRASOUND': 'Ultrasound',
        'DERMATOSCOPY': 'Dermatoscopy',
        'LAB': 'Lab',
        'GENETIC': 'Genetic'
    }
    
    type_name = type_names.get(analysis_type.upper(), analysis_type)
    
    if timestamp:
        # Очищаем timestamp для имени файла
        clean_timestamp = timestamp.replace(' ', '_').replace(':', '-')
        return f"{type_name}_report_{clean_timestamp}.docx"
    else:
        current_time = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")
        return f"{type_name}_report_{current_time}.docx"
